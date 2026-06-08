"""
pinakia_service.py — Greek court schedule (πινάκια) extraction and case matching.

Supports: PDF (Claude vision), DOCX (python-docx), XLSX (openpyxl), images.
"""

import base64
import io
import json
import re
import logging
from datetime import datetime, timezone
from typing import Optional

import anthropic

logger = logging.getLogger("nomos_one.pinakia")

_EXTRACT_SYSTEM = (
    "Είσαι ειδικός στην ανάλυση ελληνικών δικαστικών πινακείων. "
    "Εξάγεις δομημένα δεδομένα και επιστρέφεις ΜΟΝΟ έγκυρο JSON χωρίς markdown."
)

_EXTRACT_PROMPT = """Αναλύσε αυτό το πινάκειο δικαστηρίου και εξάγαγε:

Επέστρεψε ΜΟΝΟ έγκυρο JSON (χωρίς markdown, χωρίς ```):
{
  "court_name": "πλήρης ονομασία δικαστηρίου",
  "hearing_date": "YYYY-MM-DD",
  "hearings": [
    {
      "aa": 1,
      "case_number": "ΑΒΜ 123/2025 ή ΑΓ 456/2024 κτλ",
      "parties": ["ΟΝΟΜΑ ΕΠΩΝΥΜΟ", "ΟΝΟΜΑ2 ΕΠΩΝΥΜΟ2"],
      "time": "09:00",
      "notes": ""
    }
  ]
}

Κανόνες:
- court_name: πλήρες όνομα δικαστηρίου (π.χ. "Τριμελές Πλημμελειοδικείο Σύρου")
- hearing_date: μορφή YYYY-MM-DD μόνο
- parties: ΟΛΑ τα ονόματα κατηγορουμένων/εναγομένων/διαδίκων — κεφαλαία
- case_number: ΑΒΜ, ΑΓ, ΑΠ, Γ, ΠΠ κτλ — ό,τι υπάρχει
- time: "HH:MM" ή null αν δεν αναφέρεται
- notes: αναβολή, διαγραφή, ή κενό string
- Αν κοινή ώρα για όλες, βάλε την στην 1η και null στις υπόλοιπες
"""


def _extract_text_docx(file_bytes: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def _extract_text_xlsx(file_bytes: bytes) -> str:
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
    parts = []
    for ws in wb.worksheets:
        for row in ws.iter_rows(values_only=True):
            row_text = " | ".join(str(v).strip() for v in row if v is not None and str(v).strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def _parse_json(text: str) -> dict:
    cleaned = re.sub(r',\s*([}\]])', r'\1', text)
    start = cleaned.find('{')
    if start != -1:
        try:
            obj, _ = json.JSONDecoder().raw_decode(cleaned, start)
            return obj
        except json.JSONDecodeError:
            pass
    logger.warning(f"pinakia _parse_json failed, raw: {text[:300]}")
    return {"court_name": "Άγνωστο Δικαστήριο", "hearing_date": None, "hearings": []}


def extract_pinakio(api_key: str, file_bytes: bytes, media_type: str,
                    model: str = "claude-sonnet-4-6") -> dict:
    """
    Extract court schedule from a document file.
    Returns: {court_name, hearing_date, hearings: [{aa, case_number, parties, time, notes}]}
    """
    client = anthropic.Anthropic(api_key=api_key)

    DOCX_TYPES = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    )
    XLSX_TYPES = (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/vnd.ms-excel",
    )

    if media_type in DOCX_TYPES:
        text = _extract_text_docx(file_bytes)
        messages = [{"role": "user", "content": f"{_EXTRACT_PROMPT}\n\n---\n{text[:10000]}"}]

    elif media_type in XLSX_TYPES:
        text = _extract_text_xlsx(file_bytes)
        messages = [{"role": "user", "content": f"{_EXTRACT_PROMPT}\n\n---\n{text[:10000]}"}]

    elif media_type == "application/pdf":
        b64 = base64.standard_b64encode(file_bytes).decode()
        file_block = {
            "type": "document",
            "source": {"type": "base64", "media_type": "application/pdf", "data": b64}
        }
        messages = [{"role": "user", "content": [file_block, {"type": "text", "text": _EXTRACT_PROMPT}]}]

    elif media_type in ("text/plain", "text/csv"):
        text = file_bytes.decode("utf-8", errors="replace")
        messages = [{"role": "user", "content": f"{_EXTRACT_PROMPT}\n\n---\n{text[:10000]}"}]

    else:
        # Image (JPEG, PNG, WEBP)
        b64 = base64.standard_b64encode(file_bytes).decode()
        img_type = media_type if media_type in {"image/jpeg", "image/png", "image/webp"} else "image/jpeg"
        file_block = {"type": "image", "source": {"type": "base64", "media_type": img_type, "data": b64}}
        messages = [{"role": "user", "content": [file_block, {"type": "text", "text": _EXTRACT_PROMPT}]}]

    resp = client.messages.create(
        model=model,
        max_tokens=8000,
        system=_EXTRACT_SYSTEM,
        messages=messages,
    )
    raw = resp.content[0].text
    logger.info(f"Pinakio extract raw[:300]: {raw[:300]}")
    result = _parse_json(raw)
    result["_tokens"] = resp.usage.input_tokens + resp.usage.output_tokens
    return result


async def match_hearings(db, hearings: list) -> list:
    """
    Cross-reference each hearing against open cases in MongoDB.
    Adds matched_case_id, matched_case_title, matched_case_number to each hearing.
    """
    CLOSED = {"closed_won", "closed_lost"}
    result = []

    for h in hearings:
        matched_case_id = None
        matched_case_title = None
        matched_case_number = None

        # 1. Try case_number match
        case_num = (h.get("case_number") or "").strip()
        if case_num and len(case_num) >= 3:
            case = await db.cases.find_one({
                "$or": [
                    {"case_number": {"$regex": re.escape(case_num), "$options": "i"}},
                    {"title": {"$regex": re.escape(case_num), "$options": "i"}},
                    {"description": {"$regex": re.escape(case_num), "$options": "i"}},
                ]
            })
            if case:
                matched_case_id = str(case["_id"])
                matched_case_title = case.get("title", "—")
                matched_case_number = case.get("case_number")

        # 2. Try party name match (open cases only)
        if not matched_case_id:
            for party in h.get("parties", []):
                if not party or len(party.strip()) < 3:
                    continue
                # Try full name, then last name (first word of the name)
                parts = party.strip().split()
                candidates = [party.strip()]
                if len(parts) >= 2:
                    candidates.append(parts[0])   # last name (Greek convention: ΕΠΩΝΥΜΟ ΟΝΟΜΑ)

                for term in candidates:
                    if len(term) < 3:
                        continue
                    case = await db.cases.find_one({
                        "$or": [
                            {"title": {"$regex": re.escape(term), "$options": "i"}},
                            {"client_name": {"$regex": re.escape(term), "$options": "i"}},
                            {"description": {"$regex": re.escape(term), "$options": "i"}},
                        ],
                        "status": {"$nin": list(CLOSED)}
                    })
                    if case:
                        matched_case_id = str(case["_id"])
                        matched_case_title = case.get("title", "—")
                        matched_case_number = case.get("case_number")
                        break
                if matched_case_id:
                    break

        result.append({
            **h,
            "matched_case_id": matched_case_id,
            "matched_case_title": matched_case_title,
            "matched_case_number": matched_case_number,
        })

    return result


def is_pinakio_document(caption: str, filename: str, media_type: str) -> bool:
    """Heuristic: is this document likely a court schedule?"""
    DOCX_XLSX = {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/vnd.ms-excel",
    }
    if media_type in DOCX_XLSX:
        return True

    keywords = ["πινακ", "πινάκ", "δικασιμ", "δικάσιμ", "ακροατηριο", "ακροατήριο",
                "pinakio", "dikas", "πλημ", "εφετ", "αρειο"]
    text = f"{caption} {filename}".lower()
    return any(kw in text for kw in keywords)
