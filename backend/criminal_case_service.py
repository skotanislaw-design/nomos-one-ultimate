"""Criminal Case Management — AI, health-score and export service.

Combines CCMSS ai_service.py + health_score.py + export_service.py but uses
Nomos One's existing anthropic.AsyncAnthropic client instead of emergentintegrations.
"""

# ── Standard library ──────────────────────────────────────────────────────────
import io
import os
from datetime import date, datetime, timezone
from typing import Optional

# ── Third-party ───────────────────────────────────────────────────────────────
import anthropic

# ── PDF / DOCX ────────────────────────────────────────────────────────────────
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib.colors import HexColor
from reportlab.lib.enums import TA_LEFT
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

from docx import Document as DocxDocument
from docx.shared import RGBColor

# ─────────────────────────────────────────────────────────────────────────────
# Anthropic client
# ─────────────────────────────────────────────────────────────────────────────
ANTHROPIC_API_KEY = os.getenv("ANTHROPIC_API_KEY", "")
MODEL = "claude-sonnet-4-6"

# ─────────────────────────────────────────────────────────────────────────────
# Safety preambles (verbatim from CCMSS ai_service.py)
# ─────────────────────────────────────────────────────────────────────────────
SAFETY_PREAMBLE_EL = """ΣΗΜΑΝΤΙΚΟ: Είσαι βοηθός σύνταξης για δικηγόρο που χειρίζεται ΠΟΙΝΙΚΗ ΥΠΟΘΕΣΗ. Η έξοδός σου είναι ΠΑΝΤΑ "Draft for lawyer review" και ΟΧΙ τελική νομική συμβουλή.

Αυστηροί κανόνες:
1. ΔΕΝ εφευρίσκεις γεγονότα, ονόματα, ημερομηνίες ή πρόσωπα που δεν αναφέρονται στα δεδομένα.
2. ΔΕΝ επικαλείσαι συγκεκριμένα άρθρα νόμου, νομολογία ή νομικές διατάξεις εκτός αν αναφέρονται ρητά στα δεδομένα.
3. Ξεχωρίζεις πάντα τα γεγονότα από τις υποθέσεις. Χρησιμοποιείς φράσεις όπως "με βάση τα διαθέσιμα στοιχεία", "πιθανό", "χρειάζεται επιβεβαίωση", "ασαφές".
4. Σημειώνεις ρητά τις ελλείψεις και τις αβεβαιότητες.
5. Η γλώσσα εξόδου είναι ΕΛΛΗΝΙΚΑ εκτός αν ζητηθεί διαφορετικά.
6. Δομή: καθαρές επικεφαλίδες, bullet lists, σύντομες παράγραφοι.
7. Στο τέλος προσθέτεις τη γραμμή: "— Draft for lawyer review. This document is not final legal advice and must be reviewed by a qualified lawyer before use."
"""

SAFETY_PREAMBLE_EN = """IMPORTANT: You are a drafting assistant for a lawyer handling a CRIMINAL CASE. Your output is ALWAYS a "Draft for lawyer review" and is NOT final legal advice.

Strict rules:
1. NEVER invent facts, names, dates, or persons not present in the provided data.
2. NEVER cite specific law articles, case law, or legal provisions unless they are explicitly present in the provided data.
3. Always separate facts from assumptions. Use phrases like "based on available information", "likely", "requires confirmation", "unclear".
4. Explicitly flag missing information and uncertainties.
5. Output language is ENGLISH unless otherwise requested.
6. Structure: clear headings, bullet lists, short paragraphs.
7. End with the line: "— Draft for lawyer review. This document is not final legal advice and must be reviewed by a qualified lawyer before use."
"""

# ─────────────────────────────────────────────────────────────────────────────
# Prompts (verbatim from CCMSS ai_service.py)
# ─────────────────────────────────────────────────────────────────────────────
PROMPTS = {
    "case_summary": {
        "el": "Παρήγαγε δύο εκδοχές περίληψης της ποινικής υπόθεσης: (1) σύντομη περίληψη 3-5 γραμμές, (2) αναλυτική περίληψη με τίτλους: Πλαίσιο, Κρίσιμα Γεγονότα, Αβέβαια Σημεία, Άμεσες Ενέργειες.",
        "en": "Produce two versions of the criminal case summary: (1) short summary 3-5 lines, (2) detailed summary with headings: Context, Critical Facts, Uncertain Points, Immediate Actions.",
    },
    "chronology": {
        "el": "Δημιούργησε χρονολόγιο γεγονότων σε μορφή πίνακα markdown με στήλες: Ημερομηνία | Γεγονός | Πηγή | Αξιοπιστία (confirmed/alleged/unclear) | Εκκρεμότητα. Χρησιμοποίησε μόνο γεγονότα που υπάρχουν στα δεδομένα. Σημείωσε κενά.",
        "en": "Create an events chronology as a markdown table with columns: Date | Event | Source | Reliability (confirmed/alleged/unclear) | Pending. Use only events present in the data. Note any gaps.",
    },
    "missing_documents": {
        "el": "Δημιούργησε checklist ΕΛΛΕΙΠΟΝΤΩΝ εγγράφων για την υπόθεση. Έλεγξε αν λείπουν: κλήση/κατηγορητήριο, δικογραφία, κατάθεση πελάτη, καταθέσεις μαρτύρων, φωτογραφίες, βίντεο, ιατρικές γνωματεύσεις, έκθεση αστυνομίας, πραγματογνωμοσύνη, αποδεικτικά επικοινωνίας, εξουσιοδότηση, πληρεξούσιο, ταυτότητα/διαβατήριο πελάτη. Σε κάθε στοιχείο: [ ] Όνομα · Σημασία (χαμηλή/μέτρια/υψηλή) · Λόγος αναγκαιότητας.",
        "en": "Create a MISSING DOCUMENTS checklist. Check for: summons/indictment, case file, client statement, witness statements, photos, video, medical reports, police report, expert opinion, communications evidence, authorization, power of attorney, client ID/passport. For each item: [ ] Name · Importance (low/medium/high) · Reason it is needed.",
    },
    "client_questions": {
        "el": "Παρήγαγε 12-20 στοχευμένες ερωτήσεις προς τον πελάτη ώστε να συμπληρωθούν τα κενά πληροφορίας. Ομαδοποίησέ τις σε θεματικές ενότητες. Καμία ερώτηση να μην προϋποθέτει γεγονός που δεν αναφέρεται.",
        "en": "Produce 12-20 targeted questions for the client to fill information gaps. Group them by topic. No question may assume a fact not stated.",
    },
    "witness_questions": {
        "el": "Παρήγαγε ερωτήσεις προς κάθε μάρτυρα της υπόθεσης ομαδοποιημένες ανά μάρτυρα. Αν δεν υπάρχει μάρτυρας στα δεδομένα, δημιούργησε γενικό template ερωτήσεων μάρτυρα.",
        "en": "Produce questions for each witness, grouped per witness. If no witness is in the data, create a generic witness questioning template.",
    },
    "risk_analysis": {
        "el": "Παρήγαγε ανάλυση κινδύνου χωρισμένη σε: 1) Δυνατά Σημεία, 2) Αδύναμα Σημεία, 3) Ασάφειες, 4) Αντιφάσεις, 5) Επείγοντα Νομικά/Πραγματικά Ζητήματα. Κάθε σημείο σε bullet με σύντομη αιτιολογία.",
        "en": "Produce a risk analysis split into: 1) Strengths, 2) Weaknesses, 3) Ambiguities, 4) Contradictions, 5) Urgent Legal/Factual Issues. Each item as a bullet with short reasoning.",
    },
    "court_brief": {
        "el": "Παρήγαγε One-Page Court Preparation Brief με τα εξής τμήματα: Τίτλος Υπόθεσης, Βασικά Πραγματικά Περιστατικά (5-7 bullets), Κρίσιμα Έγγραφα, Μάρτυρες, Θέματα Προς Προσοχή, Πιθανές Ερωτήσεις Δικαστηρίου, Checklist Ημέρας Δικαστηρίου.",
        "en": "Produce a One-Page Court Preparation Brief with sections: Case Title, Key Facts (5-7 bullets), Critical Documents, Witnesses, Points of Caution, Likely Court Questions, Court Day Checklist.",
    },
    "client_email": {
        "el": "Συνέταξε draft email προς τον πελάτη που τον ενημερώνει για την πορεία της υπόθεσης και ζητά εκκρεμότητες. ΜΗ συμπεριλάβεις οριστική νομική συμβουλή. Τόνος: επαγγελματικός, σαφής. Δομή: Θέμα, Χαιρετισμός, Σώμα, Επόμενα Βήματα, Καταληκτική φράση. ΣΗΜΕΙΩΣΗ: Αυτό είναι draft και ΔΕΝ θα αποσταλεί αυτόματα.",
        "en": "Draft an email to the client updating on case progress and requesting pending items. DO NOT include final legal advice. Tone: professional, clear. Structure: Subject, Salutation, Body, Next Steps, Closing. NOTE: This is a draft and will NOT be sent automatically.",
    },
    "internal_memo": {
        "el": "Συνέταξε εσωτερικό memo προς την ομάδα του γραφείου για την υπόθεση. Δομή: Υπόθεση, Στάδιο, Κρίσιμες Ημερομηνίες, Εκκρεμότητες, Προτεινόμενες Ενέργειες, Σημεία Προσοχής.",
        "en": "Draft an internal memo to the firm's team on the case. Structure: Case, Stage, Key Dates, Pending Items, Suggested Actions, Points of Caution.",
    },
    "defence_strategy": {
        "el": "Παρήγαγε ΠΡΟΣΧΕΔΙΟ γενικής στρατηγικής υπεράσπισης βάσει διαθέσιμων στοιχείων. Δομή: Κεντρικά Επιχειρήματα, Αποδεικτική Βάση Κάθε Επιχειρήματος, Αντίθετα Σενάρια, Στοιχεία που Απαιτούνται. ΧΩΡΙΣ νομικές διατάξεις ή νομολογία.",
        "en": "Produce a DRAFT general defence strategy based on available data. Structure: Core Arguments, Evidentiary Basis per Argument, Counter-Scenarios, Required Items. NO law articles or case law.",
    },
    "prosecution_support": {
        "el": "Παρήγαγε ΠΡΟΣΧΕΔΙΟ ανάλυσης υποστήριξης κατηγορίας/πολιτικής αγωγής. Δομή: Κεντρικά Σημεία, Αποδεικτικά Στοιχεία, Κενά, Επόμενα Βήματα.",
        "en": "Produce a DRAFT prosecution-support / civil-claim analysis. Structure: Key Points, Evidence, Gaps, Next Steps.",
    },
    "legal_issues": {
        "el": "Εντόπισε γενικά νομικά ζητήματα/θεματικές που προκύπτουν από τα δεδομένα (ΧΩΡΙΣ να αναφέρεις άρθρα νόμου). Σε κάθε ένα: Τίτλος, Γεγονότα που Στηρίζουν, Στοιχεία που Λείπουν, Επίπεδο Κινδύνου.",
        "en": "Identify general legal issues/themes arising from the data (WITHOUT citing law articles). For each: Title, Supporting Facts, Missing Facts, Risk Level.",
    },
}

# ─────────────────────────────────────────────────────────────────────────────
# build_case_context (verbatim from CCMSS ai_service.py)
# ─────────────────────────────────────────────────────────────────────────────
def build_case_context(case_doc: dict, parties: list, events: list,
                       documents: list, evidence: list, issues: list) -> str:
    lines = ["=== CASE DATA ==="]
    lines.append(f"Title: {case_doc.get('case_title')}")
    lines.append(f"Matter type: {case_doc.get('matter_type')}")
    lines.append(f"Client: {case_doc.get('client_name')} ({case_doc.get('client_role')})")
    lines.append(f"Client email: {case_doc.get('client_email') or '—'}")
    lines.append(f"Client phone: {case_doc.get('client_phone') or '—'}")
    lines.append(f"Opposing party: {case_doc.get('opposing_party') or '—'}")
    lines.append(f"Court: {case_doc.get('court') or '—'}")
    lines.append(f"Hearing date: {case_doc.get('hearing_date') or '—'}")
    lines.append(f"Urgency: {case_doc.get('urgency_level')}")
    lines.append(f"Status: {case_doc.get('status')}")
    lines.append("")
    lines.append("Short description:")
    lines.append(case_doc.get("short_description") or "—")
    lines.append("")

    lines.append("=== PARTIES ===")
    if parties:
        for p in parties:
            lines.append(f"- {p.get('name')} [{p.get('role')}] — {p.get('contact_details') or ''} — notes: {p.get('notes') or ''}")
    else:
        lines.append("(none)")
    lines.append("")

    lines.append("=== TIMELINE ===")
    if events:
        for e in events:
            lines.append(f"- {e.get('event_date')} {e.get('event_time') or ''} | {e.get('event_description')} | source: {e.get('source') or '—'} | reliability: {e.get('confidence_level')}")
    else:
        lines.append("(none)")
    lines.append("")

    lines.append("=== DOCUMENTS ===")
    if documents:
        for d in documents:
            summary = (d.get("summary") or d.get("extracted_text") or "")[:300]
            lines.append(f"- [{d.get('category')}] {d.get('file_name')} — importance: {d.get('importance_level')} — summary: {summary}")
    else:
        lines.append("(none)")
    lines.append("")

    lines.append("=== EVIDENCE ===")
    if evidence:
        for ev in evidence:
            lines.append(f"- {ev.get('title')} | supports: {ev.get('supports')} | reliability: {ev.get('reliability')} | {ev.get('description') or ''}")
    else:
        lines.append("(none)")
    lines.append("")

    lines.append("=== LEGAL ISSUES (lawyer-flagged) ===")
    if issues:
        for li in issues:
            lines.append(f"- {li.get('issue_title')} | risk: {li.get('risk_level')} | supporting: {li.get('facts_supporting') or '—'} | missing: {li.get('missing_facts') or '—'}")
    else:
        lines.append("(none)")
    return "\n".join(lines)


# ─────────────────────────────────────────────────────────────────────────────
# generate_output — rewritten to use anthropic.AsyncAnthropic
# ─────────────────────────────────────────────────────────────────────────────
async def generate_output(
    output_type: str,
    language: str,
    case_context: str,
    extra: Optional[str] = None,
) -> str:
    lang = language if language in ("el", "en") else "el"
    preamble = SAFETY_PREAMBLE_EL if lang == "el" else SAFETY_PREAMBLE_EN
    prompt_map = PROMPTS.get(output_type)
    if not prompt_map:
        raise ValueError(f"Unknown output_type: {output_type}")
    task_prompt = prompt_map[lang]

    system_msg = preamble + "\n\n" + task_prompt

    user_text = case_context
    if extra:
        user_text += f"\n\nADDITIONAL CONTEXT:\n{extra}"

    client = anthropic.AsyncAnthropic(api_key=ANTHROPIC_API_KEY)
    response = await client.messages.create(
        model=MODEL,
        max_tokens=4096,
        system=system_msg,
        messages=[{"role": "user", "content": user_text}],
    )
    return response.content[0].text


# ─────────────────────────────────────────────────────────────────────────────
# summarize_document — rewritten to use anthropic.AsyncAnthropic
# ─────────────────────────────────────────────────────────────────────────────
async def summarize_document(file_name: str, text: str, language: str = "el") -> str:
    """Generate a short summary of an uploaded document."""
    if not text or not text.strip():
        return ""
    lang = language if language in ("el", "en") else "el"
    if lang == "el":
        sys = "Είσαι βοηθός σύνταξης νομικού γραφείου. Παρήγαγε σύντομη περίληψη 3-6 γραμμών για το έγγραφο, χωρίς να εφευρίσκεις γεγονότα. Σημείωσε αν λείπει κάτι κρίσιμο. Στο τέλος βάλε γραμμή: '— Draft for lawyer review.'"
    else:
        sys = "You are a legal-office drafting assistant. Produce a 3-6 line short summary of the document without inventing facts. Note if anything critical is missing. End with line: '— Draft for lawyer review.'"

    truncated = text[:6000]
    client = anthropic.AsyncAnthropic(api_key=ANTHROPIC_API_KEY)
    response = await client.messages.create(
        model=MODEL,
        max_tokens=512,
        system=sys,
        messages=[{"role": "user", "content": f"Filename: {file_name}\n\nContent:\n{truncated}"}],
    )
    return response.content[0].text


# ─────────────────────────────────────────────────────────────────────────────
# health_score (verbatim from CCMSS health_score.py)
# ─────────────────────────────────────────────────────────────────────────────
def _parse_date(s: Optional[str]) -> Optional[date]:
    if not s:
        return None
    try:
        return date.fromisoformat(s[:10])
    except Exception:
        return None


def days_until(due: Optional[str]) -> Optional[int]:
    d = _parse_date(due)
    if not d:
        return None
    today = datetime.now(timezone.utc).date()
    return (d - today).days


def compute_health(
    case_doc: dict,
    document_count: int,
    overdue_tasks_count: int,
    unapproved_critical_outputs: int,
    missing_critical_docs: bool,
) -> dict:
    """Return a dict {level, reasons, hearing_days_left, escalated}.

    Rules:
      RED: hearing in <=3 days AND (unapproved court_brief OR missing critical docs
           OR overdue tasks); OR overdue tasks with hearing in <=7 days.
      YELLOW: missing docs OR pending approvals OR hearing 4-14 days away
              OR overdue tasks.
      GREEN: otherwise.
    """
    reasons = []
    hd = days_until(case_doc.get("hearing_date"))

    is_critical = False
    if hd is not None and hd <= 3:
        critical_conditions = []
        if unapproved_critical_outputs > 0:
            critical_conditions.append("Μη εγκεκριμένο court brief")
        if missing_critical_docs:
            critical_conditions.append("Λείπουν κρίσιμα έγγραφα")
        if overdue_tasks_count > 0:
            critical_conditions.append(f"{overdue_tasks_count} καθυστερημένα tasks")
        if critical_conditions:
            is_critical = True
            reasons.append(f"Δικάσιμος σε {hd} ημέρες · " + ", ".join(critical_conditions))

    if is_critical:
        return {
            "level": "red",
            "reasons": reasons,
            "hearing_days_left": hd,
            "escalated": True,
        }

    yellow = False
    if hd is not None and hd <= 14 and hd >= 0:
        yellow = True
        reasons.append(f"Δικάσιμος σε {hd} ημέρες")
    if overdue_tasks_count > 0:
        yellow = True
        reasons.append(f"{overdue_tasks_count} καθυστερημένα tasks")
    if document_count == 0:
        yellow = True
        reasons.append("Δεν έχουν ανέβει έγγραφα")
    if unapproved_critical_outputs > 0:
        yellow = True
        reasons.append(f"{unapproved_critical_outputs} μη εγκεκριμένα κρίσιμα drafts")

    if yellow:
        # If hearing also overdue (negative days) AND case not closed, escalate to red
        if hd is not None and hd < 0 and case_doc.get("status") != "closed":
            return {
                "level": "red", "reasons": reasons + ["Δικάσιμος έχει παρέλθει"],
                "hearing_days_left": hd, "escalated": True,
            }
        return {
            "level": "yellow", "reasons": reasons,
            "hearing_days_left": hd, "escalated": False,
        }

    return {
        "level": "green", "reasons": ["Όλα εντάξει"],
        "hearing_days_left": hd, "escalated": False,
    }


# ─────────────────────────────────────────────────────────────────────────────
# export_service (verbatim from CCMSS export_service.py)
# ─────────────────────────────────────────────────────────────────────────────
DISCLAIMER = (
    "Draft for lawyer review. This document is not final legal advice and must "
    "be reviewed by a qualified lawyer before use."
)

_FONT_NAME = "Helvetica"
_FONT_BOLD = "Helvetica-Bold"
try:
    pdfmetrics.registerFont(TTFont("DejaVuSans", "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"))
    pdfmetrics.registerFont(TTFont("DejaVuSans-Bold", "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"))
    _FONT_NAME = "DejaVuSans"
    _FONT_BOLD = "DejaVuSans-Bold"
except Exception:
    pass


def _styles():
    base = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "Title", parent=base["Title"], fontName=_FONT_BOLD,
            fontSize=18, textColor=HexColor("#0f172a"), alignment=TA_LEFT,
            spaceAfter=12,
        ),
        "h2": ParagraphStyle(
            "H2", parent=base["Heading2"], fontName=_FONT_BOLD,
            fontSize=12, textColor=HexColor("#0f172a"), spaceAfter=6,
        ),
        "body": ParagraphStyle(
            "Body", parent=base["BodyText"], fontName=_FONT_NAME,
            fontSize=10, textColor=HexColor("#1e293b"), leading=14, spaceAfter=6,
        ),
        "meta": ParagraphStyle(
            "Meta", parent=base["BodyText"], fontName=_FONT_NAME,
            fontSize=9, textColor=HexColor("#475569"), spaceAfter=2,
        ),
        "disclaimer": ParagraphStyle(
            "Disclaimer", parent=base["BodyText"], fontName=_FONT_BOLD,
            fontSize=9, textColor=HexColor("#b45309"), spaceBefore=12, leading=12,
        ),
        "draft_tag": ParagraphStyle(
            "Draft", parent=base["BodyText"], fontName=_FONT_BOLD,
            fontSize=10, textColor=HexColor("#1d4ed8"), spaceAfter=8,
        ),
    }


def _para_safe(text: str) -> str:
    """Escape minimal HTML for ReportLab Paragraph."""
    if text is None:
        return ""
    return (
        text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
    )


def render_pdf(
    case_title: str,
    output_title: str,
    content: str,
    status: str = "draft",
    language: str = "el",
) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=2*cm, rightMargin=2*cm,
        topMargin=2*cm, bottomMargin=2*cm,
        title=output_title,
    )
    styles = _styles()
    story = []

    story.append(Paragraph(_para_safe(output_title), styles["title"]))
    story.append(Paragraph(f"<b>Υπόθεση:</b> {_para_safe(case_title)}", styles["meta"]))
    story.append(Paragraph(
        f"<b>Δημιουργήθηκε:</b> {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}",
        styles["meta"],
    ))
    story.append(Paragraph(f"<b>Status:</b> {status.upper()}", styles["meta"]))
    story.append(Spacer(1, 0.3*cm))

    story.append(Paragraph("DRAFT FOR LAWYER REVIEW", styles["draft_tag"]))
    story.append(Spacer(1, 0.2*cm))

    for block in (content or "").split("\n\n"):
        block = block.strip()
        if not block:
            continue
        safe = _para_safe(block).replace("\n", "<br/>")
        story.append(Paragraph(safe, styles["body"]))
        story.append(Spacer(1, 0.15*cm))

    story.append(Spacer(1, 0.5*cm))
    story.append(Paragraph(DISCLAIMER, styles["disclaimer"]))

    doc.build(story)
    return buf.getvalue()


def render_docx(
    case_title: str,
    output_title: str,
    content: str,
    status: str = "draft",
    language: str = "el",
) -> bytes:
    docx = DocxDocument()

    t = docx.add_heading(output_title, level=1)
    for run in t.runs:
        run.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)

    p = docx.add_paragraph()
    p.add_run(f"Υπόθεση: {case_title}").bold = True
    docx.add_paragraph(f"Δημιουργήθηκε: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}")
    docx.add_paragraph(f"Status: {status.upper()}")

    tag = docx.add_paragraph()
    r = tag.add_run("DRAFT FOR LAWYER REVIEW")
    r.bold = True
    r.font.color.rgb = RGBColor(0x1d, 0x4e, 0xd8)

    docx.add_paragraph("")

    for block in (content or "").split("\n\n"):
        block = block.strip()
        if not block:
            continue
        docx.add_paragraph(block)

    docx.add_paragraph("")
    disc = docx.add_paragraph()
    r2 = disc.add_run(DISCLAIMER)
    r2.bold = True
    r2.font.color.rgb = RGBColor(0xb4, 0x53, 0x09)

    buf = io.BytesIO()
    docx.save(buf)
    return buf.getvalue()
