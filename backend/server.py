"""
Nomos One - Law Firm Management System
Production-ready FastAPI backend — Phase 1 Security Hardened
"""

from fastapi import FastAPI, HTTPException, Depends, status, UploadFile, File, Form, Query, Request, Header
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, StreamingResponse
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from motor.motor_asyncio import AsyncIOMotorClient
from pydantic import BaseModel, Field
from typing import Optional, List, Any, Dict
from datetime import datetime, timedelta
from bson import ObjectId
from bson.errors import InvalidId
from collections import defaultdict
import asyncio
import jwt
import bcrypt
import os
import shutil
import json
import zipfile
import logging
import uuid
import time
import secrets
import re
from enum import Enum
from pathlib import Path

# ── PWA & Mobile Services ─────────────────────────────────────────────────────
from device_service import get_device_service
from push_service import get_push_service
from ai_service import extract_document as ai_extract_document, intake_analyze

# ── Two-Factor Authentication (Phase 1.6) ──────────────────────────────────────
from two_factor_service import TwoFactorService, OTPSessionType
from encryption_service import EncryptionService
from email_service import send_otp_email

# ── WebSocket & Real-time Messaging (Phase 1.7) ────────────────────────────────
from websocket_service import get_websocket_manager
from websocket_routes import router as ws_router, set_jwt_secret
from telegram_intake_service import router as telegram_router, register_webhook, SESSION_TTL_HOURS
from email_intake_service import email_intake_loop
from pinakia_service import extract_pinakio, match_hearings as match_pinakio_hearings
from lexis_service import SPECIALISTS, route_question
from nomologia_service import retrieve_relevant_nomologia
from solon_service import search_solon, format_solon_for_prompt

# ── Logging ──────────────────────────────────────────────────────────────────
logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(name)s: %(message)s")
logger = logging.getLogger("nomos_one")

# ── Config ────────────────────────────────────────────────────────────────────
MONGO_URI = os.getenv("MONGO_URI", "mongodb://localhost:27017")
DB_NAME = os.getenv("DB_NAME", "nomos_one")
JWT_SECRET = os.getenv("JWT_SECRET", "")
JWT_ALGORITHM = "HS256"
JWT_EXPIRY_HOURS = int(os.getenv("JWT_EXPIRY_HOURS", "8"))
DOCUMENT_STORAGE_PATH = Path(os.getenv("DOCUMENT_STORAGE_PATH", "/data/documents"))
ALLOWED_ORIGINS = os.getenv("CORS_ORIGINS", "http://localhost:3000").split(",")
MAX_FILE_SIZE_MB = int(os.getenv("MAX_FILE_SIZE_MB", "50"))
ALLOWED_FILE_TYPES = {
    "application/pdf", "application/msword",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "image/jpeg", "image/png", "image/tiff",
    "text/plain", "application/zip"
}

# ── Security Config ───────────────────────────────────────────────────────────
MIN_PASSWORD_LENGTH = int(os.getenv("MIN_PASSWORD_LENGTH", "8"))
MAX_LOGIN_ATTEMPTS = int(os.getenv("MAX_LOGIN_ATTEMPTS", "5"))
LOGIN_LOCKOUT_MINUTES = int(os.getenv("LOGIN_LOCKOUT_MINUTES", "15"))
STAGNANT_DAYS = int(os.getenv("STAGNANT_DAYS", "30"))
ANTHROPIC_API_KEY = os.getenv("ANTHROPIC_API_KEY", "")
# Model routing: Haiku for bulk extraction (cheap), Sonnet for legal reasoning & chat (quality)
MODEL_EXTRACTION = "claude-haiku-4-5-20251001"
MODEL_CHAT = "claude-sonnet-4-6"
# API key for server-to-server intake calls (Google Apps Script, etc.)
INTAKE_API_KEY = os.getenv("INTAKE_API_KEY", "")

# Validate JWT_SECRET at startup
if not JWT_SECRET or JWT_SECRET.startswith("CHANGE"):
    JWT_SECRET = secrets.token_hex(32)
    logger.warning(
        "JWT_SECRET not set or is default! Generated ephemeral secret. "
        "Set JWT_SECRET in .env for persistent sessions across restarts."
    )

DOCUMENT_STORAGE_PATH.mkdir(parents=True, exist_ok=True)

# ── Rate Limiter ─────────────────────────────────────────────────────────────
class LoginRateLimiter:
    def __init__(self):
        self._attempts: dict = defaultdict(list)
        self._lockouts: dict = {}

    def is_locked(self, key: str) -> bool:
        if key in self._lockouts:
            if time.time() < self._lockouts[key]:
                return True
            del self._lockouts[key]
            self._attempts.pop(key, None)
        return False

    def record_attempt(self, key: str) -> None:
        now = time.time()
        window = now - (LOGIN_LOCKOUT_MINUTES * 60)
        self._attempts[key] = [t for t in self._attempts[key] if t > window]
        self._attempts[key].append(now)
        if len(self._attempts[key]) >= MAX_LOGIN_ATTEMPTS:
            self._lockouts[key] = now + (LOGIN_LOCKOUT_MINUTES * 60)
            logger.warning(f"Account locked: {key}")

    def clear(self, key: str) -> None:
        self._attempts.pop(key, None)
        self._lockouts.pop(key, None)

    def remaining_lockout(self, key: str) -> int:
        if key in self._lockouts:
            return max(0, int((self._lockouts[key] - time.time()) / 60) + 1)
        return 0

rate_limiter = LoginRateLimiter()

# ── Password Policy ──────────────────────────────────────────────────────────
def validate_password(pw: str):
    if len(pw) < MIN_PASSWORD_LENGTH:
        return False, f"Ο κωδικός πρέπει να έχει τουλάχιστον {MIN_PASSWORD_LENGTH} χαρακτήρες"
    if not re.search(r"[A-Za-z]", pw):
        return False, "Ο κωδικός πρέπει να περιέχει τουλάχιστον ένα γράμμα"
    if not re.search(r"\d", pw):
        return False, "Ο κωδικός πρέπει να περιέχει τουλάχιστον ένα ψηφίο"
    return True, ""

# ── Input Validation ─────────────────────────────────────────────────────────
def validate_phone(phone):
    if not phone: return True
    cleaned = re.sub(r"[\s\-\+\(\)]", "", phone)
    return bool(re.match(r"^(30)?\d{10}$", cleaned) or re.match(r"^\d{10,15}$", cleaned))

def validate_tax_id(tax_id):
    if not tax_id: return True
    return bool(re.match(r"^\d{9}$", tax_id.strip()))

def sanitize_string(s):
    return s.strip() if s else s

# ── App ───────────────────────────────────────────────────────────────────────
app = FastAPI(title="Nomos One API", version="1.1.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Intake channel routers (registered at module level so routes are available immediately)
app.include_router(telegram_router)

# ── DB ────────────────────────────────────────────────────────────────────────
client: AsyncIOMotorClient = None
db = None

# ── 2FA Services ──────────────────────────────────────────────────────────────
encryption_service: EncryptionService = None
two_factor_service: TwoFactorService = None

@app.on_event("startup")
async def startup():
    global client, db, encryption_service, two_factor_service
    client = AsyncIOMotorClient(MONGO_URI)
    db = client[DB_NAME]
    await db.users.create_index("email", unique=True)
    await db.cases.create_index([("case_number", 1)], unique=True)
    await db.cases.create_index([("assigned_lawyer_id", 1)])
    await db.audit_logs.create_index([("timestamp", -1)])
    await db.documents.create_index([("case_id", 1), ("uploaded_at", -1)])
    await db.notes.create_index([("case_id", 1), ("created_at", -1)])
    await db.financials.create_index([("case_id", 1)])
    # Phase 2: Deadlines
    await db.deadlines.create_index([("date", 1)])
    await db.deadlines.create_index([("case_id", 1)])
    # Phase 3: Case parties
    await db.case_parties.create_index([("case_id", 1)])
    # Phase 4: Invoices
    await db.invoices.create_index([("invoice_number", 1)], unique=True)
    await db.expenses_log.create_index([("case_id", 1)])
    await db.invoices.create_index([("case_id", 1)])
    await db.expenses_log.create_index([("case_id", 1)])
    # Pinakia (court schedules)
    await db.pinakia.create_index([("hearing_date", -1)])
    await db.pinakia.create_index([("court_name", 1)])
    # ── Criminal Cases Module ──
    await db.cc_cases.create_index([("id", 1)], unique=True)
    await db.cc_cases.create_index([("created_at", -1)])
    await db.cc_parties.create_index([("case_id", 1)])
    await db.cc_events.create_index([("case_id", 1), ("event_date", 1)])
    await db.cc_documents.create_index([("case_id", 1)])
    await db.cc_evidence.create_index([("case_id", 1)])
    await db.cc_issues.create_index([("case_id", 1)])
    await db.cc_tasks.create_index([("case_id", 1)])
    await db.cc_outputs.create_index([("case_id", 1), ("created_at", -1)])
    # Initialize atomic counter
    existing = await db.counters.find_one({"_id": "case_number"})
    if not existing:
        year = datetime.utcnow().year
        count = await db.cases.count_documents({"case_number": {"$regex": f"^{year}-"}})
        await db.counters.update_one({"_id": "case_number"}, {"$set": {"year": year, "seq": count}}, upsert=True)

    # ── Phase 1.6: Two-Factor Authentication ─────────────────────────────────
    # Initialize TwoFactorService
    encryption_service = EncryptionService()
    two_factor_service = TwoFactorService(db, encryption_service)
    logger.info("TwoFactorService initialized")

    # Create OTP sessions collection with TTL index for auto-cleanup
    await db.otp_sessions.create_index([("expires_at", 1)], expireAfterSeconds=0)
    logger.info("OTP sessions collection initialized with TTL index")

    # ── Phase 1.7: WebSocket Real-time Messaging ───────────────────────────────
    # Initialize WebSocket manager
    ws_manager = get_websocket_manager()
    logger.info(f"WebSocket manager initialized: {ws_manager}")

    # Configure WebSocket routes with JWT secret
    set_jwt_secret(JWT_SECRET)

    # Include WebSocket router
    app.include_router(ws_router)
    logger.info("WebSocket routes registered")

    logger.info("Database connected and indexes created")
    await seed_default_admin()

    # ── Intake channels: Telegram & Email ────────────────────────────────────
    await db.telegram_sessions.create_index(
        [("updated_at", 1)], expireAfterSeconds=SESSION_TTL_HOURS * 3600
    )
    await db.pending_intakes.create_index([("submitted_at", -1)])

    # ── Portal access logs: indexes + 7-year TTL retention ───────────────────
    PORTAL_LOG_TTL_SECONDS = 7 * 365 * 24 * 3600  # 7 years
    await db.portal_access_logs.create_index([("case_id", 1), ("timestamp", -1)])
    await db.portal_access_logs.create_index([("code_hash", 1)])
    await db.portal_access_logs.create_index([("timestamp", -1)], expireAfterSeconds=PORTAL_LOG_TTL_SECONDS)
    base_url = os.getenv("BASE_URL", "https://nomos.skotanislaw.gr")
    await register_webhook(base_url)
    asyncio.create_task(email_intake_loop(db, ANTHROPIC_API_KEY, MODEL_EXTRACTION))

    # Pre-warm Solon browser in background so first LEXIS query is fast
    async def _prewarm_solon():
        try:
            from solon_service import _ensure_browser
            await _ensure_browser()
            logger.info("Solon browser pre-warmed")
        except Exception as e:
            logger.warning(f"Solon pre-warm failed: {e}")
    asyncio.create_task(_prewarm_solon())

@app.on_event("shutdown")
async def shutdown():
    client.close()


# ── Enums ─────────────────────────────────────────────────────────────────────
class UserRole(str, Enum):
    ADMIN = "administrator"
    LAWYER = "lawyer"
    SECRETARY = "secretary"

class CaseStatus(str, Enum):
    ACTIVE = "active"
    PENDING = "pending"
    CLOSED = "closed"
    ARCHIVED = "archived"

class PaymentStatus(str, Enum):
    PENDING = "pending"
    PAID = "paid"
    PARTIAL = "partial"
    OVERDUE = "overdue"

# ── Helpers ───────────────────────────────────────────────────────────────────
def serialize(doc) -> dict:
    if doc is None: return None
    doc = dict(doc)
    if "_id" in doc: doc["id"] = str(doc.pop("_id"))
    for k, v in doc.items():
        if isinstance(v, ObjectId): doc[k] = str(v)
    return doc

def make_id(val: str) -> ObjectId:
    try: return ObjectId(val)
    except: raise HTTPException(400, "Invalid ID format")

def hash_password(pw): return bcrypt.hashpw(pw.encode(), bcrypt.gensalt()).decode()
def verify_password(pw, hashed): return bcrypt.checkpw(pw.encode(), hashed.encode())

def create_token(payload):
    data = payload.copy()
    data["exp"] = datetime.utcnow() + timedelta(hours=JWT_EXPIRY_HOURS)
    return jwt.encode(data, JWT_SECRET, algorithm=JWT_ALGORITHM)

def decode_token(token):
    try: return jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALGORITHM])
    except jwt.ExpiredSignatureError: raise HTTPException(401, "Token expired")
    except jwt.InvalidTokenError: raise HTTPException(401, "Invalid token")

security = HTTPBearer()

async def get_current_user(creds: HTTPAuthorizationCredentials = Depends(security)):
    payload = decode_token(creds.credentials)
    user = await db.users.find_one({"_id": make_id(payload["sub"])})
    if not user: raise HTTPException(401, "User not found")
    if not user.get("is_active", True): raise HTTPException(403, "Ο λογαριασμός είναι ανενεργός")
    return serialize(user)

def require_role(*roles):
    async def checker(user=Depends(get_current_user)):
        if user["role"] not in [r.value for r in roles]: raise HTTPException(403, "Insufficient permissions")
        return user
    return checker

async def audit(action, user_id, resource, resource_id=None, details=None):
    await db.audit_logs.insert_one({
        "action": action, "user_id": user_id, "resource": resource,
        "resource_id": resource_id, "details": details or {}, "timestamp": datetime.utcnow()
    })

async def get_user_name(user_id):
    try:
        u = await db.users.find_one({"_id": make_id(user_id)}, {"name": 1})
        return u["name"] if u else "Άγνωστος"
    except: return "Άγνωστος"

async def get_client_name(client_id):
    try:
        c = await db.clients.find_one({"_id": make_id(client_id)}, {"name": 1})
        return c["name"] if c else "Άγνωστος"
    except: return "Άγνωστος"

async def _check_payment_gate(case_id: str, action_label: str):
    """Block any legal action if the case has an outstanding balance. Notify secretariat + handler."""
    if not case_id:
        return
    case = await db.cases.find_one({"_id": make_id(case_id)}, {"title": 1, "case_number": 1, "assigned_lawyer_id": 1, "client_id": 1})
    if not case:
        return
    invoices = await db.invoices.find({"case_id": case_id}).to_list(None)
    total_invoiced = sum(float(i.get("total_payable", i.get("total", 0))) for i in invoices)
    total_paid = sum(float(i.get("amount_paid", 0)) for i in invoices)
    outstanding = round(total_invoiced - total_paid, 2)
    if outstanding <= 0:
        return

    case_title = case.get("title", "—")
    case_number = case.get("case_number", "—")

    # Collect recipients: all secretariat/admin + case handler
    recipients = await db.users.find(
        {"role": {"$in": [UserRole.SECRETARY.value, UserRole.ADMIN.value]}, "is_active": True},
        {"_id": 1, "name": 1}
    ).to_list(None)
    lawyer_id = case.get("assigned_lawyer_id")
    if lawyer_id:
        lawyer = await db.users.find_one({"_id": make_id(lawyer_id)}, {"_id": 1, "name": 1})
        if lawyer and not any(str(r["_id"]) == str(lawyer["_id"]) for r in recipients):
            recipients.append(lawyer)

    # Store in-app notifications
    for r in recipients:
        await db.notifications.insert_one({
            "user_id": str(r["_id"]),
            "type": "payment_required",
            "title": "⚠️ Απαιτείται εξόφληση αμοιβής",
            "message": (
                f"Η ενέργεια '{action_label}' στην υπόθεση «{case_title}» (#{case_number}) "
                f"δεν επιτράπηκε. Εκκρεμεί υπόλοιπο αμοιβής: {outstanding:,.2f}€. "
                f"Ουδεμία ενέργεια επιτρέπεται πριν εξοφληθεί η αμοιβή στο σύνολό της."
            ),
            "case_id": case_id,
            "outstanding_balance": outstanding,
            "created_at": datetime.utcnow(),
            "read": False,
        })

    # Telegram alert to all authorized chats
    try:
        import httpx as _httpx_gate
        bot_token = os.getenv("TELEGRAM_BOT_TOKEN", "")
        allowed_raw = os.getenv("TELEGRAM_ALLOWED_CHATS", "")
        if bot_token and allowed_raw.strip():
            chat_ids = [int(x.strip()) for x in allowed_raw.split(",") if x.strip().lstrip("-").isdigit()]
            text = (
                f"🚫 *Αποκλεισμός ενέργειας — Εκκρεμής Αμοιβή*\n\n"
                f"Υπόθεση: *{case_title}* (#{case_number})\n"
                f"Ενέργεια: {action_label}\n"
                f"Εκκρεμές υπόλοιπο: *{outstanding:,.2f}€*\n\n"
                f"_Ουδεμία ενέργεια επιτρέπεται πριν εξοφληθεί η αμοιβή στο σύνολό της._"
            )
            api_url = f"https://api.telegram.org/bot{bot_token}/sendMessage"
            async with _httpx_gate.AsyncClient(timeout=8) as client:
                for chat_id in chat_ids:
                    try:
                        await client.post(api_url, json={"chat_id": chat_id, "text": text, "parse_mode": "Markdown"})
                    except Exception:
                        pass
    except Exception as tg_err:
        logger.warning(f"Payment gate Telegram notify failed: {tg_err}")

    raise HTTPException(
        status_code=402,
        detail={
            "code": "PAYMENT_REQUIRED",
            "message": (
                f"Εκκρεμεί αμοιβή {outstanding:,.2f}€ στην υπόθεση «{case_title}». "
                f"Ουδεμία ενέργεια επιτρέπεται πριν εξοφληθεί η αμοιβή στο σύνολό της."
            ),
            "outstanding_balance": outstanding,
            "case_id": case_id,
        }
    )

# ── Seed ──────────────────────────────────────────────────────────────────────
ADMIN_EMAIL = os.getenv("ADMIN_EMAIL", "christos@skotanislaw.com")
ADMIN_NAME = os.getenv("ADMIN_NAME", "Χρήστος Σκοτάνης")
ADMIN_INITIAL_PASSWORD = os.getenv("ADMIN_INITIAL_PASSWORD", "")

async def seed_default_admin():
    existing = await db.users.find_one({"role": UserRole.ADMIN.value})
    if not existing:
        if ADMIN_INITIAL_PASSWORD:
            pw = ADMIN_INITIAL_PASSWORD
            must_change = False  # User set their own password via env
        else:
            pw = secrets.token_urlsafe(16)
            must_change = True
        await db.users.insert_one({
            "email": ADMIN_EMAIL, "name": ADMIN_NAME,
            "password": hash_password(pw), "role": UserRole.ADMIN.value,
            "created_at": datetime.utcnow(), "is_active": True,
            "must_change_password": must_change
        })
        if must_change:
            logger.info(f"Admin created: {ADMIN_EMAIL} / {pw}")
            logger.info("SAVE THIS PASSWORD — it will not be shown again!")
        else:
            logger.info(f"Admin created: {ADMIN_EMAIL} (password set via ADMIN_INITIAL_PASSWORD)")
            logger.info("IMPORTANT: Remove ADMIN_INITIAL_PASSWORD from .env after first boot!")


# ══════════════════════════════════════════════════════════════════════════════
# AI INTAKE PIPELINE
# ══════════════════════════════════════════════════════════════════════════════

@app.post("/api/intake/analyze")
async def intake_analyze_doc(
    file: UploadFile = File(...),
    user=Depends(get_current_user)
):
    """Step 1: Upload doc, get AI extraction with confidence/summary/missing fields."""
    if not ANTHROPIC_API_KEY:
        raise HTTPException(503, "Η υπηρεσία AI δεν έχει ρυθμιστεί")
    file_bytes = await file.read()
    if len(file_bytes) > 20 * 1024 * 1024:
        raise HTTPException(400, "Το αρχείο είναι πολύ μεγάλο (max 20MB)")
    media_type = file.content_type or "image/jpeg"
    try:
        from ai_service import intake_analyze
        result = intake_analyze(ANTHROPIC_API_KEY, file_bytes, media_type, model=MODEL_EXTRACTION)
    except Exception as e:
        logger.error(f"Intake analyze error: {e}")
        raise HTTPException(500, f"Σφάλμα ανάλυσης: {str(e)[:200]}")
    await audit("INTAKE_ANALYZE", user["id"], "document", file.filename or "unknown")
    return result


class IntakeConfirmRequest(BaseModel):
    extracted: dict
    file_b64: Optional[str] = None
    filename: str = "document.pdf"
    media_type: str = "application/pdf"
    client_id: Optional[str] = None


@app.post("/api/intake/confirm", status_code=201)
async def intake_confirm(req: IntakeConfirmRequest, user=Depends(get_current_user)):
    """Step 2: Create client+case+deadlines from confirmed AI data, upload to Drive."""
    import base64 as b64mod
    now = datetime.utcnow()
    result = {}

    # ── 1. Create or find client ──────────────────────────────────────────────
    cl_data = req.extracted.get("client") or {}
    client_id_str = None
    if req.client_id:
        client_id_str = req.client_id
        _ec = await db.clients.find_one({"_id": make_id(req.client_id)})
        result["client"] = {"id": client_id_str,
                            "full_name": (_ec or {}).get("full_name", ""),
                            "existing": True}
    elif cl_data.get("full_name"):
        # Check if client with same AFM exists
        existing_client = None
        if cl_data.get("afm"):
            existing_client = await db.clients.find_one({"afm": cl_data["afm"]})
        if not existing_client and cl_data.get("full_name"):
            existing_client = await db.clients.find_one(
                {"full_name": {"$regex": cl_data["full_name"][:10], "$options": "i"}}
            )

        if existing_client:
            client_id_str = str(existing_client["_id"])
            result["client"] = {"id": client_id_str, "full_name": existing_client.get("full_name"), "existing": True}
        else:
            client_doc = {
                "full_name": sanitize_string(cl_data.get("full_name", "")),
                "afm": cl_data.get("afm"),
                "phone": cl_data.get("phone"),
                "email": cl_data.get("email"),
                "address": sanitize_string(cl_data.get("address") or ""),
                "client_type": cl_data.get("client_type", "individual"),
                "is_active": True,
                "source": "ai_intake",
                "created_at": now,
                "created_by": user["id"],
            }
            cr = await db.clients.insert_one(client_doc)
            client_id_str = str(cr.inserted_id)
            await audit("CREATE_CLIENT", user["id"], "client", client_id_str)
            result["client"] = {"id": client_id_str, "full_name": cl_data.get("full_name"), "existing": False}

    # ── 2. Create case ────────────────────────────────────────────────────────
    case_data = req.extracted.get("case") or {}
    case_id_str = None
    if case_data.get("title") or cl_data.get("full_name"):
        cn = await case_number_gen()
        case_title = case_data.get("title") or f"Υπόθεση {cl_data.get('full_name','')}"
        case_doc = {
            "title": sanitize_string(case_title),
            "client_id": client_id_str,
            "assigned_lawyer_id": user["id"],
            "status": "active",
            "legal_category": case_data.get("category", "αστικό"),
            "court": sanitize_string(case_data.get("court") or ""),
            "description": sanitize_string(
                req.extracted.get("summary") or case_data.get("summary") or ""
            ),
            "next_action": "",
            "case_number": cn,
            "opposing_party": sanitize_string(case_data.get("opposing_party") or ""),
            "source": "ai_intake",
            "review_status": "pending_review",
            "ai_confidence": req.extracted.get("confidence", "low"),
            "ai_key_facts": req.extracted.get("key_facts", []),
            "created_at": now,
            "created_by": user["id"],
            "updated_at": now,
            "last_activity": now,
        }
        cr2 = await db.cases.insert_one(case_doc)
        case_id_str = str(cr2.inserted_id)
        await audit("CREATE_CASE", user["id"], "case", case_id_str)
        result["case"] = {"id": case_id_str, "case_number": cn, "title": case_title}

    # ── 3. Create deadlines ───────────────────────────────────────────────────
    deadlines_created = []
    for dl in (req.extracted.get("deadlines") or []):
        if not dl.get("title") and not dl.get("due_date"):
            continue
        try:
            dl_date = datetime.fromisoformat(dl["due_date"]) if dl.get("due_date") else now + timedelta(days=30)
        except Exception:
            dl_date = now + timedelta(days=30)
        dl_doc = {
            "title": sanitize_string(dl.get("title") or "Προθεσμία"),
            "case_id": case_id_str,
            "client_id": client_id_str,
            "date": dl_date,
            "deadline_type": dl.get("type", "other"),
            "completed": False,
            "source": "ai_intake",
            "created_at": now,
            "created_by": user["id"],
        }
        dr = await db.deadlines.insert_one(dl_doc)
        deadlines_created.append({"id": str(dr.inserted_id), "title": dl_doc["title"],
                                   "due_date": dl.get("due_date")})
        await audit("CREATE_DEADLINE", user["id"], "deadline", str(dr.inserted_id))
    result["deadlines"] = deadlines_created

    # ── 4. Upload to Google Drive ─────────────────────────────────────────────
    drive_link = None
    try:
        from gdrive_service import is_configured, upload_document
        if is_configured() and req.file_b64:
            file_bytes = b64mod.b64decode(req.file_b64)
            client_name = cl_data.get("full_name", "Άγνωστος") or "Άγνωστος"
            year = str(now.year)
            case_number = result.get("case", {}).get("case_number", "")
            case_title_short = (case_data.get("title") or "Υπόθεση")[:30]
            case_folder = f"{case_number} - {case_title_short}" if case_number else case_title_short
            drive_result = upload_document(
                file_bytes=file_bytes,
                filename=req.filename or "document.pdf",
                mime_type=req.media_type,
                year=year,
                client_name=client_name[:50],
                case_folder=case_folder,
            )
            drive_link = drive_result.get("folder_link")
            # Store drive link on case
            if case_id_str and drive_link:
                await db.cases.update_one(
                    {"_id": make_id(case_id_str)},
                    {"$set": {"drive_folder": drive_link}}
                )
            result["drive"] = drive_result
    except Exception as e:
        logger.warning(f"Drive upload skipped: {e}")
        result["drive"] = None

    await audit("INTAKE_CONFIRM", user["id"], "intake", case_id_str or "unknown")
    return result


# ── Google Drive integration management ──────────────────────────────────────


@app.post("/api/integrations/gdrive/folder")
async def gdrive_set_folder(req: dict, user=Depends(require_role(UserRole.ADMIN))):
    """Save the Drive root folder ID (extracted from URL or raw ID)."""
    from gdrive_service import set_root_folder_id, get_root_folder_id
    url_or_id = req.get("folder_id", "").strip()
    if not url_or_id:
        raise HTTPException(400, "folder_id is required")
    # Extract folder ID from URL if full URL provided
    import re as _re
    m = _re.search(r"/folders/([a-zA-Z0-9_-]+)", url_or_id)
    folder_id = m.group(1) if m else url_or_id
    if len(folder_id) < 10:
        raise HTTPException(400, "Μη έγκυρο folder ID")
    set_root_folder_id(folder_id)
    await audit("GDRIVE_FOLDER", user["id"], "integration", folder_id)
    return {"ok": True, "folder_id": folder_id}

@app.get("/api/integrations/gdrive/status")
async def gdrive_status(user=Depends(require_role(UserRole.ADMIN))):
    try:
        from gdrive_service import is_configured, is_oauth_client_ready, get_root_folder_id
        configured = is_configured()
        client_ready = is_oauth_client_ready()
        folder_id = get_root_folder_id()
        return {
            "configured": configured,
            "oauth_client_ready": client_ready,
            "folder_configured": folder_id is not None,
            "auth_url": "/api/integrations/gdrive/oauth/start" if client_ready and not configured else None,
        }
    except Exception as e:
        return {"configured": False, "error": str(e)}


@app.post("/api/integrations/gdrive/oauth/client")
async def gdrive_upload_oauth_client(
    file: UploadFile = File(...),
    user=Depends(require_role(UserRole.ADMIN))
):
    from gdrive_service import save_oauth_client
    content = await file.read()
    ok = save_oauth_client(content)
    if not ok:
        raise HTTPException(400, "Μη εγκυρο OAuth client JSON. Κατεβαστε το απο Google Cloud Console.")
    await audit("GDRIVE_OAUTH_CLIENT", user["id"], "integration", "gdrive")
    return {"ok": True, "next": "Visit /api/integrations/gdrive/oauth/start to authorize"}


@app.get("/api/integrations/gdrive/oauth/start")
async def gdrive_oauth_start():
    from gdrive_service import get_auth_url
    try:
        url = get_auth_url()
        from fastapi.responses import RedirectResponse
        return RedirectResponse(url)
    except RuntimeError as e:
        raise HTTPException(400, str(e))


@app.get("/api/integrations/gdrive/oauth/callback")
async def gdrive_oauth_callback(code: str = None, error: str = None, state: str = None):
    if error:
        from fastapi.responses import HTMLResponse
        return HTMLResponse(f"<h2>Authorization failed: {error}</h2><p>Close this tab and try again.</p>")
    if not code:
        raise HTTPException(400, "Missing authorization code")
    try:
        from gdrive_service import exchange_code
        exchange_code(code)
        from fastapi.responses import HTMLResponse
        return HTMLResponse("""<html><body style="font-family:sans-serif;text-align:center;padding:60px">
        <h2 style="color:#16a34a">&#10003; Google Drive connected!</h2>
        <p>Authorization successful. You can close this tab and return to Nomos One.</p>
        <script>setTimeout(()=>window.close(),3000)</script>
        </body></html>""")
    except Exception as e:
        raise HTTPException(500, f"OAuth exchange failed: {str(e)}")


@app.post("/api/integrations/gdrive/setup")
async def gdrive_setup(
    file: UploadFile = File(...),
    user=Depends(require_role(UserRole.ADMIN))
):
    from gdrive_service import save_credentials, get_service_account_email
    content = await file.read()
    ok = save_credentials(content)
    if not ok:
        raise HTTPException(400, "Μη εγκυρο αρχειο service account.")
    email = get_service_account_email()
    await audit("GDRIVE_SETUP", user["id"], "integration", "gdrive")
    return {"ok": True, "service_account_email": email}


# ══════════════════════════════════════════════════════════════════════════════
# AUTH ROUTES
# ══════════════════════════════════════════════════════════════════════════════
class LoginRequest(BaseModel):
    email: str
    password: str
    device_id: Optional[str] = None  # For 2FA device trust

class ChangePasswordRequest(BaseModel):
    current_password: str
    new_password: str

# ── Client Portal Models ──
PORTAL_CODE_TTL_DAYS = int(os.getenv("PORTAL_CODE_TTL_DAYS", "90"))

class PortalLoginRequest(BaseModel):
    name: str
    case_category: str
    portal_code: str
    source: Optional[str] = None

class PortalForgotPasswordRequest(BaseModel):
    name: str
    case_category: str

class PortalMessageRequest(BaseModel):
    content: str
    subject: Optional[str] = None

class PortalAccessRequest(BaseModel):
    permissions: List[str] = Field(default_factory=lambda: [
        'case_title', 'case_number', 'case_status', 'client_name',
        'lawyer_name', 'lawyer_email', 'lawyer_phone', 'total_fees', 'outstanding_balance'
    ])
    case_id: Optional[str] = None

@app.post("/api/auth/login")
async def login(req: LoginRequest, request: Request):
    email = req.email.strip().lower()
    if not re.match(r"[^@]+@[^@]+\.[^@]+", email):
        raise HTTPException(400, "Μη έγκυρη μορφή email")
    if rate_limiter.is_locked(email):
        mins = rate_limiter.remaining_lockout(email)
        raise HTTPException(429, f"Πολλές αποτυχημένες προσπάθειες. Δοκιμάστε σε {mins} λεπτά.")
    user = await db.users.find_one({"email": email})
    if not user or not verify_password(req.password, user["password"]):
        rate_limiter.record_attempt(email)
        raise HTTPException(401, "Λάθος email ή κωδικός")
    if not user.get("is_active", True):
        raise HTTPException(403, "Ο λογαριασμός είναι ανενεργός")
    if not user.get("is_approved", True):
        raise HTTPException(403, "Ο λογαριασμός σας δεν έχει εγκριθεί ακόμα από τον διαχειριστή")
    rate_limiter.clear(email)

    # ── Phase 1.6: Check 2FA Status ────────────────────────────────────────
    try:
        user_id_str = str(user["_id"])
        two_fa_status = await two_factor_service.get_2fa_status(user_id_str)

        # Check if device is trusted (skip 2FA if it is)
        device_id = req.device_id or "unknown"
        is_device_trusted = await two_factor_service.is_device_trusted(user_id_str, device_id)

        # If 2FA enabled and device not trusted, create OTP session
        if two_fa_status.get("enabled", False) and not is_device_trusted:
            # Create OTP session
            otp_session = await two_factor_service.create_otp_session(
                user_id=user_id_str,
                device_id=device_id,
                session_type=(
                    OTPSessionType.EMAIL_LOGIN
                    if two_fa_status.get("method") == "email"
                    else OTPSessionType.TOTP_LOGIN
                ),
                ip_address=request.client.host if request.client else "unknown",
                user_agent=request.headers.get("user-agent", "unknown")
            )

            # Send OTP email if email method
            if two_fa_status.get("method") == "email":
                await send_otp_email(
                    user["email"],
                    user.get("name", "User"),
                    otp_session["otp_code"],
                    expires_minutes=10
                )

                # Audit log
                await db.audit_logs.insert_one({
                    "timestamp": datetime.utcnow(),
                    "user_id": user_id_str,
                    "action": "2fa.otp_sent",
                    "details": {"method": "email", "device_id": device_id}
                })

            # Return OTP challenge instead of token
            def mask_email(email: str) -> str:
                try:
                    local, domain = email.split('@')
                    masked = local[:3] + "***" if len(local) > 3 else local[0] + "***"
                    return f"{masked}@{domain}"
                except:
                    return "***@***"

            return {
                "requires_2fa": True,
                "otp_session_id": otp_session["session_id"],
                "method": two_fa_status.get("method"),
                "email_masked": mask_email(user["email"]),
                "expires_in": otp_session.get("expires_in_seconds", 600)
            }
    except Exception as e:
        logger.warning(f"2FA check failed: {str(e)}, allowing login without 2FA")

    # Issue token (2FA not enabled or device trusted)
    token = create_token({"sub": str(user["_id"]), "role": user["role"]})
    await audit("LOGIN", str(user["_id"]), "auth")
    u = serialize(user)
    u.pop("password", None)
    return {"token": token, "user": u, "must_change_password": user.get("must_change_password", False)}

@app.get("/api/auth/me")
async def me(user=Depends(get_current_user)):
    u = dict(user); u.pop("password", None); return u

@app.post("/api/auth/verify-password")
async def verify_current_password(req: ChangePasswordRequest, user=Depends(get_current_user)):
    full_user = await db.users.find_one({"_id": make_id(user["id"])})
    if not full_user or not verify_password(req.current_password, full_user["password"]):
        raise HTTPException(401, "Λάθος κωδικός")
    return {"ok": True}

# ── Phase 1.6: Two-Factor Authentication Endpoints ───────────────────────────────

class OTPVerifyRequest(BaseModel):
    """OTP verification request"""
    otp_session_id: str
    code: str
    trust_device: bool = False

class BackupCodeVerifyRequest(BaseModel):
    """Backup code verification request"""
    otp_session_id: str
    code: str

@app.post("/api/auth/verify-otp")
async def verify_otp(req: OTPVerifyRequest, request: Request):
    """Verify OTP code and issue JWT token"""
    try:
        from bson import ObjectId

        # Get OTP session
        otp_session = await db.otp_sessions.find_one({"_id": ObjectId(req.otp_session_id)})
        if not otp_session:
            raise HTTPException(status_code=400, detail="Invalid OTP session")

        user_id_str = str(otp_session["user_id"])
        user_id = otp_session["user_id"]

        # Check rate limiting
        is_locked, locked_until = await two_factor_service.is_otp_locked(user_id_str)
        if is_locked:
            raise HTTPException(status_code=429, detail="Account locked. Try again in 15 minutes.")

        # Verify OTP code
        session_type = otp_session.get("otp_type", "email_login")
        if "totp" in session_type:
            valid, error = await two_factor_service.verify_totp_code(user_id_str, req.code)
        else:
            valid, error = await two_factor_service.verify_email_otp(str(req.otp_session_id), req.code)

        if not valid:
            await two_factor_service.increment_failed_otp_attempts(user_id_str)
            await db.audit_logs.insert_one({
                "timestamp": datetime.utcnow(),
                "user_id": user_id_str,
                "action": "2fa.otp_attempt_failed",
                "details": {"session_id": req.otp_session_id}
            })
            raise HTTPException(status_code=401, detail=error or "Invalid code")

        # Reset failed attempts
        await two_factor_service.reset_failed_otp_attempts(user_id_str)

        # Get user for token
        user = await db.users.find_one({"_id": user_id})

        # Issue JWT token
        token = create_token({"sub": str(user["_id"]), "role": user["role"]})

        # Mark device as trusted if requested
        trust_expires = None
        if req.trust_device:
            await two_factor_service.mark_device_as_trusted(
                user_id_str,
                otp_session["device_id"],
                otp_session.get("device_name", "Unknown Device")
            )
            device = await db.devices.find_one({"_id": otp_session["device_id"]})
            trust_expires = device.get("trust_expires_at") if device else None

        # Audit log successful verification
        await db.audit_logs.insert_one({
            "timestamp": datetime.utcnow(),
            "user_id": user_id_str,
            "action": "2fa.verified",
            "details": {
                "method": "totp" if "totp" in session_type else "email",
                "device_id": str(otp_session["device_id"]),
                "trusted": req.trust_device
            }
        })

        # Mark OTP session as verified
        await db.otp_sessions.update_one(
            {"_id": ObjectId(req.otp_session_id)},
            {"$set": {"verified": True, "verified_at": datetime.utcnow()}}
        )

        u = serialize(user)
        u.pop("password", None)
        return {
            "token": token,
            "user": u,
            "device_trusted": req.trust_device,
            "trust_expires": trust_expires
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"OTP verification error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"OTP verification error: {str(e)}")


@app.post("/api/auth/verify-backup-code")
async def verify_backup_code(req: BackupCodeVerifyRequest):
    """Verify backup code and issue JWT token"""
    try:
        from bson import ObjectId

        # Get OTP session
        otp_session = await db.otp_sessions.find_one({"_id": ObjectId(req.otp_session_id)})
        if not otp_session:
            raise HTTPException(status_code=400, detail="Invalid OTP session")

        user_id_str = str(otp_session["user_id"])
        user_id = otp_session["user_id"]

        # Verify backup code
        valid, error, remaining = await two_factor_service.use_backup_code(user_id_str, req.code)
        if not valid:
            await two_factor_service.increment_failed_otp_attempts(user_id_str)
            raise HTTPException(status_code=401, detail=error)

        # Get user for token
        user = await db.users.find_one({"_id": user_id})

        # Issue JWT token
        token = create_token({"sub": str(user["_id"]), "role": user["role"]})

        # Audit log
        await db.audit_logs.insert_one({
            "timestamp": datetime.utcnow(),
            "user_id": user_id_str,
            "action": "2fa.backup_code_used",
            "details": {"codes_remaining": remaining}
        })

        # Mark OTP session as verified
        await db.otp_sessions.update_one(
            {"_id": ObjectId(req.otp_session_id)},
            {"$set": {"verified": True, "verified_at": datetime.utcnow()}}
        )

        u = serialize(user)
        u.pop("password", None)
        return {
            "token": token,
            "user": u,
            "codes_remaining": remaining
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Backup code verification error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Backup code error: {str(e)}")


# ── 2FA Management Endpoints ─────────────────────────────────────────────────

class TOTPSetupVerifyRequest(BaseModel):
    code: str

class RegenerateCodesReq(BaseModel):
    pass

@app.post("/api/auth/2fa/setup/totp")
async def setup_totp(user=Depends(get_current_user)):
    """Generate TOTP secret and QR code, store pending secret on user."""
    try:
        import base64 as _b64
        secret = await two_factor_service.generate_totp_secret()
        qr_bytes = await two_factor_service.get_totp_qr_code(
            user["id"], user["email"], secret
        )
        qr_url = "data:image/png;base64," + _b64.b64encode(qr_bytes).decode()

        # Store encrypted pending secret so verify step can use it
        encrypted = two_factor_service.encryption.encrypt_data(secret)
        await db.users.update_one(
            {"_id": make_id(user["id"])},
            {"$set": {"two_factor_auth.pending_totp_secret": encrypted}}
        )

        return {"secret": secret, "qr_code_url": qr_url}
    except Exception as e:
        raise HTTPException(500, f"TOTP setup error: {str(e)}")


@app.post("/api/auth/2fa/setup/totp/verify")
async def verify_totp_setup(req: TOTPSetupVerifyRequest, user=Depends(get_current_user)):
    """Verify TOTP code against pending secret and activate 2FA."""
    try:
        full_user = await db.users.find_one({"_id": make_id(user["id"])})
        pending_enc = (full_user.get("two_factor_auth") or {}).get("pending_totp_secret")
        if not pending_enc:
            raise HTTPException(400, "TOTP setup not initiated")

        secret = two_factor_service.encryption.decrypt_data(pending_enc)

        # Verify code against the pending secret directly
        import pyotp as _pyotp
        totp = _pyotp.TOTP(secret)
        if not totp.verify(req.code, valid_window=1):
            raise HTTPException(400, "Μη έγκυρος κωδικός TOTP")

        # Enable 2FA with this secret
        result = await two_factor_service.enable_2fa(user["id"], OTPMethod.TOTP, secret)

        # Clear pending secret
        await db.users.update_one(
            {"_id": make_id(user["id"])},
            {"$unset": {"two_factor_auth.pending_totp_secret": ""}}
        )

        await audit("2fa.enabled", user["id"], "user", details={"method": "totp"})
        return {
            "backup_codes": result["backup_codes"],
            "download_link": "/api/auth/2fa/backup-codes/download"
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, f"TOTP verify error: {str(e)}")


@app.post("/api/auth/2fa/setup/email")
async def setup_email_2fa(user=Depends(get_current_user)):
    """Enable email OTP 2FA."""
    try:
        await two_factor_service.enable_2fa(user["id"], OTPMethod.EMAIL)
        await audit("2fa.enabled", user["id"], "user", details={"method": "email"})
        return {"status": "email_2fa_enabled", "otp_sent": False}
    except Exception as e:
        raise HTTPException(500, f"Email 2FA setup error: {str(e)}")


@app.get("/api/auth/2fa/status")
async def get_2fa_status(user=Depends(get_current_user)):
    """Return current 2FA configuration for the logged-in user."""
    try:
        status = await two_factor_service.get_2fa_status(user["id"])
        return status
    except Exception as e:
        raise HTTPException(500, f"2FA status error: {str(e)}")


@app.post("/api/auth/2fa/disable")
async def disable_2fa(user=Depends(get_current_user)):
    """Disable 2FA for the logged-in user."""
    try:
        await two_factor_service.disable_2fa(user["id"])
        await audit("2fa.disabled", user["id"], "user")
        return {"status": "disabled"}
    except Exception as e:
        raise HTTPException(500, f"Disable 2FA error: {str(e)}")


@app.post("/api/auth/2fa/regenerate-codes")
async def regenerate_backup_codes(user=Depends(get_current_user)):
    """Regenerate backup codes (old ones become invalid immediately)."""
    try:
        new_codes = await two_factor_service.regenerate_backup_codes(user["id"])
        await audit("2fa.backup_codes_regenerated", user["id"], "user")
        return {"backup_codes": new_codes}
    except Exception as e:
        raise HTTPException(500, f"Regenerate codes error: {str(e)}")


@app.get("/api/auth/trusted-devices")
async def list_trusted_devices(user=Depends(get_current_user)):
    """List all trusted devices for the logged-in user."""
    try:
        devices_cursor = db.devices.find({
            "user_id": user["id"],
            "trusted": True
        })
        devices = []
        async for d in devices_cursor:
            trust_exp = d.get("trust_expires_at")
            if trust_exp and datetime.utcnow() > trust_exp:
                continue  # skip expired
            devices.append({
                "device_id": str(d["_id"]),
                "device_name": d.get("device_name", "Unknown"),
                "device_type": d.get("device_type", "web"),
                "last_seen": d.get("last_seen"),
                "trust_expires_at": trust_exp,
            })
        return {"devices": devices, "count": len(devices)}
    except Exception as e:
        raise HTTPException(500, f"Error listing devices: {str(e)}")


@app.post("/api/auth/trusted-devices/{device_id}/revoke")
async def revoke_trusted_device(device_id: str, user=Depends(get_current_user)):
    """Revoke trust for a specific device."""
    try:
        await two_factor_service.revoke_device_trust(user["id"], device_id)
        await audit("device.trust_revoked", user["id"], "device", resource_id=device_id)
        return {"status": "revoked"}
    except Exception as e:
        raise HTTPException(500, f"Revoke device error: {str(e)}")



@app.get("/api/notifications")
async def get_notifications(user=Depends(get_current_user)):
    now = datetime.utcnow()
    notifications = []

    inv_query = {"payment_status": {"$in": ["pending", "partial"]}, "due_date": {"$lt": now}}
    overdue_count = await db.invoices.count_documents(inv_query)
    if overdue_count > 0:
        label = "ληξιπρόθεσμο" if overdue_count == 1 else "ληξιπρόθεσμα"
        notifications.append({"id": "overdue-invoices", "type": "overdue",
            "msg": f"{overdue_count} τιμολόγια", "path": "/billing"})

    end = now + timedelta(days=3)
    dl_query = {"date": {"$gte": now, "$lte": end}, "completed": {"$ne": True}}
    urgent = await db.deadlines.find(dl_query).sort("date", 1).limit(3).to_list(None)
    for d in urgent:
        days_left = (d["date"] - now).days
        when = "σήμερα" if days_left == 0 else ("αύριο" if days_left == 1 else f"σε {days_left} μέρες")
        notifications.append({"id": f"deadline-{str(d['_id'])}", "type": "deadline",
            "msg": f"Προθεσμία {when}: {d.get('title', 'Χωρίς τίτλο')}", "path": "/calendar"})

    if user["role"] == UserRole.ADMIN.value:
        unread = await db.portal_messages.count_documents({"read_by_lawyer": {"$ne": True}})
        if unread > 0:
            notifications.append({"id": "portal-messages", "type": "message",
                "msg": f"{unread} αναγνωσμένα μηνύματα από πελάτες", "path": "/admin-portal"})
        pending = await db.portal_reset_requests.count_documents({"used": False})
        if pending > 0:
            notifications.append({"id": "portal-resets", "type": "warning",
                "msg": f"{pending} αιτήματα επαναφοράς κωδικού πύλης", "path": "/admin-portal"})

    return notifications



@app.post("/api/ai/extract-document")
async def ai_extract_doc(
    file: UploadFile = File(...),
    document_type: str = Form(default="auto"),
    user=Depends(get_current_user)
):
    if not ANTHROPIC_API_KEY:
        raise HTTPException(503, "Η υπηρεσία AI δεν έχει ρυθμιστεί")
    file_bytes = await file.read()
    if len(file_bytes) > 20 * 1024 * 1024:
        raise HTTPException(400, "Το αρχείο είναι πολύ μεγάλο (max 20MB)")
    media_type = file.content_type or "image/jpeg"
    try:
        result = ai_extract_document(ANTHROPIC_API_KEY, file_bytes, media_type, document_type, model=MODEL_EXTRACTION)
    except Exception as e:
        logger.error(f"AI extract error: {e}")
        raise HTTPException(500, f"Σφάλμα εξαγωγής: {str(e)[:200]}")
    await audit("AI_EXTRACT", user["id"], "document", file.filename or "unknown")
    return result



# ── Nomos AI Bot (Claude) ────────────────────────────────────────────────────
class BotChatEntry(BaseModel):
    role: str
    content: str

class BotChatRequest(BaseModel):
    message: str
    history: List[BotChatEntry] = []

@app.post("/api/bot/chat")
async def bot_chat(req: BotChatRequest, user=Depends(get_current_user)):
    if not ANTHROPIC_API_KEY:
        raise HTTPException(503, "Η υπηρεσία AI δεν έχει ρυθμιστεί")

    system_prompt = (
        "Είσαι ο Nomos AI, νομικός βοηθός της δικηγορικής εταιρείας «Σκοτάνης & Συνεργάτες» στη Μύκονο. "
        "Ειδικεύεσαι στο ελληνικό δίκαιο: ποινικό (ΠΚ, ΚΠΔ), διοικητικό (ΣτΕ, ΔΕφ), "
        "αστικό (ΑΚ), εμπορικό, εργατικό, περιβαλλοντικό (ΑΕΚΚ, Ν.1650/1986, Ν.4042/2012). "
        "Απαντάς ΠΑΝΤΑ στα ελληνικά εκτός αν ο χρήστης γράψει σε άλλη γλώσσα. "
        "Χρησιμοποιείς νομική ορολογία, αναφέρεις νομοθεσία και νομολογία (ΑΠ, ΣτΕ κλπ.) όταν χρειάζεται. "
        "Είσαι συνοπτικός και πρακτικός. "
        "Βοηθάς με: σύνταξη δικογράφων/υπομνημάτων, ανάλυση νομικών ζητημάτων, "
        "εξήγηση νόμων, στρατηγικές υποθέσεων, υπολογισμό προθεσμιών."
    )

    messages = [
        {"role": h.role, "content": h.content}
        for h in req.history[-20:]
        if h.role in ("user", "assistant") and h.content.strip()
    ]
    messages.append({"role": "user", "content": req.message})

    async def stream_response():
        import anthropic as ant
        client = ant.AsyncAnthropic(api_key=ANTHROPIC_API_KEY)
        try:
            async with client.messages.stream(
                model=MODEL_CHAT,
                max_tokens=2048,
                system=system_prompt,
                messages=messages,
            ) as stream:
                async for text in stream.text_stream:
                    yield f"data: {json.dumps({'text': text})}\n\n"
            yield "data: [DONE]\n\n"
        except Exception as e:
            logger.error(f"Bot chat error: {e}")
            yield f"data: {json.dumps({'error': str(e)[:200]})}\n\n"

    return StreamingResponse(
        stream_response(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )

@app.post("/api/linda/chat")
async def linda_chat(req: BotChatRequest, user=Depends(get_current_user)):
    """Linda — Personal AI assistant for the lawyer."""
    if not ANTHROPIC_API_KEY:
        raise HTTPException(503, "Η υπηρεσία AI δεν έχει ρυθμιστεί")

    # Collect brief stats for context
    try:
        active_cases   = await db.cases.count_documents({"status": "active"})
        pending_cases  = await db.cases.count_documents({"status": "pending"})
        total_clients  = await db.clients.count_documents({})
        pending_docs   = await db.documents.count_documents({"uploaded_by": "portal_client", "status": "pending_review"})
        today_str      = datetime.utcnow().strftime("%d/%m/%Y")
    except Exception:
        active_cases = pending_cases = total_clients = pending_docs = 0
        today_str = datetime.utcnow().strftime("%d/%m/%Y")

    system_prompt = (
        f"Σήμερα είναι {today_str}. "
        "Λέγεσαι Λίντα. Είσαι η έξυπνη, αποτελεσματική και εξαιρετικά οργανωμένη "
        "προσωπική νομική βοηθός του δικηγόρου Χρήστου Σκοτάνη, Δικηγόρου Μυκόνου. "
        "Έχεις ζεστό, επαγγελματικό τόνο — σαν έμπιστη συνεργάτιδα που ξέρει την υπόθεση από μέσα. "
        "Μιλάς ΠΑΝΤΑ ελληνικά εκτός αν ο χρήστης γράψει σε άλλη γλώσσα. "
        "\n\nΤρέχουσα κατάσταση γραφείου:\n"
        f"• Ενεργές υποθέσεις: {active_cases}\n"
        f"• Σε εκκρεμότητα: {pending_cases}\n"
        f"• Εντολείς: {total_clients}\n"
        f"• Έγγραφα portal προς έλεγχο: {pending_docs}\n"
        "\nΜπορείς να βοηθήσεις με:\n"
        "– Οργάνωση ημερήσιας ατζέντας και προτεραιοτήτων\n"
        "– Σύνταξη νομικών εγγράφων, επιστολών, υπομνημάτων\n"
        "– Ανάλυση νομικών ζητημάτων με αναφορά στην ελληνική νομοθεσία\n"
        "– Εντοπισμό επικείμενων προθεσμιών και δικασίμων\n"
        "– Σύνοψη υποθέσεων για τρίτους\n"
        "– Οποιαδήποτε άλλη υποστήριξη χρειάζεται ο Χρήστος"
    )

    messages = [
        {"role": h.role, "content": h.content}
        for h in req.history[-20:]
        if h.role in ("user", "assistant") and h.content.strip()
    ]
    messages.append({"role": "user", "content": req.message})

    async def stream_response():
        import anthropic as ant
        client = ant.AsyncAnthropic(api_key=ANTHROPIC_API_KEY)
        try:
            async with client.messages.stream(
                model=MODEL_CHAT,
                max_tokens=2048,
                system=system_prompt,
                messages=messages,
            ) as stream:
                async for text in stream.text_stream:
                    yield f"data: {json.dumps({'text': text})}\n\n"
            yield "data: [DONE]\n\n"
        except Exception as e:
            logger.error(f"Linda chat error: {e}")
            yield f"data: {json.dumps({'error': str(e)[:200]})}\n\n"

    return StreamingResponse(
        stream_response(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )

# ── LEXIS — 12 Specialist Legal AI ──────────────────────────────────────────

class LexisRequest(BaseModel):
    message: str
    history: list = []
    specialist_id: str = ""   # empty = auto-route

@app.get("/api/lexis/specialists")
async def lexis_specialists(user=Depends(get_current_user)):
    return [{"id": s["id"], "name": s["name"], "short": s["short"],
             "color": s["color"], "icon": s["icon"]} for s in SPECIALISTS.values()]

@app.post("/api/lexis/chat")
async def lexis_chat(req: LexisRequest, user=Depends(get_current_user)):
    if not ANTHROPIC_API_KEY:
        raise HTTPException(503, "Η υπηρεσία AI δεν έχει ρυθμιστεί")

    sid = req.specialist_id if req.specialist_id else route_question(req.message)
    spec = SPECIALISTS.get(sid, SPECIALISTS["civil"])

    messages = [{"role": h["role"], "content": h["content"]}
                for h in req.history[-20:]
                if h.get("role") in ("user", "assistant") and h.get("content", "").strip()]
    messages.append({"role": "user", "content": req.message})

    async def stream_response():
        import anthropic as ant
        client = ant.AsyncAnthropic(api_key=ANTHROPIC_API_KEY)
        yield f"data: {json.dumps({'specialist_id': sid, 'specialist_name': spec['name']})}\n\n"
        try:
            # Fetch live νομολογία — HUDOC/EUR-Lex + Solon concurrently
            yield f"data: {json.dumps({'status': 'Αναζήτηση νομολογίας...'})}\n\n"
            try:
                nomologia_ctx, solon_results = await asyncio.wait_for(
                    asyncio.gather(
                        retrieve_relevant_nomologia(req.message, sid),
                        search_solon(req.message, sid, max_results=4),
                        return_exceptions=True,
                    ),
                    timeout=12.0
                )
            except asyncio.TimeoutError:
                nomologia_ctx, solon_results = "", []
                logger.warning("nomologia retrieval timed out")

            if isinstance(nomologia_ctx, Exception):
                nomologia_ctx = ""
            if isinstance(solon_results, Exception):
                solon_results = []

            system = spec["system_prompt"]
            if nomologia_ctx:
                system = system + "\n\n" + nomologia_ctx
            if solon_results:
                solon_ctx = await format_solon_for_prompt(solon_results)
                if solon_ctx:
                    system = system + "\n\n" + solon_ctx

            async with client.messages.stream(
                model=MODEL_CHAT,
                max_tokens=2048,
                system=system,
                messages=messages,
            ) as stream:
                async for text in stream.text_stream:
                    yield f"data: {json.dumps({'text': text})}\n\n"
            yield "data: [DONE]\n\n"
        except Exception as e:
            logger.error(f"Lexis chat error: {e}")
            yield f"data: {json.dumps({'error': str(e)[:200]})}\n\n"

    return StreamingResponse(stream_response(), media_type="text/event-stream",
                             headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})


@app.get("/api/lexis/search-nomologia")
async def lexis_search_nomologia(q: str, specialist: str = "echr", user=Depends(get_current_user)):
    """Direct νομολογία search — returns raw results list."""
    from nomologia_service import _search_hudoc, _search_eurlex
    hudoc, eurlex = await asyncio.gather(
        _search_hudoc(q, specialist, max_results=5),
        _search_eurlex(q, specialist, max_results=3),
        return_exceptions=True,
    )
    return {
        "echr": hudoc if not isinstance(hudoc, Exception) else [],
        "eurlex": eurlex if not isinstance(eurlex, Exception) else [],
    }


@app.post("/api/auth/change-password")
async def change_password(req: ChangePasswordRequest, user=Depends(get_current_user)):
    full_user = await db.users.find_one({"_id": make_id(user["id"])})
    if not verify_password(req.current_password, full_user["password"]):
        raise HTTPException(400, "Ο τρέχων κωδικός είναι λάθος")
    is_valid, err = validate_password(req.new_password)
    if not is_valid: raise HTTPException(400, err)
    if req.current_password == req.new_password:
        raise HTTPException(400, "Ο νέος κωδικός πρέπει να είναι διαφορετικός")
    await db.users.update_one({"_id": make_id(user["id"])}, {"$set": {
        "password": hash_password(req.new_password),
        "must_change_password": False,
        "password_changed_at": datetime.utcnow()
    }})
    await audit("CHANGE_PASSWORD", user["id"], "auth")
    return {"ok": True, "message": "Ο κωδικός άλλαξε επιτυχώς"}

# ══════════════════════════════════════════════════════════════════════════════
# USERS
# ══════════════════════════════════════════════════════════════════════════════
class CreateUserRequest(BaseModel):
    email: str; name: str; password: str; role: UserRole

class RegisterUserRequest(BaseModel):
    email: str; name: str; password: str

class UpdateUserRequest(BaseModel):
    name: Optional[str] = None; email: Optional[str] = None
    role: Optional[UserRole] = None; is_active: Optional[bool] = None
    password: Optional[str] = None

# ══════════════════════════════════════════════════════════════════════════════
# PWA & MOBILE - API v1
# ══════════════════════════════════════════════════════════════════════════════
class RegisterDeviceRequest(BaseModel):
    device_name: str = Field(..., min_length=1, max_length=255)
    device_type: str = Field(..., description="ios | android | web | desktop")
    push_token: str = Field(..., min_length=1)
    app_version: str = Field(..., description="e.g., 1.0.0")

class TrustDeviceRequest(BaseModel):
    device_name: Optional[str] = None

@app.get("/api/users")
async def list_users(user=Depends(require_role(UserRole.ADMIN))):
    users = await db.users.find({}).to_list(None)
    return [{k: v for k, v in serialize(u).items() if k != "password"} for u in users]

@app.post("/api/users", status_code=201)
async def create_user(req: CreateUserRequest, user=Depends(require_role(UserRole.ADMIN))):
    email = req.email.strip().lower()
    if not re.match(r"[^@]+@[^@]+\.[^@]+", email): raise HTTPException(400, "Μη έγκυρη μορφή email")
    is_valid, err = validate_password(req.password)
    if not is_valid: raise HTTPException(400, err)
    if await db.users.find_one({"email": email}): raise HTTPException(409, "Το email χρησιμοποιείται ήδη")
    doc = {"email": email, "name": sanitize_string(req.name), "password": hash_password(req.password),
           "role": req.role.value, "is_active": True, "is_approved": True, "created_at": datetime.utcnow(), "must_change_password": True}
    result = await db.users.insert_one(doc)
    await audit("CREATE_USER", user["id"], "user", str(result.inserted_id))
    doc["_id"] = result.inserted_id
    s = serialize(doc); s.pop("password", None); return s

# ── Self-update profile (any logged-in user) — MUST be before /{user_id} ──
class UpdateProfileRequest(BaseModel):
    full_name: Optional[str] = None
    phone:     Optional[str] = None
    email:     Optional[str] = None

@app.put("/api/users/me")
async def update_my_profile(req: UpdateProfileRequest, user=Depends(get_current_user)):
    """Any user can update their own name, phone, email."""
    update: dict = {"updated_at": datetime.utcnow()}
    if req.full_name and req.full_name.strip():
        update["full_name"] = sanitize_string(req.full_name.strip())
        update["name"]      = update["full_name"]
    if req.phone is not None:
        update["phone"] = sanitize_string(req.phone.strip())
    if req.email and req.email.strip():
        email = req.email.strip().lower()
        if not re.match(r"[^@]+@[^@]+\.[^@]+", email):
            raise HTTPException(400, "Μη έγκυρη μορφή email")
        if await db.users.find_one({"email": email, "_id": {"$ne": make_id(user["id"])}}):
            raise HTTPException(409, "Το email χρησιμοποιείται ήδη")
        update["email"] = email
    await db.users.update_one({"_id": make_id(user["id"])}, {"$set": update})
    await audit("UPDATE_PROFILE", user["id"], "user", user["id"])
    updated = await db.users.find_one({"_id": make_id(user["id"])})
    return {
        "ok":    True,
        "name":  updated.get("full_name") or updated.get("name", ""),
        "email": updated.get("email", ""),
        "phone": updated.get("phone", ""),
    }

@app.put("/api/users/{user_id}")
async def update_user(user_id: str, req: UpdateUserRequest, user=Depends(require_role(UserRole.ADMIN))):
    update = {}
    if req.name: update["name"] = sanitize_string(req.name)
    if req.email:
        email = req.email.strip().lower()
        if not re.match(r"[^@]+@[^@]+\.[^@]+", email): raise HTTPException(400, "Μη έγκυρη μορφή email")
        if await db.users.find_one({"email": email, "_id": {"$ne": make_id(user_id)}}):
            raise HTTPException(409, "Το email χρησιμοποιείται ήδη")
        update["email"] = email
    if req.role: update["role"] = req.role.value
    if req.is_active is not None: update["is_active"] = req.is_active
    if req.password:
        is_valid, err = validate_password(req.password)
        if not is_valid: raise HTTPException(400, err)
        update["password"] = hash_password(req.password)
        update["must_change_password"] = True
    if not update: raise HTTPException(400, "Δεν υπάρχουν πεδία")
    update["updated_at"] = datetime.utcnow()
    await db.users.update_one({"_id": make_id(user_id)}, {"$set": update})
    await audit("UPDATE_USER", user["id"], "user", user_id)
    return {"ok": True}

@app.delete("/api/users/{user_id}")
async def delete_user(user_id: str, user=Depends(require_role(UserRole.ADMIN))):
    if user_id == user["id"]: raise HTTPException(400, "Δεν μπορείτε να διαγράψετε τον εαυτό σας")
    await db.users.delete_one({"_id": make_id(user_id)})
    await audit("DELETE_USER", user["id"], "user", user_id)
    return {"ok": True}

# ── User Registration (Public) ──
@app.post("/api/auth/register", status_code=201)
async def register_user(req: RegisterUserRequest):
    """Δημόσια εγγραφή χρήστη — δημιουργεί pending user που χρειάζεται έγκριση"""
    email = req.email.strip().lower()
    if not re.match(r"[^@]+@[^@]+\.[^@]+", email):
        raise HTTPException(400, "Μη έγκυρη μορφή email")
    is_valid, err = validate_password(req.password)
    if not is_valid:
        raise HTTPException(400, err)
    if await db.users.find_one({"email": email}):
        raise HTTPException(409, "Το email χρησιμοποιείται ήδη")
    doc = {"email": email, "name": sanitize_string(req.name), "password": hash_password(req.password),
           "role": UserRole.LAWYER.value, "is_active": True, "is_approved": False,
           "created_at": datetime.utcnow(), "must_change_password": True}
    result = await db.users.insert_one(doc)
    await audit("REGISTER_USER", str(result.inserted_id), "user", str(result.inserted_id))
    return {"message": "Εγγραφή επιτυχής. Περιμένετε την έγκριση διαχειριστή.", "email": email}

# ── Admin: List Pending Users ──
@app.get("/api/admin/pending-users")
async def list_pending_users(user=Depends(require_role(UserRole.ADMIN))):
    """Λίστα χρηστών που περιμένουν έγκριση"""
    users = await db.users.find({"is_approved": False}).to_list(None)
    return [{k: v for k, v in serialize(u).items() if k != "password"} for u in users]

# ── Admin: Approve User ──
@app.post("/api/admin/users/{user_id}/approve")
async def approve_user(user_id: str, user=Depends(require_role(UserRole.ADMIN))):
    """Έγκριση χρήστη για πρόσβαση στο σύστημα"""
    result = await db.users.update_one({"_id": make_id(user_id)}, {"$set": {"is_approved": True}})
    if result.matched_count == 0:
        raise HTTPException(404, "Χρήστης δεν βρέθηκε")
    await audit("APPROVE_USER", user["id"], "user", user_id)
    return {"ok": True, "message": "Χρήστης εγκρίθηκε"}

# ── Admin: Reject User ──
@app.post("/api/admin/users/{user_id}/reject")
async def reject_user(user_id: str, user=Depends(require_role(UserRole.ADMIN))):
    """Απόρριψη και διαγραφή χρήστη"""
    if user_id == user["id"]:
        raise HTTPException(400, "Δεν μπορείτε να απορρίψετε τον εαυτό σας")
    result = await db.users.delete_one({"_id": make_id(user_id)})
    if result.deleted_count == 0:
        raise HTTPException(404, "Χρήστης δεν βρέθηκε")
    await audit("REJECT_USER", user["id"], "user", user_id)
    return {"ok": True, "message": "Χρήστης απορρίφθηκε"}

# ══════════════════════════════════════════════════════════════════════════════
# CLIENT PORTAL
# ══════════════════════════════════════════════════════════════════════════════

def create_portal_token(data: dict, expires_in_hours: int = 730) -> str:
    """Create JWT token for portal access (30 days)"""
    payload = {**data, "exp": datetime.utcnow() + timedelta(hours=expires_in_hours), "type": "portal"}
    return jwt.encode(payload, JWT_SECRET, algorithm="HS256")

async def get_portal_user(credentials: HTTPAuthorizationCredentials = Depends(HTTPBearer(auto_error=False))) -> dict:
    """Get portal user from JWT token"""
    if not credentials:
        raise HTTPException(401, "Δεν υπάρχει πρόσβαση")
    try:
        payload = jwt.decode(credentials.credentials, JWT_SECRET, algorithms=["HS256"])
        if payload.get("type") != "portal":
            raise HTTPException(401, "Μη έγκυρο token")
        return payload
    except jwt.ExpiredSignatureError:
        raise HTTPException(401, "Το token έχει λήξει")
    except jwt.InvalidTokenError:
        raise HTTPException(401, "Μη έγκυρο token")

@app.post("/api/portal/auth")
async def portal_login(req: PortalLoginRequest, request: Request):
    """Client portal authentication with name and code"""
    client_ip = request.client.host if request.client else "unknown"
    user_agent = request.headers.get("user-agent", "unknown")

    portal_access = await db.portal_access.find_one({
        "portal_code": {"$regex": f"^{re.escape(req.portal_code.strip())}$", "$options": "i"},
        "is_active": True
    })
    if not portal_access:
        raise HTTPException(401, "Μη έγκυρος κωδικός πρόσβασης")

    # Check TTL expiry
    expires_at = portal_access.get("expires_at")
    if expires_at and datetime.utcnow() > expires_at:
        raise HTTPException(401, "Ο κωδικός πρόσβασης έχει λήξει. Επικοινωνήστε με το γραφείο.")

    # Compare name against stored client_name (set when generating access)
    stored_name = (portal_access.get("client_name") or "").strip().lower()
    submitted_name = req.name.strip().lower()
    if stored_name and stored_name != submitted_name:
        raise HTTPException(401, "Μη συμφωνία στοιχείων")

    case = await db.cases.find_one({"_id": ObjectId(portal_access["case_id"])})

    now = datetime.utcnow()
    await db.portal_access.update_one(
        {"_id": portal_access["_id"]},
        {"$set": {"accessed_at": now, "last_ip": client_ip, "last_user_agent": user_agent[:200]}}
    )

    # Full audit trail
    import hashlib
    code_hash = hashlib.sha256(req.portal_code.strip().encode()).hexdigest()
    await db.portal_access_logs.insert_one({
        "portal_access_id": str(portal_access["_id"]),
        "code_hash": code_hash,
        "case_id": portal_access["case_id"],
        "client_id": str(portal_access["client_id"]),
        "client_name": req.name,
        "ip": client_ip,
        "user_agent": user_agent[:300],
        "source": req.source or "direct",
        "timestamp": now,
    })

    client_name = portal_access.get("client_name") or req.name
    token = create_portal_token({
        "client_id": portal_access["client_id"],
        "case_id": portal_access["case_id"],
        "client_name": client_name,
        "permissions": portal_access.get("permissions", [])
    })

    await audit("PORTAL_LOGIN", str(portal_access["client_id"]), "portal", portal_access["case_id"],
                details={"ip": client_ip, "source": req.source or "direct", "code_hash": code_hash})
    return {
        "token": token,
        "client_name": client_name,
        "case_title": case.get("title") if case else "",
        "mandate_accepted": portal_access.get("mandate_accepted", False),
    }

@app.post("/api/portal/accept-mandate")
async def portal_accept_mandate(request: Request, user=Depends(get_portal_user)):
    """Record client acceptance of the mandate (εντολή) on first portal login."""
    client_ip = request.client.host if request.client else "unknown"
    accepted_at = datetime.utcnow()

    # Find portal access record
    portal_access = await db.portal_access.find_one({
        "case_id": user["case_id"],
        "client_id": user["client_id"],
        "is_active": True,
    })
    if not portal_access:
        raise HTTPException(404, "Portal access not found")
    if portal_access.get("mandate_accepted"):
        return {"ok": True, "already_accepted": True}

    # Get case info for the email
    case = await db.cases.find_one({"_id": ObjectId(user["case_id"])})
    case_title    = (case or {}).get("title") or (case or {}).get("offense") or "—"
    case_category = (case or {}).get("category") or (case or {}).get("legal_category") or "—"
    case_subject  = (case or {}).get("description") or (case or {}).get("offense") or "—"
    client_name   = user.get("client_name") or portal_access.get("client_name", "—")
    client_email  = portal_access.get("client_email", "")

    # Persist acceptance
    await db.portal_access.update_one(
        {"_id": portal_access["_id"]},
        {"$set": {
            "mandate_accepted": True,
            "mandate_accepted_at": accepted_at,
            "mandate_ip": client_ip,
        }}
    )
    await audit("PORTAL_MANDATE_ACCEPT", user["client_id"], "portal", user["case_id"],
                {"ip": client_ip, "client_name": client_name})

    # Send confirmation email to client
    if client_email:
        try:
            settings_doc = await db.settings.find_one({"_id": "global"}) or {}
            smtp_host = settings_doc.get("smtp_host") or os.getenv("SMTP_HOST", "")
            smtp_port = int(settings_doc.get("smtp_port") or os.getenv("SMTP_PORT", "587"))
            smtp_user = settings_doc.get("smtp_user") or os.getenv("SMTP_USER", "")
            smtp_pass = settings_doc.get("smtp_pass") or os.getenv("SMTP_PASS", "")
            from_email = settings_doc.get("notification_email") or smtp_user or FIRM_EMAIL_DISPLAY
            noreply_addr = os.getenv("SMTP_FROM", from_email)
            ts_str = accepted_at.strftime("%d/%m/%Y %H:%M") + " UTC"

            mandate_html = f"""
<!DOCTYPE html><html lang="el">
<head><meta charset="UTF-8">
<style>
  body{{font-family:'Segoe UI',Arial,sans-serif;background:#f4f6f9;margin:0;padding:0}}
  .wrap{{max-width:640px;margin:32px auto;background:#fff;border-radius:12px;overflow:hidden;box-shadow:0 2px 16px rgba(0,0,0,.1)}}
  .hdr{{background:linear-gradient(135deg,#071220,#0a1929);padding:32px 40px;text-align:center}}
  .hdr h1{{color:#C6A75E;font-size:22px;margin:0 0 6px}}
  .hdr p{{color:#8aa0b8;font-size:13px;margin:0}}
  .body{{padding:36px 40px}}
  .body p{{color:#374151;font-size:14px;line-height:1.7;margin:0 0 14px}}
  .mandate-box{{background:#f8fafc;border:1px solid #e2e8f0;border-left:4px solid #C6A75E;border-radius:8px;padding:20px 24px;margin:20px 0}}
  .mandate-box p{{margin:0 0 10px;font-size:13px;color:#374151}}
  .mandate-box p:last-child{{margin:0}}
  .info-row{{display:flex;gap:8px;margin:6px 0}}
  .info-label{{font-weight:600;color:#6b7280;font-size:12px;min-width:130px}}
  .info-val{{color:#111827;font-size:13px}}
  .stamp{{background:#f0fdf4;border:1px solid #bbf7d0;border-radius:8px;padding:14px 20px;margin:20px 0;text-align:center}}
  .stamp p{{color:#166534;font-size:13px;font-weight:600;margin:0}}
  .footer{{background:#f8fafc;border-top:1px solid #e5e7eb;padding:20px 40px;text-align:center}}
  .footer p{{color:#9ca3af;font-size:11px;margin:0}}
</style></head>
<body><div class="wrap">
  <div class="hdr">
    <h1>Σκοτάνης &amp; Συνεργάτες</h1>
    <p>Επιβεβαίωση Εντολής &amp; Αποδοχής Αμοιβής</p>
  </div>
  <div class="body">
    <p>Αγαπητέ/ή <strong>{client_name}</strong>,</p>
    <p>Σας αποστέλλουμε επιβεβαίωση της ψηφιακής αποδοχής της εντολής που χορηγήσατε στο δικηγορικό μας γραφείο μέσω της ηλεκτρονικής πλατφόρμας Client Portal.</p>

    <div class="mandate-box">
      <p><strong>ΣΤΟΙΧΕΙΑ ΕΝΤΟΛΗΣ</strong></p>
      <div class="info-row"><span class="info-label">Εντολέας:</span><span class="info-val">{client_name}</span></div>
      <div class="info-row"><span class="info-label">Τίτλος Υπόθεσης:</span><span class="info-val">{case_title}</span></div>
      <div class="info-row"><span class="info-label">Κατηγορία:</span><span class="info-val">{case_category}</span></div>
      <div class="info-row"><span class="info-label">Αντικείμενο:</span><span class="info-val">{case_subject}</span></div>
      <div class="info-row"><span class="info-label">Ημ/νία Αποδοχής:</span><span class="info-val">{ts_str}</span></div>
      <div class="info-row"><span class="info-label">Διεύθυνση IP:</span><span class="info-val">{client_ip}</span></div>
    </div>

    <p>Με την αποδοχή σας επιβεβαιώνετε ότι:</p>
    <ul style="color:#374151;font-size:14px;line-height:1.8;padding-left:20px">
      <li>Εντέλλεσθε το δικηγορικό γραφείο <strong>Σκοτάνης &amp; Συνεργάτες</strong> να αναλάβει και να χειριστεί την ανωτέρω υπόθεσή σας.</li>
      <li>Αποδέχεστε και δεσμεύεστε να καταβάλετε τη <strong>συμφωνηθείσα αμοιβή</strong> του γραφείου.</li>
      <li>Αναγνωρίζετε ότι τυχόν δικαστικά έξοδα, γραμμάτια προκαταβολής εισφορών και λοιπές δαπάνες υπόθεσης <strong>βαρύνουν εσάς</strong> ως εντολέα.</li>
    </ul>

    <div class="stamp">
      <p>✓ Η εντολή αποδέχθηκε ψηφιακά — {ts_str}</p>
    </div>

    <p style="font-size:12px;color:#6b7280">Το παρόν αποτελεί απόδειξη ψηφιακής αποδοχής βάσει του ν. 3979/2011 και του Κανονισμού EU 910/2014 (eIDAS). Φυλάξτε το για τα αρχεία σας.</p>
  </div>
  <div class="footer">
    <p>Σκοτάνης &amp; Συνεργάτες | christos@skotanislaw.com</p>
  </div>
</div></body></html>"""

            if smtp_host and smtp_user:
                import smtplib
                from email.mime.multipart import MIMEMultipart
                from email.mime.text import MIMEText as MIMETextLocal
                msg = MIMEMultipart("alternative")
                msg["Subject"] = "Επιβεβαίωση Εντολής — Σκοτάνης & Συνεργάτες"
                msg["From"]    = f"Σκοτάνης & Συνεργάτες <{noreply_addr}>"
                msg["To"]      = f"{client_name} <{client_email}>"
                msg.attach(MIMETextLocal(mandate_html, "html", "utf-8"))
                with smtplib.SMTP(smtp_host, smtp_port, timeout=10) as srv:
                    srv.ehlo(); srv.starttls(); srv.login(smtp_user, smtp_pass)
                    srv.sendmail(noreply_addr, client_email, msg.as_string())
        except Exception as e:
            logger.warning(f"Mandate email failed: {e}")

    return {"ok": True, "accepted_at": accepted_at.isoformat()}

@app.post("/api/portal/forgot-password")
async def portal_forgot_password(req: PortalForgotPasswordRequest):
    """Client forgot password - notify admin"""
    # Find client case
    case = await db.cases.find_one({
        "$and": [
            {"client_name": {"$regex": f"^{re.escape(req.name)}$", "$options": "i"}},
            {"category": req.case_category}
        ]
    })

    if not case:
        # Don't reveal if case exists
        return {"message": "Αν η υπόθεση υπάρχει, ο διαχειριστής θα λάβει ειδοποίηση"}

    # Create password reset request
    reset_code = secrets.token_urlsafe(32)
    await db.portal_reset_requests.insert_one({
        "client_name": req.name,
        "case_id": str(case["_id"]),
        "reset_code": reset_code,
        "created_at": datetime.utcnow(),
        "expires_at": datetime.utcnow() + timedelta(hours=24),
        "used": False
    })

    await audit("PORTAL_FORGOT_PASSWORD", str(case["_id"]), "portal", str(case["_id"]))
    return {"message": "Αίτημα επαναφοράς κωδικού καταχωρήθηκε. Ο διαχειριστής θα επικοινωνήσει."}

@app.get("/api/portal/my-case")
async def portal_get_case(user=Depends(get_portal_user)):
    """Get case data visible to client (permission-filtered)"""
    try:
        case = await db.cases.find_one({"_id": ObjectId(user["case_id"])})
    except InvalidId:
        raise HTTPException(404, "Υπόθεση δεν βρέθηκε")
    if not case:
        raise HTTPException(404, "Υπόθεση δεν βρέθηκε")

    permissions = user.get("permissions", [])

    # Get mandate status from portal_access
    portal_access_doc = await db.portal_access.find_one({
        "case_id": user["case_id"], "client_id": user["client_id"], "is_active": True
    })
    mandate_accepted = (portal_access_doc or {}).get("mandate_accepted", False)

    filtered_case: dict = {
        "id": str(case["_id"]),
        "title":            case.get("title", "—")       if "case_title"  in permissions else "—",
        "case_number":      case.get("case_number", "—") if "case_number" in permissions else "—",
        "status":           case.get("status", "—")      if "case_status" in permissions else "—",
        "category":         case.get("legal_category") or case.get("category", "—") if "case_status" in permissions else "—",
        "next_action":      case.get("next_action", "")  if "case_status" in permissions else None,
        "mandate_accepted": mandate_accepted,
        "case_subject":     case.get("description") or case.get("offense") or "—",
    }

    # Lawyer info
    lawyer_obj = None
    if case.get("assigned_lawyer_id"):
        try:
            lawyer_obj = await db.users.find_one({"_id": ObjectId(case["assigned_lawyer_id"])})
        except Exception:
            pass

    if "lawyer_name" in permissions and lawyer_obj:
        filtered_case["lawyer_name"]  = lawyer_obj.get("full_name") or lawyer_obj.get("name", "—")
        filtered_case["lawyer_email"] = lawyer_obj.get("email", "") if "lawyer_email" in permissions else None
        filtered_case["lawyer_phone"] = lawyer_obj.get("phone", "") if "lawyer_phone" in permissions else None
        # Build nested object for frontend compatibility
        filtered_case["lawyer"] = {
            "name":  filtered_case["lawyer_name"],
            "email": filtered_case.get("lawyer_email"),
            "phone": filtered_case.get("lawyer_phone"),
        }

    # Financial summary
    if "total_fees" in permissions or "outstanding_balance" in permissions:
        invoices = await db.invoices.find({"case_id": user["case_id"]}).to_list(None)
        total_invoiced = sum(float(i.get("total_payable", i.get("total", 0))) for i in invoices)
        total_paid     = sum(float(i.get("amount_paid", 0)) for i in invoices)
        filtered_case["total_fees"]          = round(total_invoiced, 2) if "total_fees" in permissions else None
        filtered_case["paid_fees"]           = round(total_paid, 2)
        filtered_case["outstanding_balance"] = round(total_invoiced - total_paid, 2) if "outstanding_balance" in permissions else None

    return filtered_case

@app.get("/api/portal/case-events")
async def portal_get_events(user=Depends(get_portal_user)):
    """Get case timeline events visible to client"""
    # Get non-sensitive audit events
    events = await db.audit_logs.find({
        "$or": [
            {"case_id": user["case_id"], "action": {"$in": ["CREATE_NOTE", "CREATE_DEADLINE", "UPDATE_STATUS", "CREATE_INVOICE"]}},
            {"entity_id": user["case_id"]}
        ]
    }).sort("created_at", -1).to_list(None)

    return [serialize(e) for e in events[:20]]

@app.post("/api/portal/messages")
async def portal_send_message(req: PortalMessageRequest, user=Depends(get_portal_user)):
    """Client send message to lawyer and admin"""
    case = await db.cases.find_one({"_id": ObjectId(user["case_id"])})
    if not case:
        raise HTTPException(404, "Υπόθεση δεν βρέθηκε")

    message = {
        "case_id": user["case_id"],
        "client_name": user.get("client_name"),
        "content": sanitize_string(req.content),
        "subject": req.subject or "Μήνυμα από πελάτη",
        "created_at": datetime.utcnow(),
        "read": False
    }

    result = await db.portal_messages.insert_one(message)

    # Notify lawyer via push
    lawyer_id = case.get("assigned_lawyer_id")
    if lawyer_id:
        try:
            push_svc = get_push_service(db)
            client_name = case.get("client_name", "Πελάτης")
            msg_text = message.get("message", "")[:80]
            await push_svc.send_to_user(
                str(lawyer_id),
                title="Νέο μήνυμα από πελάτη",
                body=f"{client_name}: {msg_text}",
                path="/admin-portal",
            )
        except Exception as _pe:
            logger.warning(f"Push notification failed: {_pe}")

    await audit("PORTAL_MESSAGE", user.get("client_id", ""), "portal", str(result.inserted_id))
    return {"ok": True, "message_id": str(result.inserted_id)}

@app.post("/api/portal/upload")
async def portal_upload_document(file: UploadFile = File(...), user=Depends(get_portal_user)):
    """Client stages a document — AI analysis runs only after confirm."""
    if not file.filename:
        raise HTTPException(400, "Δεν υπάρχει αρχείο")

    file_bytes = await file.read()
    doc_id     = str(ObjectId())
    case_id    = user["case_id"]
    doc_dir    = Path(f"documents/{case_id}")
    doc_dir.mkdir(parents=True, exist_ok=True)
    file_path  = doc_dir / f"{doc_id}_{file.filename}"

    with open(file_path, "wb") as fh:
        fh.write(file_bytes)

    doc = {
        "case_id":          case_id,
        "filename":         file.filename,
        "file_path":        str(file_path),
        "mime_type":        file.content_type or "application/octet-stream",
        "uploaded_by":      "portal_client",
        "uploaded_by_name": user.get("client_name", ""),
        "client_id":        user.get("client_id", ""),
        "created_at":       datetime.utcnow(),
        "size":             len(file_bytes),
        "ai_summary":       "",
        "status":           "staged",
    }
    result = await db.documents.insert_one(doc)
    doc_id_str = str(result.inserted_id)

    await audit("PORTAL_UPLOAD_STAGE", user.get("client_id", ""), "document", doc_id_str)
    return {"ok": True, "document_id": doc_id_str, "filename": file.filename}

@app.post("/api/portal/upload-confirm")
async def portal_upload_confirm(user=Depends(get_portal_user)):
    """Client confirms batch upload → silent AI analysis → notify lawyer."""
    case_id = user["case_id"]
    staged = await db.documents.find({"case_id": case_id, "status": "staged"}).to_list(None)
    if not staged:
        return {"ok": True, "confirmed": 0}

    case = await db.cases.find_one({"_id": ObjectId(case_id)})
    filenames = [d["filename"] for d in staged]
    doc_ids = [str(d["_id"]) for d in staged]

    # ── AI summary for each document (background, silent) ────────────────────
    summaries = {}
    if ANTHROPIC_API_KEY:
        try:
            import anthropic as ant
            cli = ant.Anthropic(api_key=ANTHROPIC_API_KEY)
            for doc in staged:
                fn_lower = (doc.get("filename") or "").lower()
                text_preview = ""
                try:
                    file_bytes = open(doc["file_path"], "rb").read()
                    if fn_lower.endswith(".txt"):
                        text_preview = file_bytes.decode("utf-8", errors="ignore")[:3000]
                    elif fn_lower.endswith(".docx"):
                        from docx import Document as DocxDoc
                        import io as _io
                        d = DocxDoc(_io.BytesIO(file_bytes))
                        text_preview = "\n".join(p.text for p in d.paragraphs if p.text)[:3000]
                except Exception:
                    pass
                prompt = (
                    f"Αρχείο από πελάτη: «{doc['filename']}» ({doc.get('size', 0)//1024} KB).\n"
                    + (f"Περιεχόμενο:\n{text_preview[:2000]}\n\n" if text_preview else "")
                    + "Γράψε 2-3 προτάσεις σύνοψης ελληνικά για τον χειριστή δικηγόρο. "
                      "Εστίασε σε τι υποβάλλει ο πελάτης και αν απαιτείται άμεση ενέργεια."
                )
                try:
                    resp = cli.messages.create(
                        model=MODEL_EXTRACTION, max_tokens=300,
                        messages=[{"role": "user", "content": prompt}]
                    )
                    summaries[str(doc["_id"])] = resp.content[0].text.strip()
                except Exception as e:
                    logger.warning(f"AI summary failed for {doc['filename']}: {e}")
        except Exception as e:
            logger.warning(f"AI init failed: {e}")

    # ── Move staged → pending_review ──────────────────────────────────────────
    for doc in staged:
        did = str(doc["_id"])
        await db.documents.update_one(
            {"_id": doc["_id"]},
            {"$set": {"status": "pending_review", "ai_summary": summaries.get(did, ""), "confirmed_at": datetime.utcnow()}}
        )

    # ── Notify lawyer ─────────────────────────────────────────────────────────
    if case:
        ai_text = "\n".join(f"• {fn}: {summaries.get(did,'')}" for fn, did in zip(filenames, doc_ids) if summaries.get(did))
        notif = {
            "user_id":    case.get("assigned_lawyer_id", ""),
            "type":       "portal_upload",
            "title":      f"Νέα έγγραφα από {user.get('client_name', 'πελάτη')} ({len(staged)} αρχεία)",
            "message":    "Αρχεία: " + ", ".join(filenames) + (f"\n\nΑνάλυση:\n{ai_text}" if ai_text else ""),
            "case_id":    case_id,
            "read":       False,
            "created_at": datetime.utcnow(),
        }
        await db.notifications.insert_one(notif)

    await audit("PORTAL_UPLOAD_CONFIRM", user.get("client_id", ""), "document", case_id)
    return {"ok": True, "confirmed": len(staged)}

@app.get("/api/portal/messages")
async def portal_get_messages(user=Depends(get_portal_user)):
    """Client reads their message thread with lawyer."""
    msgs = await db.portal_messages.find(
        {"case_id": user["case_id"]}
    ).sort("created_at", 1).to_list(None)
    return [serialize(m) for m in msgs]

@app.get("/api/portal/progress")
async def portal_get_progress(user=Depends(get_portal_user)):
    """Get upcoming hearings and deadlines for the portal client's case."""
    case_id  = user["case_id"]
    now_str  = datetime.utcnow().date().isoformat()

    hearings_raw  = await db.hearings.find({"case_id": case_id}).sort("hearing_date", 1).to_list(None)
    deadlines_raw = await db.deadlines.find({"case_id": case_id}).sort("date", 1).to_list(None)

    upcoming_hearings = [
        serialize(h) for h in hearings_raw
        if str(h.get("hearing_date", ""))[:10] >= now_str
        and h.get("status") not in ("completed", "cancelled")
    ]
    past_hearings = [
        serialize(h) for h in hearings_raw
        if str(h.get("hearing_date", ""))[:10] < now_str
        or h.get("status") in ("completed", "cancelled")
    ]
    upcoming_deadlines = [
        serialize(d) for d in deadlines_raw
        if str(d.get("date", ""))[:10] >= now_str
    ]
    past_deadlines = [
        serialize(d) for d in deadlines_raw
        if str(d.get("date", ""))[:10] < now_str
    ]

    case = await db.cases.find_one({"_id": ObjectId(case_id)})
    next_action = (case or {}).get("next_action", "")
    last_action = ""
    last_event = await db.audit_logs.find_one({"case_id": case_id}, sort=[("created_at", -1)])
    if last_event:
        label_map = {
            "CREATE_NOTE": "Σημείωση",
            "CREATE_DEADLINE": "Προθεσμία",
            "UPDATE_STATUS": "Ενημέρωση κατάστασης",
            "CREATE_INVOICE": "Τιμολόγιο",
            "CREATE_HEARING": "Δικάσιμος",
            "PORTAL_UPLOAD": "Ανέβασμα εγγράφου",
        }
        act = last_event.get("action", "")
        last_action = label_map.get(act, act)

    return {
        "upcoming_hearings":  upcoming_hearings[:5],
        "past_hearings":      past_hearings[-3:],
        "upcoming_deadlines": upcoming_deadlines[:5],
        "past_deadlines":     past_deadlines[-3:],
        "next_action":        next_action,
        "last_action":        last_action,
    }

@app.get("/api/portal/financials")
async def portal_get_financials(user=Depends(get_portal_user)):
    """Detailed financial breakdown visible to the portal client."""
    case_id = user["case_id"]
    permissions = user.get("permissions", [])
    if "total_fees" not in permissions and "outstanding_balance" not in permissions:
        raise HTTPException(403, "Δεν έχετε πρόσβαση στα οικονομικά")

    invoices = await db.invoices.find({"case_id": case_id}).sort("date", -1).to_list(None)
    financials = await db.financials.find({"case_id": case_id}).sort("date", -1).to_list(None)

    total = sum(float(i.get("total_payable", i.get("total", 0))) for i in invoices)
    paid  = sum(float(i.get("amount_paid", 0)) for i in invoices)

    inv_list = []
    for i in invoices:
        inv_list.append({
            "id":             str(i["_id"]),
            "invoice_number": i.get("invoice_number", ""),
            "description":    i.get("description") or i.get("subject", "—"),
            "total":          float(i.get("total_payable", i.get("total", 0))),
            "paid":           float(i.get("amount_paid", 0)),
            "status":         i.get("payment_status", "pending"),
            "date":           str(i.get("date", i.get("created_at", "")))[:10],
        })

    fee_list = []
    for f in financials:
        fee_list.append({
            "id":          str(f["_id"]),
            "description": f.get("description", "—"),
            "amount":      float(f.get("amount", 0)),
            "type":        f.get("type", "fee"),
            "date":        str(f.get("date", f.get("created_at", "")))[:10],
        })

    settings_doc = await db.settings.find_one({"_id": "global"}) or {}
    bank_accounts = settings_doc.get("bank_accounts", [])

    return {
        "total":         round(total, 2),
        "paid":          round(paid, 2),
        "outstanding":   round(total - paid, 2),
        "invoices":      inv_list,
        "fee_entries":   fee_list,
        "bank_accounts": bank_accounts,
    }

@app.get("/api/admin/portal-documents")
async def list_portal_documents(user=Depends(require_role(UserRole.ADMIN, UserRole.LAWYER))):
    """List portal-uploaded documents awaiting lawyer review."""
    q: dict = {"uploaded_by": "portal_client"}
    if user["role"] == UserRole.LAWYER.value:
        case_ids = [str(c["_id"]) async for c in db.cases.find({"assigned_lawyer_id": user["id"]}, {"_id": 1})]
        q["case_id"] = {"$in": case_ids}
    docs = await db.documents.find(q).sort("created_at", -1).to_list(None)
    result = []
    for d in docs:
        s = serialize(d)
        case = await db.cases.find_one({"_id": make_id(s.get("case_id", ""))})
        s["case_title"]  = case.get("title", "—") if case else "—"
        s["case_number"] = case.get("case_number", "") if case else ""
        result.append(s)
    return result

@app.post("/api/admin/portal-documents/{doc_id}/approve")
async def approve_portal_document(doc_id: str, user=Depends(require_role(UserRole.ADMIN, UserRole.LAWYER))):
    """Lawyer approves portal document → uploads to Google Drive."""
    doc = await db.documents.find_one({"_id": make_id(doc_id)})
    if not doc:
        raise HTTPException(404, "Έγγραφο δεν βρέθηκε")

    drive_link = None
    try:
        from gdrive_service import is_configured, upload_document as drive_upload
        if is_configured():
            case = await db.cases.find_one({"_id": make_id(doc["case_id"])})
            client_name = doc.get("uploaded_by_name", "Πελάτης")
            year = str(datetime.utcnow().year)
            case_number = case.get("case_number", "") if case else ""
            case_title  = (case.get("title") or "Υπόθεση")[:30] if case else "Υπόθεση"
            case_folder = f"{case_number} - {case_title}" if case_number else case_title

            with open(doc["file_path"], "rb") as fh:
                file_bytes = fh.read()

            drive_result = drive_upload(
                file_bytes=file_bytes,
                filename=doc["filename"],
                mime_type=doc.get("mime_type", "application/octet-stream"),
                year=year,
                client_name=client_name[:50],
                case_folder=case_folder,
            )
            drive_link = drive_result.get("folder_link") or drive_result.get("file_link")
    except Exception as e:
        logger.warning(f"Drive upload failed: {e}")

    await db.documents.update_one(
        {"_id": make_id(doc_id)},
        {"$set": {
            "status":      "approved",
            "reviewed_by": user["id"],
            "reviewed_at": datetime.utcnow(),
            "drive_link":  drive_link,
        }}
    )
    await audit("APPROVE_PORTAL_DOC", user["id"], "document", doc_id)
    return {"ok": True, "drive_link": drive_link}

@app.post("/api/admin/portal-documents/{doc_id}/reject")
async def reject_portal_document(doc_id: str, user=Depends(require_role(UserRole.ADMIN, UserRole.LAWYER))):
    """Lawyer rejects (ignores) portal document."""
    await db.documents.update_one(
        {"_id": make_id(doc_id)},
        {"$set": {"status": "rejected", "reviewed_by": user["id"], "reviewed_at": datetime.utcnow()}}
    )
    await audit("REJECT_PORTAL_DOC", user["id"], "document", doc_id)
    return {"ok": True}

# ── Admin Portal Management ──
@app.post("/api/admin/clients/{client_id}/generate-portal-access")
async def generate_portal_access(client_id: str, req: PortalAccessRequest, user=Depends(require_role(UserRole.ADMIN))):
    """Generate portal access code for a specific client case."""
    try:
        ObjectId(client_id)
    except InvalidId:
        raise HTTPException(400, "Μη έγκυρο client ID")

    # Get client name
    client = await db.clients.find_one({"_id": make_id(client_id)})
    if not client:
        raise HTTPException(404, "Εντολέας δεν βρέθηκε")
    client_name = client.get("full_name") or client.get("name", "")

    # Use specific case if provided, else first active case
    if req.case_id:
        case = await db.cases.find_one({"_id": make_id(req.case_id)})
        if not case:
            raise HTTPException(404, "Υπόθεση δεν βρέθηκε")
    else:
        cases = await db.cases.find(
            {"$or": [{"client_id": client_id}, {"client_ids": client_id}],
             "status": {"$in": ["active", "pending"]}}
        ).to_list(None)
        if not cases:
            raise HTTPException(404, "Δεν υπάρχουν ενεργές υποθέσεις")
        case = cases[0]

    portal_code = secrets.token_urlsafe(12)
    access_record = {
        "client_id":    client_id,
        "client_name":  client_name,
        "case_id":      str(case["_id"]),
        "case_title":   case.get("title") or case.get("offense") or "—",
        "case_number":  case.get("case_number") or "",
        "case_category": case.get("category") or case.get("legal_category") or "",
        "portal_code":  portal_code,
        "permissions":     req.permissions,
        "is_active":       True,
        "created_at":      datetime.utcnow(),
        "expires_at":      datetime.utcnow() + timedelta(days=PORTAL_CODE_TTL_DAYS),
        "created_by":      user["id"],
        "client_email":    client.get("email") or "",
        "mandate_accepted": False,
    }
    result = await db.portal_access.insert_one(access_record)

    await audit("CREATE_PORTAL_ACCESS", user["id"], "portal", str(result.inserted_id))
    return {
        "ok":          True,
        "portal_code": portal_code,
        "case_id":     str(case["_id"]),
        "case_title":  case.get("title"),
        "client_name": client_name,
    }

@app.patch("/api/admin/cases/{case_id}/portal-permissions")
async def update_portal_permissions(case_id: str, req: PortalAccessRequest, user=Depends(require_role(UserRole.ADMIN))):
    """Update what fields client can see"""
    portal_access = await db.portal_access.find_one({"case_id": case_id})
    if not portal_access:
        raise HTTPException(404, "Portal access not found for this case")

    await db.portal_access.update_one(
        {"_id": portal_access["_id"]},
        {"$set": {"permissions": req.permissions, "updated_at": datetime.utcnow()}}
    )

    await audit("UPDATE_PORTAL_PERMISSIONS", user["id"], "portal", case_id)
    return {"ok": True, "message": "Portal permissions updated"}

@app.get("/api/admin/portal-access")
async def list_portal_access(user=Depends(require_role(UserRole.ADMIN))):
    """List all portal access codes"""
    codes = await db.portal_access.find().sort("created_at", -1).to_list(None)
    return [serialize(c) for c in codes]

@app.delete("/api/admin/portal-access/{code_id}")
async def delete_portal_access(code_id: str, user=Depends(require_role(UserRole.ADMIN))):
    """Delete a portal access code"""
    result = await db.portal_access.delete_one({"_id": make_id(code_id)})
    if result.deleted_count == 0:
        raise HTTPException(404, "Portal access not found")
    await audit("DELETE_PORTAL_ACCESS", user["id"], "portal", code_id)
    return {"ok": True}

@app.get("/api/admin/portal-reset-requests")
async def list_portal_reset_requests(user=Depends(require_role(UserRole.ADMIN))):
    """List all portal password reset requests"""
    requests = await db.portal_reset_requests.find({"used": False}).sort("created_at", -1).to_list(None)
    return [serialize(r) for r in requests]

@app.post("/api/admin/portal-reset-requests/{request_id}/approve")
async def approve_portal_reset(request_id: str, user=Depends(require_role(UserRole.ADMIN))):
    """Approve reset: generate new portal code and notify client"""
    req = await db.portal_reset_requests.find_one({"_id": make_id(request_id)})
    if not req:
        raise HTTPException(404, "Request not found")

    case_id = req.get("case_id")
    portal_access = await db.portal_access.find_one({"case_id": case_id})
    if not portal_access:
        raise HTTPException(404, "Portal access not found for this case")

    new_code = secrets.token_urlsafe(12)
    await db.portal_access.update_one(
        {"_id": portal_access["_id"]},
        {"$set": {"portal_code": new_code, "updated_at": datetime.utcnow()}}
    )
    await db.portal_reset_requests.update_one(
        {"_id": make_id(request_id)},
        {"$set": {"used": True, "resolved_at": datetime.utcnow(), "resolved_by": user["id"]}}
    )
    await audit("APPROVE_PORTAL_RESET", user["id"], "portal", case_id)
    return {"ok": True, "new_portal_code": new_code}

@app.post("/api/admin/portal-reset-requests/{request_id}/reject")
async def reject_portal_reset(request_id: str, user=Depends(require_role(UserRole.ADMIN))):
    """Reject a portal password reset request"""
    result = await db.portal_reset_requests.update_one(
        {"_id": make_id(request_id)},
        {"$set": {"used": True, "resolved_at": datetime.utcnow(), "resolved_by": user["id"], "rejected": True}}
    )
    if result.matched_count == 0:
        raise HTTPException(404, "Request not found")
    await audit("REJECT_PORTAL_RESET", user["id"], "portal", request_id)
    return {"ok": True}

# ── Admin: Portal Message Inbox ───────────────────────────────────────────────

@app.get("/api/admin/portal-messages")
async def list_portal_messages(user=Depends(require_role(UserRole.ADMIN, UserRole.LAWYER))):
    """List all portal messages with case/client info for the lawyer inbox."""
    q: dict = {}
    if user["role"] == UserRole.LAWYER.value:
        case_ids = [str(c["_id"]) async for c in db.cases.find({"assigned_lawyer_id": user["id"]}, {"_id": 1})]
        q["case_id"] = {"$in": case_ids}
    msgs = await db.portal_messages.find(q).sort("created_at", -1).to_list(None)
    result = []
    for m in msgs:
        s = serialize(m)
        case = await db.cases.find_one({"_id": make_id(s.get("case_id", ""))})
        s["case_title"]  = (case.get("title") or "—") if case else "—"
        s["case_number"] = (case.get("case_number") or "") if case else ""
        result.append(s)
    return result

@app.patch("/api/admin/portal-messages/{msg_id}/read")
async def mark_portal_message_read(msg_id: str, user=Depends(require_role(UserRole.ADMIN, UserRole.LAWYER))):
    """Mark a portal message as read."""
    await db.portal_messages.update_one(
        {"_id": make_id(msg_id)},
        {"$set": {"read": True, "read_by_lawyer": True, "read_at": datetime.utcnow()}}
    )
    return {"ok": True}

class PortalReplyRequest(BaseModel):
    content: str

@app.post("/api/admin/portal-messages/{msg_id}/reply")
async def reply_portal_message(msg_id: str, req: PortalReplyRequest, user=Depends(require_role(UserRole.ADMIN, UserRole.LAWYER))):
    """Lawyer replies to a portal message — stores reply and emails the client."""
    import smtplib
    from email.mime.text import MIMEText
    from email.mime.multipart import MIMEMultipart

    msg = await db.portal_messages.find_one({"_id": make_id(msg_id)})
    if not msg:
        raise HTTPException(404, "Μήνυμα δεν βρέθηκε")

    reply = {
        "author":     user.get("full_name") or user.get("name") or "Δικηγόρος",
        "content":    sanitize_string(req.content),
        "created_at": datetime.utcnow().isoformat(),
        "role":       "lawyer",
    }

    await db.portal_messages.update_one(
        {"_id": make_id(msg_id)},
        {
            "$push": {"replies": reply},
            "$set":  {"read": True, "read_by_lawyer": True},
        }
    )

    # Send email to client if we have their email
    case_id = msg.get("case_id", "")
    client_email = ""
    client_name  = msg.get("client_name", "Πελάτης")
    if case_id:
        pa = await db.portal_access.find_one({"case_id": case_id})
        if pa:
            client_obj = await db.clients.find_one({"_id": make_id(pa.get("client_id", ""))})
            if client_obj:
                client_email = client_obj.get("email", "")

    if client_email:
        try:
            settings_doc = await db.settings.find_one({"_id": "global"}) or {}
            smtp_host = settings_doc.get("smtp_host") or os.getenv("SMTP_HOST", "")
            smtp_port = int(settings_doc.get("smtp_port") or os.getenv("SMTP_PORT", "587"))
            smtp_user = settings_doc.get("smtp_user") or os.getenv("SMTP_USER", "")
            smtp_pass = settings_doc.get("smtp_pass") or os.getenv("SMTP_PASS", "")
            from_email = settings_doc.get("notification_email") or smtp_user or FIRM_EMAIL_DISPLAY

            noreply_addr = os.getenv("SMTP_FROM", from_email)
            lawyer_name  = user.get("full_name") or user.get("name") or reply["author"]
            lawyer_email = user.get("email") or ""

            if smtp_host and smtp_user:
                body_html = f"""
<html><body style="font-family:Arial,sans-serif;color:#222;max-width:600px;margin:auto">
<div style="background:#071220;padding:24px;border-radius:8px;margin-bottom:24px">
  <h2 style="color:#C6A75E;margin:0">NOMOS ONE — Σκοτάνης &amp; Συνεργάτες</h2>
  <p style="color:#8aa0b8;margin:4px 0 0">Απάντηση στο μήνυμά σας</p>
</div>
<p>Αγαπητέ/ή {client_name},</p>
<p>Λάβατε απάντηση από τον/την {lawyer_name}:</p>
<div style="background:#f4f4f4;border-left:4px solid #C6A75E;padding:16px;border-radius:4px;margin:16px 0">
  <p style="margin:0;white-space:pre-wrap">{reply['content']}</p>
</div>
<p>Μπορείτε να συνδεθείτε στην Πύλη Πελάτη για να δείτε όλη την αλληλογραφία.</p>
{f'<p>Για άμεση επικοινωνία: <a href="mailto:{lawyer_email}">{lawyer_email}</a></p>' if lawyer_email else ''}
<p style="color:#888;font-size:12px;margin-top:32px">Σκοτάνης &amp; Συνεργάτες — Εμπιστευτική Επικοινωνία</p>
</body></html>"""
                mail = MIMEMultipart("alternative")
                mail["Subject"]  = f"Απάντηση από {lawyer_name} — Πύλη Πελάτη NOMOS ONE"
                mail["From"]     = f"{lawyer_name} — Σκοτάνης & Συνεργάτες <{noreply_addr}>"
                mail["To"]       = f"{client_name} <{client_email}>"
                if lawyer_email:
                    mail["Reply-To"] = lawyer_email
                mail.attach(MIMEText(body_html, "html", "utf-8"))
                with smtplib.SMTP(smtp_host, smtp_port, timeout=10) as srv:
                    srv.ehlo(); srv.starttls(); srv.login(smtp_user, smtp_pass)
                    srv.sendmail(noreply_addr, client_email, mail.as_string())
                await db.email_logs.insert_one({
                    "to": client_email, "subject": f"Απάντηση από {lawyer_name}",
                    "sent_by": user["id"], "sent_at": datetime.utcnow(), "status": "sent"
                })
        except Exception as e:
            logger.warning(f"Reply email failed: {e}")

    await audit("PORTAL_REPLY", user["id"], "portal", msg_id)
    return {"ok": True, "email_sent": bool(client_email)}


# ══════════════════════════════════════════════════════════════════════════════
# CLIENTS
# ══════════════════════════════════════════════════════════════════════════════
class ClientRequest(BaseModel):
    full_name: str; email: Optional[str] = None; phone: Optional[str] = None
    address: Optional[str] = None; afm: Optional[str] = None; client_type: Optional[str] = None
    profession: Optional[str] = None; notes: Optional[str] = None; is_active: Optional[bool] = None

@app.get("/api/clients")
async def list_clients(user=Depends(get_current_user)):
    clients = await db.clients.find({}).to_list(None)
    result = []
    for c in clients:
        s = serialize(c)
        cid = s["id"]
        s["cases_count"] = await db.cases.count_documents(
            {"$or": [{"client_id": cid}, {"client_ids": cid}]}
        )
        result.append(s)
    return result

@app.get("/api/clients/{client_id}/cases")
async def get_client_cases(client_id: str, user=Depends(get_current_user)):
    """All cases where this client appears (primary or co-client)."""
    cases = await db.cases.find(
        {"$or": [{"client_id": client_id}, {"client_ids": client_id}]}
    ).sort("created_at", -1).to_list(None)
    result = []
    for c in cases:
        s = serialize(c)
        s["client_name"] = await get_client_name(s.get("client_id", ""))
        # Populate all client names for multi-client cases
        extra_ids = s.get("client_ids", [])
        if len(extra_ids) > 1:
            names = []
            for cid in extra_ids:
                n = await get_client_name(cid)
                if n and n not in names:
                    names.append(n)
            s["client_names"] = names
        result.append(s)
    return result

@app.get("/api/clients/{client_id}")
async def get_client(client_id: str, user=Depends(get_current_user)):
    doc = await db.clients.find_one({"_id": make_id(client_id)})
    if not doc: raise HTTPException(404, "Ο εντολέας δεν βρέθηκε")
    s = serialize(doc)
    s["cases_count"] = await db.cases.count_documents(
        {"$or": [{"client_id": client_id}, {"client_ids": client_id}]}
    )
    return s

@app.get("/api/clients/{client_id}/360")
async def client_360(client_id: str, user=Depends(get_current_user)):
    """Full 360° profile for a client — all linked cases, hearings, deadlines, financials."""
    client = await db.clients.find_one({"_id": make_id(client_id)})
    if not client:
        raise HTTPException(404, "Εντολέας δεν βρέθηκε")

    cases = await db.cases.find(
        {"$or": [{"client_id": client_id}, {"client_ids": client_id}]}
    ).sort("created_at", -1).to_list(None)
    case_ids = [str(c["_id"]) for c in cases]

    enriched_cases = []
    for c in cases:
        s = serialize(c)
        s["client_name"] = await get_client_name(s.get("client_id", ""))
        s["assigned_lawyer_name"] = await get_user_name(s.get("assigned_lawyer_id", ""))
        enriched_cases.append(s)

    if case_ids:
        hearings_raw   = await db.hearings.find({"case_id": {"$in": case_ids}}).sort("hearing_date", 1).to_list(None)
        deadlines_raw  = await db.deadlines.find({"case_id": {"$in": case_ids}}).sort("date", 1).to_list(None)
        invoices_raw   = await db.invoices.find({"case_id": {"$in": case_ids}}).to_list(None)
        doc_count      = await db.documents.count_documents({"case_id": {"$in": case_ids}})
        note_count     = await db.notes.count_documents({"case_id": {"$in": case_ids}})
    else:
        hearings_raw = deadlines_raw = invoices_raw = []
        doc_count = note_count = 0

    case_map = {str(c["_id"]): c for c in cases}
    def _case_label(cid: str) -> str:
        c = case_map.get(cid, {})
        return c.get("offense") or c.get("title") or cid[:8]

    hearings = []
    for h in hearings_raw:
        s = serialize(h); s["case_label"] = _case_label(s.get("case_id", "")); hearings.append(s)

    deadlines = []
    for d in deadlines_raw:
        s = serialize(d); s["case_label"] = _case_label(s.get("case_id", "")); deadlines.append(s)

    total_invoiced = sum(float(i.get("total_payable", i.get("total", 0))) for i in invoices_raw)
    total_paid     = sum(float(i.get("amount_paid", 0)) for i in invoices_raw)

    case_fin: dict = {}
    for i in invoices_raw:
        cid = i.get("case_id", "")
        if cid not in case_fin:
            case_fin[cid] = {"invoiced": 0.0, "paid": 0.0}
        case_fin[cid]["invoiced"] += float(i.get("total_payable", i.get("total", 0)))
        case_fin[cid]["paid"]     += float(i.get("amount_paid", 0))

    now_str = datetime.utcnow().isoformat()[:10]
    upcoming_hearings  = [h for h in hearings  if str(h.get("hearing_date", ""))[:10] >= now_str and h.get("status") not in ("completed", "cancelled")]
    upcoming_deadlines = [d for d in deadlines if str(d.get("date", ""))[:10] >= now_str]

    return {
        "client": serialize(client),
        "stats": {
            "cases_total":        len(cases),
            "cases_active":       sum(1 for c in cases if not str(c.get("status","")).startswith("closed") and c.get("status") != "archived"),
            "hearings_upcoming":  len(upcoming_hearings),
            "deadlines_upcoming": len(upcoming_deadlines),
            "doc_count":          doc_count,
            "note_count":         note_count,
            "total_invoiced":     total_invoiced,
            "total_paid":         total_paid,
            "balance":            total_invoiced - total_paid,
        },
        "cases":           enriched_cases,
        "hearings":        hearings,
        "deadlines":       deadlines,
        "case_financials": case_fin,
        "invoices":        [serialize(i) for i in invoices_raw],
    }


@app.post("/api/clients", status_code=201)
async def create_client(req: ClientRequest, user=Depends(require_role(UserRole.ADMIN, UserRole.SECRETARY))):
    if req.phone and not validate_phone(req.phone): raise HTTPException(400, "Μη έγκυρος αριθμός τηλεφώνου")
    if req.afm and not validate_tax_id(req.afm): raise HTTPException(400, "Μη έγκυρο ΑΦΜ (9 ψηφία)")
    if req.email and not re.match(r"[^@]+@[^@]+\.[^@]+", req.email.strip()): raise HTTPException(400, "Μη έγκυρο email")
    doc = {"full_name": sanitize_string(req.full_name), "email": req.email.strip().lower() if req.email else None,
           "phone": req.phone.strip() if req.phone else None, "address": sanitize_string(req.address) if req.address else None,
           "afm": req.afm.strip() if req.afm else None, "client_type": req.client_type or "individual",
           "notes": sanitize_string(req.notes) if req.notes else None,
           "created_at": datetime.utcnow(), "created_by": user["id"]}
    result = await db.clients.insert_one(doc)
    await audit("CREATE_CLIENT", user["id"], "client", str(result.inserted_id))
    doc["_id"] = result.inserted_id; return serialize(doc)

@app.put("/api/clients/{client_id}")
async def update_client(client_id: str, req: ClientRequest, user=Depends(require_role(UserRole.ADMIN, UserRole.SECRETARY))):
    if req.phone and not validate_phone(req.phone): raise HTTPException(400, "Μη έγκυρος αριθμός τηλεφώνου")
    if req.afm and not validate_tax_id(req.afm): raise HTTPException(400, "Μη έγκυρο ΑΦΜ (9 ψηφία)")
    if req.email and not re.match(r"[^@]+@[^@]+\.[^@]+", req.email.strip()): raise HTTPException(400, "Μη έγκυρο email")
    data = {"full_name": sanitize_string(req.full_name), "email": req.email.strip().lower() if req.email else None,
            "phone": req.phone.strip() if req.phone else None, "address": sanitize_string(req.address) if req.address else None,
            "afm": req.afm.strip() if req.afm else None, "client_type": req.client_type,
            "notes": sanitize_string(req.notes) if req.notes else None,
            "updated_at": datetime.utcnow()}
    await db.clients.update_one({"_id": make_id(client_id)}, {"$set": data})
    await audit("UPDATE_CLIENT", user["id"], "client", client_id)
    return {"ok": True}

@app.get("/api/clients/{client_id}/export")
async def export_client(client_id: str, user=Depends(require_role(UserRole.ADMIN, UserRole.SECRETARY))):
    doc = await db.clients.find_one({"_id": make_id(client_id)})
    if not doc: raise HTTPException(404, "Ο εντολέας δεν βρέθηκε")
    cases = await db.cases.find({"$or": [{"client_id": client_id}, {"client_ids": client_id}]}).to_list(None)
    export_data = {"client": serialize(doc), "cases": [serialize(c) for c in cases], "exported_at": datetime.utcnow().isoformat()}
    await audit("EXPORT_CLIENT", user["id"], "client", client_id)
    content = json.dumps(export_data, ensure_ascii=False, indent=2, default=str)
    return StreamingResponse(iter([content.encode()]), media_type="application/json",
                             headers={"Content-Disposition": f"attachment; filename=client_{client_id}.json"})

# ══════════════════════════════════════════════════════════════════════════════
# CASES — atomic case number
# ══════════════════════════════════════════════════════════════════════════════
class CaseRequest(BaseModel):
    title: str; client_id: str; assigned_lawyer_id: Optional[str] = None
    status: CaseStatus = CaseStatus.ACTIVE; legal_category: Optional[str] = None; category: Optional[str] = None
    next_action: Optional[str] = None; next_action_date: Optional[datetime] = None; court: Optional[str] = None
    description: Optional[str] = None; summary: Optional[str] = None
    offense: Optional[str] = None; law_articles: Optional[str] = None

class CaseStatusUpdate(BaseModel):
    status: CaseStatus; next_action: Optional[str] = None; next_action_date: Optional[datetime] = None

async def case_number_gen():
    year = datetime.utcnow().year
    result = await db.counters.find_one_and_update(
        {"_id": "case_number"}, {"$set": {"year": year}, "$inc": {"seq": 1}},
        upsert=True, return_document=True
    )
    if result.get("year") != year:
        result = await db.counters.find_one_and_update(
            {"_id": "case_number"}, {"$set": {"year": year, "seq": 1}}, return_document=True)
    return f"{year}-{str(result['seq']).zfill(4)}"

def is_locked(case): return case.get("status") in [CaseStatus.CLOSED.value, CaseStatus.ARCHIVED.value]


# ── Pending Intakes (Telegram / Email intake queue) ───────────────────────────

@app.get("/api/pending-intakes")
async def list_pending_intakes(user=Depends(get_current_user)):
    items = await db.pending_intakes.find(
        {"status": {"$in": ["pending", "rejected"]}}
    ).sort("submitted_at", -1).to_list(100)
    return [serialize(i) for i in items]


@app.post("/api/pending-intakes/{intake_id}/approve")
async def approve_pending_intake(intake_id: str, user=Depends(require_role(UserRole.ADMIN, UserRole.LAWYER))):
    intake = await db.pending_intakes.find_one({"_id": make_id(intake_id)})
    if not intake:
        raise HTTPException(404, "Δεν βρέθηκε")
    if intake["status"] == "approved":
        raise HTTPException(400, "Έχει ήδη εγκριθεί")

    now = datetime.utcnow()
    extracted = intake.get("extracted") or {}
    cl = extracted.get("client") or {}
    cs = extracted.get("case") or {}

    # Support multiple clients selected via Telegram multi-select
    client_names = intake.get("client_names") or []
    primary_name = intake.get("client_name") or cl.get("full_name", "Άγνωστος")
    if not client_names:
        client_names = [primary_name]

    async def _find_or_create_client(name: str, afm: str | None, is_primary: bool) -> str:
        existing = None
        if is_primary and afm:
            existing = await db.clients.find_one({"afm": afm})
        if not existing and name:
            existing = await db.clients.find_one(
                {"full_name": {"$regex": re.escape(name[:15]), "$options": "i"}}
            )
        if existing:
            return str(existing["_id"])
        doc = {
            "full_name": name,
            "afm": afm if is_primary else None,
            "phone": cl.get("phone", "") if is_primary else "",
            "email": cl.get("email", "") if is_primary else "",
            "address": cl.get("address", "") if is_primary else "",
            "client_type": cl.get("client_type", "individual"),
            "is_active": True,
            "source": "intake_channel",
            "created_at": now,
            "created_by": user["id"],
        }
        cr = await db.clients.insert_one(doc)
        cid = str(cr.inserted_id)
        await audit("CREATE_CLIENT", user["id"], "client", cid)
        return cid

    client_ids = []
    for i, name in enumerate(client_names):
        cid = await _find_or_create_client(name, cl.get("afm") if i == 0 else None, i == 0)
        client_ids.append(cid)

    primary_client_id = client_ids[0]
    cn = await case_number_gen()
    case_title = cs.get("title") or f"Υπόθεση {primary_name}"
    case_doc = {
        "title": case_title,
        "client_id": primary_client_id,
        "client_ids": client_ids,
        "assigned_lawyer_id": user["id"],
        "status": "active",
        "legal_category": cs.get("category", "αστικό"),
        "court": cs.get("court", ""),
        "description": extracted.get("summary") or cs.get("summary", ""),
        "case_number": cn,
        "opposing_party": cs.get("opposing_party", ""),
        "source": "intake_channel",
        "review_status": "approved",
        "ai_confidence": extracted.get("confidence", "low"),
        "ai_key_facts": extracted.get("key_facts", []),
        "source_files": intake.get("filenames", []),
        "created_at": now,
        "created_by": user["id"],
        "updated_at": now,
        "last_activity": now,
    }
    cr2 = await db.cases.insert_one(case_doc)
    case_id = str(cr2.inserted_id)
    await audit("CREATE_CASE", user["id"], "case", case_id)

    await db.pending_intakes.update_one(
        {"_id": make_id(intake_id)},
        {"$set": {
            "status": "approved",
            "reviewed_by": user["id"],
            "reviewed_at": now,
            "created_client_ids": client_ids,
            "created_client_id": primary_client_id,
            "created_case_id": case_id,
            "created_case_number": cn,
        }}
    )
    return {"status": "approved", "case_number": cn, "case_id": case_id,
            "client_id": primary_client_id, "client_ids": client_ids}


@app.post("/api/pending-intakes/{intake_id}/reject")
async def reject_pending_intake(
    intake_id: str,
    body: dict = {},
    user=Depends(require_role(UserRole.ADMIN, UserRole.LAWYER))
):
    intake = await db.pending_intakes.find_one({"_id": make_id(intake_id)})
    if not intake:
        raise HTTPException(404, "Δεν βρέθηκε")
    await db.pending_intakes.update_one(
        {"_id": make_id(intake_id)},
        {"$set": {
            "status": "rejected",
            "reviewed_by": user["id"],
            "reviewed_at": datetime.utcnow(),
            "notes": body.get("notes", ""),
        }}
    )
    return {"status": "rejected"}

@app.get("/api/cases")
async def list_cases(user=Depends(get_current_user), status: Optional[str] = None):
    query = {}
    if status: query["status"] = status
    cases = await db.cases.find(query).sort("created_at", -1).to_list(None)
    result = []
    for c in cases:
        s = serialize(c)
        s["assigned_lawyer_name"] = await get_user_name(s.get("assigned_lawyer_id", ""))
        s["client_name"] = await get_client_name(s.get("client_id", ""))
        # For multi-client cases, populate all client names
        extra_ids = s.get("client_ids", [])
        if len(extra_ids) > 1:
            names = []
            for cid in extra_ids:
                n = await get_client_name(cid)
                if n and n not in names:
                    names.append(n)
            s["client_names"] = names
        result.append(s)
    return result

@app.get("/api/cases/stagnant")
async def stagnant_cases(user=Depends(get_current_user)):
    sd = datetime.utcnow() - timedelta(days=STAGNANT_DAYS)
    q = {"status": "active", "$or": [{"last_activity": {"$lt": sd}},
         {"last_activity": {"$exists": False}, "created_at": {"$lt": sd}}]}
    cases = await db.cases.find(q).to_list(None)
    result = []
    for c in cases:
        s = serialize(c); s["assigned_lawyer_name"] = await get_user_name(s.get("assigned_lawyer_id", "")); result.append(s)
    return result
@app.get("/api/cases/{case_id}")
async def get_case(case_id: str, user=Depends(get_current_user)):
    case = await db.cases.find_one({"_id": make_id(case_id)})
    if not case: raise HTTPException(404, "Η υπόθεση δεν βρέθηκε")

    s = serialize(case)
    s["assigned_lawyer_name"] = await get_user_name(s.get("assigned_lawyer_id", ""))
    s["client_name"] = await get_client_name(s.get("client_id", ""))
    # Populate full client objects for all linked clients
    client_ids = s.get("client_ids") or ([s["client_id"]] if s.get("client_id") else [])
    clients_detail = []
    for cid in client_ids:
        cd = await db.clients.find_one({"_id": make_id(cid)})
        if cd:
            clients_detail.append(serialize(cd))
    if clients_detail:
        s["clients"] = clients_detail
    return s

@app.post("/api/cases", status_code=201)
async def create_case(req: CaseRequest, user=Depends(get_current_user)):
    # Use current user as lawyer if not specified
    lawyer_id = req.assigned_lawyer_id or user["id"]
    lawyer = await db.users.find_one({"_id": make_id(lawyer_id)})
    if not lawyer: raise HTTPException(404, "Ο δικηγόρος δεν βρέθηκε")
    cl = await db.clients.find_one({"_id": make_id(req.client_id)})
    if not cl: raise HTTPException(404, "Ο εντολέας δεν βρέθηκε")
    cn = await case_number_gen()
    doc = {"title": sanitize_string(req.title), "client_id": req.client_id,
           "assigned_lawyer_id": lawyer_id, "status": req.status.value,
           "next_action": sanitize_string(req.next_action or ""), "next_action_date": req.next_action_date,
           "legal_category": req.category or req.legal_category,
           "offense": sanitize_string(req.offense) if req.offense else None,
           "law_articles": sanitize_string(req.law_articles) if req.law_articles else None,
           "court": sanitize_string(req.court) if req.court else None,
           "description": sanitize_string(req.summary or req.description or "") if (req.summary or req.description) else None,
           "case_number": cn, "created_at": datetime.utcnow(), "created_by": user["id"],
           "updated_at": datetime.utcnow(), "last_activity": datetime.utcnow()}
    result = await db.cases.insert_one(doc)
    await audit("CREATE_CASE", user["id"], "case", str(result.inserted_id))
    doc["_id"] = result.inserted_id; s = serialize(doc)
    s["assigned_lawyer_name"] = lawyer.get("full_name") or lawyer.get("name", ""); s["client_name"] = cl.get("full_name") or cl.get("name", "")
    return s

@app.put("/api/cases/{case_id}")
async def update_case(case_id: str, req: CaseRequest, user=Depends(require_role(UserRole.ADMIN, UserRole.SECRETARY))):
    case = await db.cases.find_one({"_id": make_id(case_id)})
    if not case: raise HTTPException(404, "Η υπόθεση δεν βρέθηκε")
    if is_locked(case): raise HTTPException(423, "Η υπόθεση είναι κλειστή/αρχειοθετημένη")
    data = {"title": sanitize_string(req.title), "client_id": req.client_id,
            "assigned_lawyer_id": req.assigned_lawyer_id, "status": req.status.value,
            "next_action": sanitize_string(req.next_action), "next_action_date": req.next_action_date,
            "court": sanitize_string(req.court) if req.court else None,
            "description": sanitize_string(req.description) if req.description else None,
            "updated_at": datetime.utcnow()}
    await db.cases.update_one({"_id": make_id(case_id)}, {"$set": data})
    await audit("UPDATE_CASE", user["id"], "case", case_id)
    return {"ok": True}

@app.patch("/api/cases/{case_id}/status")
async def update_case_status(case_id: str, req: CaseStatusUpdate, user=Depends(require_role(UserRole.ADMIN, UserRole.SECRETARY))):
    case = await db.cases.find_one({"_id": make_id(case_id)})
    if not case: raise HTTPException(404, "Η υπόθεση δεν βρέθηκε")
    if is_locked(case) and req.status != CaseStatus.ACTIVE:
        if user["role"] != UserRole.ADMIN.value: raise HTTPException(423, "Η υπόθεση είναι κλειστή")
    update = {"status": req.status.value, "updated_at": datetime.utcnow()}
    if req.next_action: update["next_action"] = sanitize_string(req.next_action)
    if req.next_action_date: update["next_action_date"] = req.next_action_date
    await db.cases.update_one({"_id": make_id(case_id)}, {"$set": update})
    await audit("STATUS_CHANGE", user["id"], "case", case_id, {"new_status": req.status.value})
    return {"ok": True}

@app.get("/api/cases/{case_id}/export")
async def export_case(case_id: str, user=Depends(require_role(UserRole.ADMIN, UserRole.SECRETARY))):
    case = await db.cases.find_one({"_id": make_id(case_id)})
    if not case: raise HTTPException(404, "Η υπόθεση δεν βρέθηκε")
    notes = await db.notes.find({"case_id": case_id}).to_list(None)
    financials = await db.financials.find({"case_id": case_id}).to_list(None)
    docs_meta = await db.documents.find({"case_id": case_id}).to_list(None)
    import io
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        data = {"case": serialize(case), "notes": [serialize(n) for n in notes],
                "financials": [serialize(f) for f in financials],
                "documents_metadata": [serialize(d) for d in docs_meta],
                "exported_at": datetime.utcnow().isoformat()}
        zf.writestr("case_data.json", json.dumps(data, ensure_ascii=False, indent=2, default=str))
        for dm in docs_meta:
            fp = DOCUMENT_STORAGE_PATH / dm.get("stored_filename", "")
            if fp.exists(): zf.write(fp, f"documents/{dm.get('original_filename', 'file')}")
    buf.seek(0)
    await audit("EXPORT_CASE", user["id"], "case", case_id)
    return StreamingResponse(buf, media_type="application/zip",
                             headers={"Content-Disposition": f"attachment; filename=case_{case_id}.zip"})

# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENTS
# ══════════════════════════════════════════════════════════════════════════════
@app.get("/api/cases/{case_id}/documents")
async def list_documents(case_id: str, user=Depends(get_current_user)):
    await _check_case_access(case_id, user)
    docs = await db.documents.find({"case_id": case_id}).sort("uploaded_at", -1).to_list(None)
    return [serialize(d) for d in docs]

@app.post("/api/cases/{case_id}/documents", status_code=201)
async def upload_document(case_id: str, file: UploadFile = File(...), doc_type: str = Query(...),
    doc_date: str = Query(...), court_authority: Optional[str] = Query(None),
    notes: Optional[str] = Query(None), user=Depends(get_current_user)):
    await _check_case_access(case_id, user)
    case = await db.cases.find_one({"_id": make_id(case_id)})
    if is_locked(case): raise HTTPException(423, "Η υπόθεση είναι κλειστή")
    if file.content_type not in ALLOWED_FILE_TYPES:
        raise HTTPException(400, f"Μη επιτρεπτός τύπος αρχείου: {file.content_type}")
    content = await file.read()
    if len(content) > MAX_FILE_SIZE_MB * 1024 * 1024:
        raise HTTPException(400, f"Το αρχείο υπερβαίνει το όριο {MAX_FILE_SIZE_MB}MB")
    orig = sanitize_string(file.filename) or "unnamed"
    ext = Path(orig).suffix; stored = f"{uuid.uuid4()}{ext}"
    (DOCUMENT_STORAGE_PATH / stored).write_bytes(content)
    meta = {"case_id": case_id, "original_filename": orig, "stored_filename": stored,
            "doc_type": sanitize_string(doc_type), "doc_date": doc_date,
            "court_authority": sanitize_string(court_authority) if court_authority else None,
            "notes": sanitize_string(notes) if notes else None, "content_type": file.content_type,
            "size_bytes": len(content), "uploaded_by": user["id"], "uploaded_at": datetime.utcnow(), "archived": False}
    r = await db.documents.insert_one(meta)
    await db.cases.update_one({"_id": make_id(case_id)}, {"$set": {"last_activity": datetime.utcnow()}})
    await audit("UPLOAD_DOCUMENT", user["id"], "document", str(r.inserted_id), {"filename": orig})
    meta["_id"] = r.inserted_id; return serialize(meta)

@app.get("/api/documents/{doc_id}/download")
async def download_document(doc_id: str, user=Depends(get_current_user)):
    doc = await db.documents.find_one({"_id": make_id(doc_id)})
    if not doc: raise HTTPException(404, "Το έγγραφο δεν βρέθηκε")
    await _check_case_access(doc["case_id"], user)
    fp = DOCUMENT_STORAGE_PATH / doc["stored_filename"]
    if not fp.exists(): raise HTTPException(404, "Αρχείο δεν βρέθηκε στο δίσκο")
    await audit("DOWNLOAD_DOCUMENT", user["id"], "document", doc_id)
    return FileResponse(fp, filename=doc["original_filename"], media_type=doc["content_type"])

@app.patch("/api/documents/{doc_id}/archive")
async def archive_document(doc_id: str, user=Depends(require_role(UserRole.ADMIN, UserRole.SECRETARY))):
    doc = await db.documents.find_one({"_id": make_id(doc_id)})
    if not doc: raise HTTPException(404, "Το έγγραφο δεν βρέθηκε")
    await db.documents.update_one({"_id": make_id(doc_id)}, {"$set": {"archived": True, "archived_at": datetime.utcnow()}})
    await audit("ARCHIVE_DOCUMENT", user["id"], "document", doc_id)
    return {"ok": True}

@app.delete("/api/documents/{doc_id}")
async def delete_document(doc_id: str, user=Depends(require_role(UserRole.ADMIN))):
    doc = await db.documents.find_one({"_id": make_id(doc_id)})
    if not doc: raise HTTPException(404, "Το έγγραφο δεν βρέθηκε")
    if doc.get("archived"): raise HTTPException(423, "Αρχειοθετημένα έγγραφα δεν διαγράφονται")
    fp = DOCUMENT_STORAGE_PATH / doc["stored_filename"]
    if fp.exists(): fp.unlink()
    await db.documents.delete_one({"_id": make_id(doc_id)})
    await audit("DELETE_DOCUMENT", user["id"], "document", doc_id)
    return {"ok": True}

# ══════════════════════════════════════════════════════════════════════════════
# NOTES
# ══════════════════════════════════════════════════════════════════════════════
class NoteRequest(BaseModel):
    content: str

@app.get("/api/cases/{case_id}/notes")
async def list_notes(case_id: str, user=Depends(get_current_user)):
    await _check_case_access(case_id, user)
    return [serialize(n) for n in await db.notes.find({"case_id": case_id}).sort("created_at", -1).to_list(None)]

@app.post("/api/cases/{case_id}/notes", status_code=201)
async def create_note(case_id: str, req: NoteRequest, user=Depends(get_current_user)):
    await _check_case_access(case_id, user)
    case = await db.cases.find_one({"_id": make_id(case_id)})
    if is_locked(case): raise HTTPException(423, "Η υπόθεση είναι κλειστή")
    doc = {"case_id": case_id, "content": sanitize_string(req.content),
           "author_id": user["id"], "author_name": user["name"], "created_at": datetime.utcnow()}
    r = await db.notes.insert_one(doc)
    await db.cases.update_one({"_id": make_id(case_id)}, {"$set": {"last_activity": datetime.utcnow()}})
    await audit("CREATE_NOTE", user["id"], "note", str(r.inserted_id))
    doc["_id"] = r.inserted_id; return serialize(doc)

# ══════════════════════════════════════════════════════════════════════════════
# FINANCIALS
# ══════════════════════════════════════════════════════════════════════════════
class FinancialRequest(BaseModel):
    description: str; amount: float; entry_type: str
    payment_status: PaymentStatus = PaymentStatus.PENDING; payment_method: Optional[str] = None; invoice_number: Optional[str] = None; due_date: Optional[datetime] = None; date: Optional[datetime] = None

@app.get("/api/cases/{case_id}/financials")
async def list_financials(case_id: str, user=Depends(get_current_user)):
    if user["role"] == UserRole.SECRETARY.value: raise HTTPException(403, "Δεν έχετε πρόσβαση στα οικονομικά")
    await _check_case_access(case_id, user)
    entries = await db.financials.find({"case_id": case_id}).sort("date", -1).to_list(None)
    total = sum(e.get("amount", 0) for e in entries)
    total_fees = sum(e.get("amount", 0) for e in entries if e.get("entry_type") == "fee")
    total_expenses = sum(e.get("amount", 0) for e in entries if e.get("entry_type") == "expense")
    return {"entries": [serialize(e) for e in entries], "total": total, "total_fees": total_fees, "total_expenses": total_expenses}

@app.post("/api/cases/{case_id}/financials", status_code=201)
async def create_financial(case_id: str, req: FinancialRequest, user=Depends(require_role(UserRole.ADMIN, UserRole.LAWYER))):
    await _check_case_access(case_id, user)
    case = await db.cases.find_one({"_id": make_id(case_id)})
    if is_locked(case): raise HTTPException(423, "Η υπόθεση είναι κλειστή")
    if req.amount <= 0: raise HTTPException(400, "Το ποσό πρέπει να είναι θετικό")
    if req.entry_type not in ("fee", "expense"): raise HTTPException(400, "Τύπος: 'fee' ή 'expense'")
    doc = {"description": sanitize_string(req.description), "amount": req.amount, "entry_type": req.entry_type,
           "payment_status": req.payment_status.value, "case_id": case_id, "created_by": user["id"],
           "created_at": datetime.utcnow(), "date": req.date or datetime.utcnow()}
    r = await db.financials.insert_one(doc)
    await audit("CREATE_FINANCIAL", user["id"], "financial", str(r.inserted_id))
    doc["_id"] = r.inserted_id; return serialize(doc)

@app.put("/api/cases/{case_id}/financials/{entry_id}")
async def update_financial(case_id: str, entry_id: str, req: FinancialRequest, user=Depends(require_role(UserRole.ADMIN, UserRole.LAWYER))):
    case = await db.cases.find_one({"_id": make_id(case_id)})
    if is_locked(case): raise HTTPException(423, "Η υπόθεση είναι κλειστή")
    if req.amount <= 0: raise HTTPException(400, "Το ποσό πρέπει να είναι θετικό")
    await db.financials.update_one({"_id": make_id(entry_id)}, {"$set": {
        "description": sanitize_string(req.description), "amount": req.amount,
        "entry_type": req.entry_type, "payment_status": req.payment_status.value,
        "date": req.date or datetime.utcnow(), "updated_at": datetime.utcnow()}})
    await audit("UPDATE_FINANCIAL", user["id"], "financial", entry_id)
    return {"ok": True}

@app.delete("/api/cases/{case_id}/financials/{entry_id}")
async def delete_financial(case_id: str, entry_id: str, user=Depends(require_role(UserRole.ADMIN))):
    case = await db.cases.find_one({"_id": make_id(case_id)})
    if is_locked(case): raise HTTPException(423, "Η υπόθεση είναι κλειστή")
    await db.financials.delete_one({"_id": make_id(entry_id)})
    await audit("DELETE_FINANCIAL", user["id"], "financial", entry_id)
    return {"ok": True}

# ══════════════════════════════════════════════════════════════════════════════
# SEARCH & AUDIT & DASHBOARD
# ══════════════════════════════════════════════════════════════════════════════
@app.get("/api/search")
async def search(q: str = Query(..., min_length=1), user=Depends(get_current_user)):
    results = {"cases": [], "clients": []}
    q_clean = q.strip(); regex = {"$regex": q_clean, "$options": "i"}
    cq = {"$or": [{"title": regex}, {"case_number": regex}, {"description": regex}]}
    if user["role"] == UserRole.LAWYER.value: cq["assigned_lawyer_id"] = user["id"]
    results["cases"] = [serialize(c) for c in await db.cases.find(cq).limit(10).to_list(None)]
    if user["role"] != UserRole.LAWYER.value:
        results["clients"] = [serialize(c) for c in await db.clients.find(
            {"$or": [{"name": regex}, {"email": regex}, {"phone": regex}]}).limit(10).to_list(None)]
    return results

@app.get("/api/audit-logs")
async def get_audit_logs(user=Depends(require_role(UserRole.ADMIN)), skip: int = 0, limit: int = 50, resource: Optional[str] = None):
    query = {"resource": resource} if resource else {}
    logs = await db.audit_logs.find(query).sort("timestamp", -1).skip(skip).limit(min(limit, 100)).to_list(None)
    total = await db.audit_logs.count_documents(query)
    return {"logs": [serialize(l) for l in logs], "total": total}

@app.get("/api/dashboard/stats")
async def dashboard_stats(user=Depends(get_current_user)):
    q = {"assigned_lawyer_id": user["id"]} if user["role"] == UserRole.LAWYER.value else {}
    total = await db.cases.count_documents(q)
    active = await db.cases.count_documents({**q, "status": "active"})
    pending = await db.cases.count_documents({**q, "status": "pending"})
    closed = await db.cases.count_documents({**q, "status": "closed"})
    sd = datetime.utcnow() - timedelta(days=STAGNANT_DAYS)
    stagnant = await db.cases.count_documents({**q, "status": "active", "$or": [
        {"last_activity": {"$lt": sd}}, {"last_activity": {"$exists": False}, "created_at": {"$lt": sd}}]})
    overdue = 0
    if user["role"] != UserRole.SECRETARY.value:
        overdue = await db.financials.count_documents({"payment_status": "overdue"})
    return {"total_cases": total, "active_cases": active, "pending_cases": pending,
            "closed_cases": closed, "stagnant_cases": stagnant, "overdue_financials": overdue}

    return result

# ── Lookup endpoint for populating dropdowns ──────────────────────────────────
@app.get("/api/users/lawyers")
async def list_lawyers(user=Depends(get_current_user)):
    """List active lawyers for case assignment dropdown."""
    lawyers = await db.users.find({"role": UserRole.LAWYER.value, "is_active": True}, {"name": 1, "email": 1}).to_list(None)
    return [serialize(l) for l in lawyers]

async def _check_case_access(case_id, user):
    case = await db.cases.find_one({"_id": make_id(case_id)})
    if not case: raise HTTPException(404, "Η υπόθεση δεν βρέθηκε")

    return case

# ══════════════════════════════════════════════════════════════════════════════
# PHASE 2: DEADLINES / CALENDAR
# ══════════════════════════════════════════════════════════════════════════════
class DeadlineType(str, Enum):
    COURT_DATE = "court_date"          # Δικάσιμος
    FILING_DEADLINE = "filing_deadline" # Προθεσμία κατάθεσης
    STATUTE_LIMIT = "statute_limit"    # Παραγραφή
    MEETING = "meeting"                # Συνάντηση
    PAYMENT_DUE = "payment_due"        # Πληρωμή
    OTHER = "other"                    # Λοιπά

class DeadlineRequest(BaseModel):
    case_id: Optional[str] = None
    title: str
    deadline_type: Optional[str] = None
    type: Optional[str] = None
    date: Optional[datetime] = None
    due_date: Optional[str] = None
    description: Optional[str] = None
    reminder_days: int = 3
    all_day: bool = True
    location: Optional[str] = None

class DeadlineUpdateRequest(BaseModel):
    title: Optional[str] = None
    deadline_type: Optional[DeadlineType] = None
    date: Optional[datetime] = None
    description: Optional[str] = None
    reminder_days: Optional[int] = None
    completed: Optional[bool] = None
    all_day: Optional[bool] = None
    location: Optional[str] = None

@app.get("/api/deadlines")
async def list_deadlines(
    user=Depends(get_current_user),
    from_date: Optional[str] = None,
    to_date: Optional[str] = None,
    include_completed: bool = False
):
    """List deadlines, optionally filtered by date range."""
    query = {}
    if not include_completed:
        query["completed"] = {"$ne": True}
    if user["role"] == UserRole.LAWYER.value:
        # Only deadlines for lawyer's cases + personal (no case_id)
        lawyer_case_ids = await db.cases.distinct("_id", {"assigned_lawyer_id": user["id"]})
        lawyer_case_str_ids = [str(cid) for cid in lawyer_case_ids]
        query["$or"] = [
            {"case_id": {"$in": lawyer_case_str_ids}},
            {"case_id": None},
            {"created_by": user["id"]}
        ]
    if from_date:
        try: query.setdefault("date", {})["$gte"] = datetime.fromisoformat(from_date)
        except: pass
    if to_date:
        try: query.setdefault("date", {})["$lte"] = datetime.fromisoformat(to_date)
        except: pass
    deadlines = await db.deadlines.find(query).sort("date", 1).to_list(None)
    result = []
    for d in deadlines:
        s = serialize(d)
        if s.get("case_id"):
            case = await db.cases.find_one({"_id": make_id(s["case_id"])}, {"title": 1, "case_number": 1})
            if case:
                s["case_title"] = case.get("title", "")
                s["case_number"] = case.get("case_number", "")
        result.append(s)
    return result

@app.get("/api/deadlines/upcoming")
async def upcoming_deadlines(user=Depends(get_current_user), days: int = 14):
    """Get deadlines within next N days for dashboard."""
    now = datetime.utcnow()
    end = now + timedelta(days=days)
    query = {"date": {"$gte": now, "$lte": end}, "completed": {"$ne": True}}
    if user["role"] == UserRole.LAWYER.value:
        cids = await db.cases.distinct("_id", {"assigned_lawyer_id": user["id"]})
        cids_str = [str(c) for c in cids]
        query["$or"] = [{"case_id": {"$in": cids_str}}, {"case_id": None}, {"created_by": user["id"]}]
    deadlines = await db.deadlines.find(query).sort("date", 1).limit(20).to_list(None)
    result = []
    for d in deadlines:
        s = serialize(d)
        if s.get("case_id"):
            case = await db.cases.find_one({"_id": make_id(s["case_id"])}, {"title": 1, "case_number": 1})
            if case: s["case_title"] = case.get("title", ""); s["case_number"] = case.get("case_number", "")
        # Add urgency
        days_until = (d["date"] - now).days
        s["days_until"] = days_until
        s["is_urgent"] = days_until <= d.get("reminder_days", 3)
        s["is_overdue"] = days_until < 0
        result.append(s)
    return result

@app.post("/api/deadlines", status_code=201)
async def create_deadline(req: DeadlineRequest, user=Depends(get_current_user)):
    if req.case_id:
        await _check_case_access(req.case_id, user)
        await _check_payment_gate(req.case_id, f"Προθεσμία: {req.title}")
    doc = {
        "case_id": req.case_id,
        "title": sanitize_string(req.title),
        "deadline_type": (req.deadline_type or req.type or "hearing"),
        "date": req.due_date and __import__('datetime').datetime.fromisoformat(req.due_date) or req.date or __import__('datetime').datetime.utcnow(),
        "description": sanitize_string(req.description) if req.description else None,
        "reminder_days": req.reminder_days,
        "all_day": req.all_day,
        "location": sanitize_string(req.location) if req.location else None,
        "completed": False,
        "created_by": user["id"],
        "created_at": datetime.utcnow()
    }
    r = await db.deadlines.insert_one(doc)
    await audit("CREATE_DEADLINE", user["id"], "deadline", str(r.inserted_id))
    doc["_id"] = r.inserted_id
    return serialize(doc)

@app.put("/api/deadlines/{deadline_id}")
async def update_deadline(deadline_id: str, req: DeadlineUpdateRequest, user=Depends(get_current_user)):
    dl = await db.deadlines.find_one({"_id": make_id(deadline_id)})
    if not dl: raise HTTPException(404, "Η προθεσμία δεν βρέθηκε")
    if dl.get("case_id"):
        await _check_case_access(dl["case_id"], user)
    update = {}
    if req.title is not None: update["title"] = sanitize_string(req.title)
    if req.deadline_type is not None: update["deadline_type"] = req.deadline_type.value
    if req.date is not None: update["date"] = req.date
    if req.description is not None: update["description"] = sanitize_string(req.description)
    if req.reminder_days is not None: update["reminder_days"] = req.reminder_days
    if req.completed is not None: update["completed"] = req.completed
    if req.all_day is not None: update["all_day"] = req.all_day
    if req.location is not None: update["location"] = sanitize_string(req.location)
    if update:
        update["updated_at"] = datetime.utcnow()
        await db.deadlines.update_one({"_id": make_id(deadline_id)}, {"$set": update})
    await audit("UPDATE_DEADLINE", user["id"], "deadline", deadline_id)
    return {"ok": True}

@app.delete("/api/deadlines/{deadline_id}")
async def delete_deadline(deadline_id: str, user=Depends(get_current_user)):
    dl = await db.deadlines.find_one({"_id": make_id(deadline_id)})
    if not dl: raise HTTPException(404, "Η προθεσμία δεν βρέθηκε")
    if dl.get("case_id"):
        await _check_case_access(dl["case_id"], user)
    await db.deadlines.delete_one({"_id": make_id(deadline_id)})
    await audit("DELETE_DEADLINE", user["id"], "deadline", deadline_id)
    return {"ok": True}

@app.get("/api/cases/{case_id}/deadlines")
async def case_deadlines(case_id: str, user=Depends(get_current_user)):
    await _check_case_access(case_id, user)
    dls = await db.deadlines.find({"case_id": case_id}).sort("date", 1).to_list(None)
    now = datetime.utcnow()
    result = []
    for d in dls:
        s = serialize(d)
        days_until = (d["date"] - now).days
        s["days_until"] = days_until
        s["is_urgent"] = days_until <= d.get("reminder_days", 3) and not d.get("completed")
        s["is_overdue"] = days_until < 0 and not d.get("completed")
        result.append(s)
    return result

# ══════════════════════════════════════════════════════════════════════════════
# PHASE 3: MULTI-PARTY CASES (Case Parties)
# ══════════════════════════════════════════════════════════════════════════════
class PartyRole(str, Enum):
    CLIENT = "client"           # Εντολέας
    OPPONENT = "opponent"       # Αντίδικος
    WITNESS = "witness"         # Μάρτυρας
    THIRD_PARTY = "third_party" # Τρίτος
    GUARANTOR = "guarantor"     # Εγγυητής
    EXPERT = "expert"           # Πραγματογνώμονας
    OTHER = "other"             # Λοιποί

class CasePartyRequest(BaseModel):
    name: str
    party_role: PartyRole
    email: Optional[str] = None
    phone: Optional[str] = None
    address: Optional[str] = None
    tax_id: Optional[str] = None
    lawyer_name: Optional[str] = None  # Πληρεξούσιος δικηγόρος αντιδίκου
    notes: Optional[str] = None

@app.get("/api/cases/{case_id}/parties")
async def list_case_parties(case_id: str, user=Depends(get_current_user)):
    await _check_case_access(case_id, user)
    parties = await db.case_parties.find({"case_id": case_id}).sort("party_role", 1).to_list(None)
    return [serialize(p) for p in parties]

@app.post("/api/cases/{case_id}/parties", status_code=201)
async def add_case_party(case_id: str, req: CasePartyRequest, user=Depends(get_current_user)):
    await _check_case_access(case_id, user)
    if req.phone and not validate_phone(req.phone): raise HTTPException(400, "Μη έγκυρο τηλέφωνο")
    if req.tax_id and not validate_tax_id(req.tax_id): raise HTTPException(400, "Μη έγκυρο ΑΦΜ")
    doc = {
        "case_id": case_id,
        "name": sanitize_string(req.name),
        "party_role": req.party_role.value,
        "email": req.email.strip().lower() if req.email else None,
        "phone": req.phone.strip() if req.phone else None,
        "address": sanitize_string(req.address) if req.address else None,
        "tax_id": req.tax_id.strip() if req.tax_id else None,
        "lawyer_name": sanitize_string(req.lawyer_name) if req.lawyer_name else None,
        "notes": sanitize_string(req.notes) if req.notes else None,
        "created_by": user["id"],
        "created_at": datetime.utcnow()
    }
    r = await db.case_parties.insert_one(doc)
    await audit("ADD_CASE_PARTY", user["id"], "case_party", str(r.inserted_id), {"case_id": case_id})
    doc["_id"] = r.inserted_id
    return serialize(doc)

@app.put("/api/cases/{case_id}/parties/{party_id}")
async def update_case_party(case_id: str, party_id: str, req: CasePartyRequest, user=Depends(get_current_user)):
    await _check_case_access(case_id, user)
    if req.phone and not validate_phone(req.phone): raise HTTPException(400, "Μη έγκυρο τηλέφωνο")
    if req.tax_id and not validate_tax_id(req.tax_id): raise HTTPException(400, "Μη έγκυρο ΑΦΜ")
    data = {
        "name": sanitize_string(req.name), "party_role": req.party_role.value,
        "email": req.email.strip().lower() if req.email else None,
        "phone": req.phone.strip() if req.phone else None,
        "address": sanitize_string(req.address) if req.address else None,
        "tax_id": req.tax_id.strip() if req.tax_id else None,
        "lawyer_name": sanitize_string(req.lawyer_name) if req.lawyer_name else None,
        "notes": sanitize_string(req.notes) if req.notes else None,
        "updated_at": datetime.utcnow()
    }
    await db.case_parties.update_one({"_id": make_id(party_id), "case_id": case_id}, {"$set": data})
    await audit("UPDATE_CASE_PARTY", user["id"], "case_party", party_id)
    return {"ok": True}

@app.delete("/api/cases/{case_id}/parties/{party_id}")
async def delete_case_party(case_id: str, party_id: str, user=Depends(get_current_user)):
    await _check_case_access(case_id, user)
    await db.case_parties.delete_one({"_id": make_id(party_id), "case_id": case_id})
    await audit("DELETE_CASE_PARTY", user["id"], "case_party", party_id)
    return {"ok": True}

# ══════════════════════════════════════════════════════════════════════════════
# PHASE 4: GREEK LAWYER INVOICING
# ══════════════════════════════════════════════════════════════════════════════
# ΦΠΑ 24%, Παρακράτηση Φόρου 15%, Γραμμάτιο Παράστασης
VAT_RATE = 0.24
WITHHOLDING_TAX_RATE = 0.15

# ── Κρατήσεις Γραμματίου Προείσπραξης (Δικηγορικός Σύλλογος Αθηνών) ──────────
# Πηγή: ΕΦΚΑ 2024, ΔΣΑ, ΕΑΝ
GRAMMATIO_EFKA_RATE     = 0.2695   # ΕΦΚΑ (κύρια σύνταξη + επικουρική + εφάπαξ)
GRAMMATIO_EAN_RATE      = 0.0167   # ΕΑΝ — Ειδικό Αποθεματικό Νομικής Αρωγής
GRAMMATIO_DSA_RATE      = 0.0200   # Ταμείο Δ.Σ. (ΔΣΑ/τοπικός σύλλογος)
# Σύνολο κρατήσεων ≈ 30.62% — ο δικηγόρος εισπράττει ~69.38% του γραμματίου


# ── Invoicing Context Endpoint ────────────────────────────────────────────────
@app.get("/api/cases/{case_id}/invoicing-context")
async def get_invoicing_context(case_id: str, user=Depends(get_current_user)):
    """Επιστρέφει όλα τα χρήσιμα στοιχεία για τιμολόγηση μιας υπόθεσης:
    - Στοιχεία υπόθεσης + πελάτη
    - Έξοδα υπόθεσης (ανά κατηγορία)
    - Υπάρχοντα τιμολόγια + σύνολα
    - Γραμμάτια (από expenses_log)
    """
    await _check_case_access(case_id, user)
    case = await db.cases.find_one({"_id": make_id(case_id)})
    if not case:
        raise HTTPException(404, "Η υπόθεση δεν βρέθηκε")

    client = None
    if case.get("client_id"):
        client = await db.clients.find_one({"_id": make_id(case["client_id"])})

    # Expenses for this case
    expenses_raw = await db.expenses_log.find({"case_id": case_id}).sort("date", -1).to_list(None)
    expenses = [serialize(e) for e in expenses_raw]

    # Grammatia from expenses (category = "grammatio")
    grammatia = [e for e in expenses if e.get("category") == "grammatio"]
    grammatio_total = sum(float(g.get("amount", 0)) for g in grammatia)

    # Expense totals by category
    exp_by_cat: dict = {}
    for e in expenses:
        cat = e.get("category", "other")
        exp_by_cat[cat] = exp_by_cat.get(cat, 0) + float(e.get("amount", 0))
    total_expenses = sum(float(e.get("amount", 0)) for e in expenses)

    # Existing invoices for this case
    invoices_raw = await db.invoices.find({"case_id": case_id}).sort("issue_date", -1).to_list(None)
    invoices = [serialize(i) for i in invoices_raw]
    already_invoiced = sum(float(i.get("total_payable", i.get("total", 0))) for i in invoices)
    already_paid     = sum(float(i.get("amount_paid", 0)) for i in invoices)

    return {
        "case": serialize(case),
        "client": serialize(client) if client else None,
        "expenses": expenses,
        "expenses_by_category": exp_by_cat,
        "total_expenses": round(total_expenses, 2),
        "grammatia": grammatia,
        "grammatio_total": round(grammatio_total, 2),
        "invoices": invoices,
        "already_invoiced": round(already_invoiced, 2),
        "already_paid": round(already_paid, 2),
        # Standard grammatio deduction rates (για πληροφορία)
        "grammatio_rates": {
            "efka": GRAMMATIO_EFKA_RATE,
            "ean": GRAMMATIO_EAN_RATE,
            "dsa": GRAMMATIO_DSA_RATE,
            "total_deductions": round(GRAMMATIO_EFKA_RATE + GRAMMATIO_EAN_RATE + GRAMMATIO_DSA_RATE, 4),
            "net_to_lawyer_pct": round(1 - GRAMMATIO_EFKA_RATE - GRAMMATIO_EAN_RATE - GRAMMATIO_DSA_RATE, 4),
        },
    }


class InvoiceRequest(BaseModel):
    case_id: str
    client_id: str
    items: List[dict]  # [{description, amount, is_grammatio, is_expense}]
    notes: Optional[str] = None
    issue_date: Optional[datetime] = None
    # Γραμμάτιο — αν συμπεριληφθεί ως ήδη καταβληθέν
    grammatio_gross: float = 0.0          # Αξία γραμματίου
    grammatio_efka: float = 0.0           # Κράτηση ΕΦΚΑ
    grammatio_ean: float = 0.0            # Κράτηση ΕΑΝ
    grammatio_dsa: float = 0.0            # Κράτηση Δ.Σ.
    grammatio_other_deductions: float = 0.0  # Λοιπές κρατήσεις

class FeeCalculatorRequest(BaseModel):
    net_amount: float
    include_vat: bool = True
    include_withholding: bool = True
    grammatio_amount: float = 0.0  # Γραμμάτιο Παράστασης

@app.post("/api/invoicing/calculate")
async def calculate_fees(req: FeeCalculatorRequest, user=Depends(get_current_user)):
    """Calculate Greek lawyer fee breakdown."""
    net = req.net_amount
    if net <= 0: raise HTTPException(400, "Το ποσό πρέπει να είναι θετικό")

    vat_amount = round(net * VAT_RATE, 2) if req.include_vat else 0
    withholding_amount = round(net * WITHHOLDING_TAX_RATE, 2) if req.include_withholding else 0
    grammatio = round(req.grammatio_amount, 2)
    gross_amount = round(net + vat_amount, 2)
    total_payable = round(gross_amount - withholding_amount + grammatio, 2)
    lawyer_receives = round(gross_amount - withholding_amount, 2)

    return {
        "net_amount": net,
        "vat_rate": VAT_RATE,
        "vat_amount": vat_amount,
        "gross_amount": gross_amount,
        "withholding_tax_rate": WITHHOLDING_TAX_RATE,
        "withholding_amount": withholding_amount,
        "grammatio_amount": grammatio,
        "total_payable_by_client": total_payable,
        "lawyer_receives": lawyer_receives,
        "breakdown": {
            "Καθαρή Αμοιβή": net,
            "ΦΠΑ 24%": vat_amount,
            "Μικτή Αμοιβή": gross_amount,
            "Παρακράτηση Φόρου 15%": -withholding_amount,
            "Γραμμάτιο Παράστασης": grammatio,
            "Ο Εντολέας Πληρώνει": total_payable,
            "Ο Δικηγόρος Εισπράττει": lawyer_receives
        }
    }

@app.post("/api/invoices", status_code=201)
async def create_invoice(req: InvoiceRequest, user=Depends(require_role(UserRole.ADMIN, UserRole.LAWYER))):
    await _check_case_access(req.case_id, user)
    cl = await db.clients.find_one({"_id": make_id(req.client_id)})
    if not cl: raise HTTPException(404, "Ο εντολέας δεν βρέθηκε")

    # ── Process line items ────────────────────────────────────────────────────
    total_net = 0.0        # αμοιβές (χωρίς ΦΠΑ)
    total_expenses = 0.0   # έξοδα (δεν φέρουν ΦΠΑ/παρακράτηση)
    processed_items = []

    for item in req.items:
        amt = float(item.get("amount", 0))
        if amt <= 0: raise HTTPException(400, "Μη έγκυρο ποσό γραμμής")
        is_expense  = item.get("is_expense", False)
        is_grammatio = item.get("is_grammatio", False)
        if is_expense or is_grammatio:
            total_expenses += amt
        else:
            total_net += amt
        processed_items.append({
            "description":   item.get("description", ""),
            "amount":        amt,
            "is_grammatio":  is_grammatio,
            "is_expense":    is_expense,
        })

    # ── Tax calculations on net fees only ─────────────────────────────────────
    vat        = round(total_net * VAT_RATE, 2)
    withholding = round(total_net * WITHHOLDING_TAX_RATE, 2)
    gross       = round(total_net + vat, 2)

    # ── Γραμμάτιο Προείσπραξης ─────────────────────────────────────────────────
    grammatio_gross = round(float(req.grammatio_gross), 2)
    if grammatio_gross > 0:
        # Κρατήσεις από γραμμάτιο
        efka   = round(float(req.grammatio_efka)  or grammatio_gross * GRAMMATIO_EFKA_RATE, 2)
        ean    = round(float(req.grammatio_ean)   or grammatio_gross * GRAMMATIO_EAN_RATE,  2)
        dsa    = round(float(req.grammatio_dsa)   or grammatio_gross * GRAMMATIO_DSA_RATE,  2)
        other  = round(float(req.grammatio_other_deductions), 2)
        total_deductions = efka + ean + dsa + other
        grammatio_net_to_lawyer = round(grammatio_gross - total_deductions, 2)
    else:
        efka = ean = dsa = other = total_deductions = grammatio_net_to_lawyer = 0.0

    grammatio_info = {
        "gross": grammatio_gross,
        "efka": efka, "ean": ean, "dsa": dsa,
        "other_deductions": other,
        "total_deductions": round(total_deductions, 2),
        "net_to_lawyer": grammatio_net_to_lawyer,
    } if grammatio_gross > 0 else None

    # ── Final totals ──────────────────────────────────────────────────────────
    # Ο εντολέας χρεώνεται: αμοιβή + ΦΠΑ + έξοδα - παρακράτηση
    # Το γραμμάτιο έχει ήδη καταβληθεί → αφαιρείται από το υπόλοιπο
    total_before_grammatio = round(gross - withholding + total_expenses, 2)
    total_payable           = round(total_before_grammatio - grammatio_gross, 2)
    lawyer_receives         = round(gross - withholding + total_expenses + grammatio_net_to_lawyer, 2)

    # ── Invoice number ─────────────────────────────────────────────────────────
    year = datetime.utcnow().year
    inv_counter = await db.counters.find_one_and_update(
        {"_id": "invoice_number"}, {"$set": {"year": year}, "$inc": {"seq": 1}},
        upsert=True, return_document=True)
    if inv_counter.get("year") != year:
        inv_counter = await db.counters.find_one_and_update(
            {"_id": "invoice_number"}, {"$set": {"year": year, "seq": 1}}, return_document=True)
    inv_number = f"ΤΔΑ-{year}-{str(inv_counter['seq']).zfill(4)}"

    invoice = {
        "invoice_number":       inv_number,
        "case_id":              req.case_id,
        "client_id":            req.client_id,
        "client_name":          cl.get("full_name", cl.get("name", "")),
        "client_afm":           cl.get("afm", cl.get("tax_id", "")),
        "client_address":       cl.get("address", ""),
        "client_email":         cl.get("email", ""),
        "items":                processed_items,
        # Αμοιβές
        "net_amount":           round(total_net, 2),
        "vat_rate":             VAT_RATE,
        "vat_amount":           vat,
        "gross_amount":         gross,
        "withholding_rate":     WITHHOLDING_TAX_RATE,
        "withholding_amount":   withholding,
        # Έξοδα
        "expenses_amount":      round(total_expenses, 2),
        # Γραμμάτιο
        "grammatio":            grammatio_info,
        "grammatio_gross":      grammatio_gross,
        # Σύνολα
        "total_before_grammatio": total_before_grammatio,
        "total_payable":        max(total_payable, 0),  # δεν πηγαίνει αρνητικό
        "lawyer_receives":      round(lawyer_receives, 2),
        "notes":                sanitize_string(req.notes) if req.notes else None,
        "issue_date":           req.issue_date or datetime.utcnow(),
        "created_by":           user["id"],
        "created_by_name":      user.get("name", ""),
        "created_at":           datetime.utcnow(),
        "status":               "issued",
        "payment_status":       "pending",
        "amount_paid":          0.0,
    }
    r = await db.invoices.insert_one(invoice)
    await audit("CREATE_INVOICE", user["id"], "invoice", str(r.inserted_id), {"invoice_number": inv_number})
    invoice["_id"] = r.inserted_id
    return serialize(invoice)

@app.get("/api/invoices")
async def list_invoices(user=Depends(get_current_user), case_id: Optional[str] = None):
    if user["role"] == UserRole.SECRETARY.value:
        raise HTTPException(403, "Δεν έχετε πρόσβαση στα τιμολόγια")
    query = {}
    if case_id: query["case_id"] = case_id
    if user["role"] == UserRole.LAWYER.value:
        cids = await db.cases.distinct("_id", {"assigned_lawyer_id": user["id"]})
        query["case_id"] = {"$in": [str(c) for c in cids]}
    invoices = await db.invoices.find(query).sort("created_at", -1).to_list(None)
    return [serialize(i) for i in invoices]

@app.get("/api/invoices/{invoice_id}")
async def get_invoice(invoice_id: str, user=Depends(get_current_user)):
    if user["role"] == UserRole.SECRETARY.value:
        raise HTTPException(403, "Δεν έχετε πρόσβαση")
    inv = await db.invoices.find_one({"_id": make_id(invoice_id)})
    if not inv: raise HTTPException(404, "Το τιμολόγιο δεν βρέθηκε")
    return serialize(inv)

@app.get("/api/cases/{case_id}/invoices")
async def case_invoices(case_id: str, user=Depends(get_current_user)):
    if user["role"] == UserRole.SECRETARY.value:
        raise HTTPException(403, "Δεν έχετε πρόσβαση")
    await _check_case_access(case_id, user)
    invs = await db.invoices.find({"case_id": case_id}).sort("created_at", -1).to_list(None)
    return [serialize(i) for i in invs]

# ══════════════════════════════════════════════════════════════════════════════
# PHASE 5: DOCUMENT TEMPLATES
# ══════════════════════════════════════════════════════════════════════════════
# Built-in Greek legal document templates with auto-fill
# ─────────────────────────────────────────────────────────────────────────────
# LEGAL DOCUMENT TEMPLATES  (νομικά πλήρη πρότυπα)
# Each field entry: {"name": str, "label": str, "type": "text"|"date"|"textarea", "required": bool}
# ─────────────────────────────────────────────────────────────────────────────
TEMPLATES = {
    # ── 1. ΕΞΟΥΣΙΟΔΟΤΗΣΗ ──────────────────────────────────────────────────────
    "exousiodotisi": {
        "id": "exousiodotisi",
        "name": "Εξουσιοδότηση",
        "description": "Γενική εξουσιοδότηση προς δικηγόρο για δικαστική εκπροσώπηση",
        "category": "Γενικό",
        "fields": [
            {"name":"_case_id","label":"Υπόθεση (προαιρετική — συμπληρώνει αυτόματα)","type":"linked","linked_to":"cases","required":False},
            {"name":"_client_id","label":"Εντολέας (εναλλακτικά, αν δεν υπάρχει υπόθεση)","type":"linked","linked_to":"clients","required":False},
            {"name":"client_name","label":"Ονοματεπώνυμο εντολέα","type":"text","required":True},
            {"name":"client_father","label":"Πατρώνυμο εντολέα","type":"text","required":True},
            {"name":"client_id_number","label":"Αρ. Δελτίου Ταυτότητας / Διαβατηρίου","type":"text","required":True},
            {"name":"client_tax_id","label":"ΑΦΜ εντολέα","type":"text","required":True},
            {"name":"client_address","label":"Πλήρης διεύθυνση κατοικίας","type":"text","required":True},
            {"name":"client_phone","label":"Τηλέφωνο εντολέα","type":"text","required":False},
            {"name":"lawyer_name","label":"Ονοματεπώνυμο δικηγόρου","type":"text","required":True},
            {"name":"lawyer_am","label":"Αριθμός Μητρώου (ΑΜ) δικηγόρου","type":"text","required":True},
            {"name":"lawyer_dsb","label":"Δικηγορικός Σύλλογος δικηγόρου","type":"text","required":True},
            {"name":"case_title","label":"Τίτλος / αντικείμενο υπόθεσης","type":"text","required":True},
            {"name":"court","label":"Δικαστήριο / Αρχή","type":"text","required":True},
            {"name":"city","label":"Πόλη σύνταξης","type":"text","required":True},
            {"name":"date","label":"Ημερομηνία","type":"date","required":True},
        ],
        "template": """ΕΞΟΥΣΙΟΔΟΤΗΣΗ

Εγώ, ο/η κάτωθι υπογράφων/ουσα {{client_name}}, του/της {{client_father}}, κάτοικος {{client_address}}, κάτοχος Δ.Α.Τ. / Διαβατηρίου αριθ. {{client_id_number}}, με Α.Φ.Μ. {{client_tax_id}}, τηλ. {{client_phone}},

ΕΞΟΥΣΙΟΔΟΤΩ

τον/την Δικηγόρο {{lawyer_name}}, μέλος του Δικηγορικού Συλλόγου {{lawyer_dsb}} (Α.Μ. {{lawyer_am}}), να με εκπροσωπεί νόμιμα σε κάθε ενέργεια σχετική με την υπόθεση «{{case_title}}» ενώπιον {{court}} και κάθε άλλης αρμόδιας αρχής, δικαστηρίου ή υπηρεσίας.

Ειδικότερα εξουσιοδοτώ τον/την ανωτέρω να:
α) Καταθέτει, υπογράφει και υποβάλλει κάθε δικόγραφο, αίτηση, δήλωση ή υπόμνημα.
β) Εκπροσωπεί τα συμφέροντά μου ενώπιον δικαστηρίων, εισαγγελιών, αστυνομικών, διοικητικών και λοιπών αρχών.
γ) Λαμβάνει γνώση κάθε εγγράφου και να υπογράφει πρακτικά συζητήσεων.
δ) Ασκεί κάθε νόμιμο ένδικο μέσο ή βοήθημα.
ε) Διενεργεί κάθε απαραίτητη νόμιμη ενέργεια για την προάσπιση των δικαιωμάτων και συμφερόντων μου.

Η παρούσα ισχύει μέχρι ρητής ανάκλησής της.

{{city}}, {{date}}

Ο/Η Εξουσιοδοτών/ούσα


_________________________________
{{client_name}}"""
    },

    # ── 2. ΕΙΔΙΚΟ ΠΛΗΡΕΞΟΥΣΙΟ ─────────────────────────────────────────────────
    "plirexousio": {
        "id": "plirexousio",
        "name": "Ειδικό Πληρεξούσιο",
        "description": "Ειδικό πληρεξούσιο για παράσταση σε δίκη (άρθρο 96 ΚΠολΔ)",
        "category": "Δικαστικό",
        "fields": [
            {"name":"_case_id","label":"Υπόθεση (συμπληρώνει αυτόματα)","type":"linked","linked_to":"cases","required":False},
            {"name":"client_name","label":"Ονοματεπώνυμο εντολέα","type":"text","required":True},
            {"name":"client_father","label":"Πατρώνυμο εντολέα","type":"text","required":True},
            {"name":"client_id_number","label":"Αρ. Δ.Α.Τ. / Διαβατηρίου","type":"text","required":True},
            {"name":"client_tax_id","label":"ΑΦΜ εντολέα","type":"text","required":True},
            {"name":"client_address","label":"Διεύθυνση κατοικίας","type":"text","required":True},
            {"name":"client_phone","label":"Τηλέφωνο","type":"text","required":False},
            {"name":"lawyer_name","label":"Ονοματεπώνυμο δικηγόρου","type":"text","required":True},
            {"name":"lawyer_am","label":"Αριθμός Μητρώου (ΑΜ)","type":"text","required":True},
            {"name":"lawyer_dsb","label":"Δικηγορικός Σύλλογος","type":"text","required":True},
            {"name":"lawyer_address","label":"Διεύθυνση δικηγορικού γραφείου","type":"text","required":False},
            {"name":"court","label":"Δικαστήριο","type":"text","required":True},
            {"name":"hearing_date","label":"Ημερομηνία δικασίμου","type":"date","required":True},
            {"name":"case_title","label":"Αντικείμενο / τίτλος υπόθεσης","type":"text","required":True},
            {"name":"case_number","label":"Αριθμός Υπόθεσης (ΓΑΚ/ΕΑΚ)","type":"text","required":False},
            {"name":"opponent_name","label":"Ονοματεπώνυμο αντιδίκου","type":"text","required":False},
            {"name":"city","label":"Πόλη σύνταξης","type":"text","required":True},
            {"name":"date","label":"Ημερομηνία σύνταξης","type":"date","required":True},
        ],
        "template": """ΕΙΔΙΚΟ ΠΛΗΡΕΞΟΥΣΙΟ
(κατ' άρθρο 96 Κ.Πολ.Δ.)

Εγώ, ο/η {{client_name}}, του/της {{client_father}}, κάτοικος {{client_address}}, κάτοχος Δ.Α.Τ. / Διαβατηρίου αριθ. {{client_id_number}}, Α.Φ.Μ. {{client_tax_id}}, τηλ. {{client_phone}},

ΔΙΟΡΙΖΩ

ειδικό/ή πληρεξούσιο/α Δικηγόρο μου τον/την {{lawyer_name}}, Δικηγόρο {{lawyer_dsb}}, Α.Μ. {{lawyer_am}}, (οδός {{lawyer_address}}), και του/της παρέχω πλήρη εξουσία και πληρεξουσιότητα να:

1. Παρασταθεί ενώπιον {{court}} κατά τη δικάσιμο της {{hearing_date}} για τη συζήτηση της υπόθεσης «{{case_title}}» (Αρ. {{case_number}}) μεταξύ εμού και του/της {{opponent_name}}.
2. Αναβάλει ή μετακινεί τη συζήτηση, συμφωνεί με τον αντίδικο για τόπο και χρόνο εξέτασης μαρτύρων.
3. Παρίσταται σε κάθε αναβολή ή μετά από αναβολή.
4. Αποδέχεται ή αρνείται έγγραφα του αντιδίκου, ασκεί παρεμπίπτουσες αγωγές ή αντίκρουσή τους.
5. Ασκεί κάθε ένδικο μέσο (έφεση, αναίρεση, αναψηλάφηση, αντιτιθέμενη έφεση, τριτανακοπή).
6. Εκδίδει και παραλαμβάνει αντίγραφα αποφάσεων και δικογράφων.
7. Προβαίνει σε κάθε απαραίτητη νόμιμη ενέργεια για την ευόδωση της παραπάνω υπόθεσης.

Ό,τι πράξει ο/η ανωτέρω Δικηγόρος εντός των ορίων της παρούσας πληρεξουσιότητας θεωρώ ισχυρό και δεσμευτικό για εμένα.

{{city}}, {{date}}

Ο/Η Εντολέας


_________________________________
{{client_name}}
(Υπογραφή)"""
    },

    # ── 3. ΜΗΝΥΣΗ / ΕΓΚΛΗΣΗ ───────────────────────────────────────────────────
    "minusi": {
        "id": "minusi",
        "name": "Μήνυση – Έγκληση",
        "description": "Μήνυση/Έγκληση ενώπιον Εισαγγελέα (άρθρα 40-43 ΚΠΔ)",
        "category": "Ποινικό",
        "fields": [
            {"name":"_case_id","label":"Υπόθεση (συμπληρώνει αυτόματα)","type":"linked","linked_to":"cases","required":False},
            {"name":"_client_id","label":"Μηνυτής (αν δεν υπάρχει υπόθεση)","type":"linked","linked_to":"clients","required":False},
            {"name":"client_name","label":"Ονοματεπώνυμο μηνυτή","type":"text","required":True},
            {"name":"client_father","label":"Πατρώνυμο μηνυτή","type":"text","required":True},
            {"name":"client_id_number","label":"Αρ. Δ.Α.Τ. μηνυτή","type":"text","required":True},
            {"name":"client_tax_id","label":"ΑΦΜ μηνυτή","type":"text","required":True},
            {"name":"client_address","label":"Διεύθυνση μηνυτή","type":"text","required":True},
            {"name":"client_phone","label":"Τηλέφωνο μηνυτή","type":"text","required":True},
            {"name":"client_email","label":"Email μηνυτή","type":"text","required":False},
            {"name":"opponent_name","label":"Ονοματεπώνυμο μηνυομένου","type":"text","required":True},
            {"name":"opponent_father","label":"Πατρώνυμο μηνυομένου (αν γνωστό)","type":"text","required":False},
            {"name":"opponent_address","label":"Διεύθυνση μηνυομένου","type":"text","required":True},
            {"name":"offense","label":"Αδίκημα (π.χ. Απάτη κατ' εξακολούθηση)","type":"text","required":True},
            {"name":"law_articles","label":"Σχετικά άρθρα νόμου (π.χ. άρθρα 386 ΠΚ)","type":"text","required":True},
            {"name":"incident_date","label":"Ημερομηνία/περίοδος συμβάντος","type":"text","required":True},
            {"name":"incident_place","label":"Τόπος συμβάντος","type":"text","required":True},
            {"name":"description","label":"Ιστορικό – περιγραφή αδικήματος","type":"textarea","required":True},
            {"name":"damages","label":"Ζημία / βλάβη που υπέστη ο μηνυτής","type":"textarea","required":False},
            {"name":"evidence","label":"Αποδεικτικά μέσα (έγγραφα, μάρτυρες κ.λπ.)","type":"textarea","required":False},
            {"name":"court","label":"Εισαγγελία (π.χ. Πλημμελειοδικών Αθηνών)","type":"text","required":True},
            {"name":"lawyer_name","label":"Ονοματεπώνυμο δικηγόρου","type":"text","required":False},
            {"name":"lawyer_am","label":"ΑΜ δικηγόρου","type":"text","required":False},
            {"name":"city","label":"Πόλη κατάθεσης","type":"text","required":True},
            {"name":"date","label":"Ημερομηνία","type":"date","required":True},
        ],
        "template": """ΕΝΩΠΙΟΝ ΤΟΥ κ. ΕΙΣΑΓΓΕΛΕΑ {{court}}

ΜΗΝΥΣΗ – ΕΓΚΛΗΣΗ

Του/Της: {{client_name}}, του/της {{client_father}}, κατοίκου {{client_address}},
Δ.Α.Τ. αριθ. {{client_id_number}}, Α.Φ.Μ. {{client_tax_id}},
Τηλ.: {{client_phone}}, Email: {{client_email}}
(εφεξής «Μηνυτής/τρια»)

– δια του πληρεξουσίου Δικηγόρου {{lawyer_name}}, Α.Μ. {{lawyer_am}} –

ΚΑΤΑ

Του/Της: {{opponent_name}}, του/της {{opponent_father}}, κατοίκου {{opponent_address}}
(εφεξής «Μηνυόμενος/η»)

ΓΙΑ ΤΑ ΑΔΙΚΗΜΑΤΑ: {{offense}}
({{law_articles}})

Α. ΙΣΤΟΡΙΚΟ

Κύριε Εισαγγελέα,

Τον/τη μηνυόμενο/η γνωρίζω / συνδέομαι / ήλθα σε επαφή μαζί του/της κατά τον/τους χρόνο/χρόνους {{incident_date}} στον/στη {{incident_place}}.

{{description}}

Β. ΖΗΜΙΑ / ΒΛΑΒΗ

{{damages}}

Γ. ΑΠΟΔΕΙΚΤΙΚΑ ΜΕΣΑ

{{evidence}}

Δ. ΝΟΜΙΚΗ ΒΑΣΗ

Επειδή τα ανωτέρω συμπεριφερόμενα του/της μηνυομένου/ης συνιστούν τα αδικήματα: {{offense}}, κατά τις διατάξεις {{law_articles}}.

Επειδή τα ανωτέρω αποτελούν αξιόποινες πράξεις.

Επειδή υφίσταται ανάγκη άμεσης ποινικής δίωξης για την προστασία των εννόμων συμφερόντων μου.

Ε. ΑΙΤΗΜΑ

Για τους λόγους αυτούς, και με ρητή επιφύλαξη κάθε άλλου νόμιμου δικαιώματός μου (αστική αγωγή, αξίωση αποζημίωσης κ.λπ.),

ΖΗΤΩ

1. Να ασκηθεί ποινική δίωξη κατά του/της ανωτέρω μηνυομένου/ης για τα αναφερόμενα αδικήματα.
2. Να διαταχθεί η διεξαγωγή προκαταρκτικής εξέτασης και κύριας ανάκρισης.
3. Να ληφθούν υπόψη τα προσκομιζόμενα αποδεικτικά μέσα.
4. Να κοινοποιηθεί αντίγραφο της παρούσας στον/στην εγκαλούμενο/η κατ' άρθρο 42 §2 ΚΠΔ.
5. Να κληθεί ο/η μηνυτής/τρια ως παθών/ούσα στην κύρια ανάκριση.

Επιφυλάσσομαι πλήρως των αστικών μου δικαιωμάτων.

{{city}}, {{date}}

Ο/Η Μηνυτής/τρια                             Ο Πληρεξούσιος Δικηγόρος


_____________________                        _____________________
{{client_name}}                               {{lawyer_name}}"""
    },

    # ── 4. ΑΙΤΗΣΗ ΑΝΑΣΤΟΛΗΣ ΕΚΤΕΛΕΣΗΣ ────────────────────────────────────────
    "aitisi_anastolis": {
        "id": "aitisi_anastolis",
        "name": "Αίτηση Αναστολής Εκτέλεσης",
        "description": "Αίτηση αναστολής εκτέλεσης απόφασης (άρθρα 912-915 ΚΠολΔ)",
        "category": "Δικαστικό",
        "fields": [
            {"name":"_case_id","label":"Υπόθεση (συμπληρώνει αυτόματα)","type":"linked","linked_to":"cases","required":False},
            {"name":"client_name","label":"Ονοματεπώνυμο αιτούντος","type":"text","required":True},
            {"name":"client_father","label":"Πατρώνυμο αιτούντος","type":"text","required":True},
            {"name":"client_id_number","label":"Αρ. Δ.Α.Τ.","type":"text","required":True},
            {"name":"client_tax_id","label":"ΑΦΜ αιτούντος","type":"text","required":True},
            {"name":"client_address","label":"Διεύθυνση αιτούντος","type":"text","required":True},
            {"name":"lawyer_name","label":"Ονοματεπώνυμο δικηγόρου","type":"text","required":True},
            {"name":"lawyer_am","label":"ΑΜ δικηγόρου","type":"text","required":True},
            {"name":"lawyer_dsb","label":"Δικηγορικός Σύλλογος","type":"text","required":True},
            {"name":"opponent_name","label":"Ονοματεπώνυμο καθ' ου η αίτηση","type":"text","required":True},
            {"name":"opponent_address","label":"Διεύθυνση καθ' ου","type":"text","required":True},
            {"name":"court","label":"Δικαστήριο προς το οποίο απευθύνεται","type":"text","required":True},
            {"name":"decision_number","label":"Αριθμός / Έτος αρ. απόφασης (π.χ. 1234/2025)","type":"text","required":True},
            {"name":"issuing_court","label":"Δικαστήριο που εξέδωσε την απόφαση","type":"text","required":True},
            {"name":"case_number","label":"Αριθμός υπόθεσης","type":"text","required":False},
            {"name":"execution_act","label":"Πράξη εκτέλεσης / επιταγή προς πληρωμή","type":"text","required":False},
            {"name":"description","label":"Λόγοι αναστολής – ιστορικό","type":"textarea","required":True},
            {"name":"appeal_details","label":"Στοιχεία έφεσης που ασκήθηκε","type":"text","required":False},
            {"name":"irreparable_harm","label":"Επικείμενη ανεπανόρθωτη βλάβη αιτούντος","type":"textarea","required":True},
            {"name":"hearing_date","label":"Ορισθείσα δικάσιμος","type":"date","required":False},
            {"name":"city","label":"Πόλη κατάθεσης","type":"text","required":True},
            {"name":"date","label":"Ημερομηνία","type":"date","required":True},
        ],
        "template": """ΕΝΩΠΙΟΝ ΤΟΥ {{court}}

ΑΙΤΗΣΗ ΑΝΑΣΤΟΛΗΣ ΕΚΤΕΛΕΣΗΣ
(κατ' άρθρα 912-915 Κ.Πολ.Δ.)

ΑΙΤΩΝ/ΑΙΤΟΥΣΑ: {{client_name}}, του/της {{client_father}}, κάτοικος {{client_address}}, Δ.Α.Τ. {{client_id_number}}, Α.Φ.Μ. {{client_tax_id}}, ο/η οποίος/α παρίσταται δια του πληρεξουσίου Δικηγόρου {{lawyer_name}} (ΑΜ {{lawyer_am}}, Δ.Σ. {{lawyer_dsb}}).

ΚΑΘ' ΟΥ Η ΑΙΤΗΣΗ: {{opponent_name}}, κάτοικος {{opponent_address}}.

Α. ΙΣΤΟΡΙΚΟ – ΝΟΜΙΚΗ ΒΑΣΗ

Με την υπ' αριθ. {{decision_number}} απόφαση του {{issuing_court}} (Υπόθεση αριθ. {{case_number}}) ο/η καθ' ου απέκτησε τίτλο εκτελεστό εναντίον μου.

Βάσει της εν λόγω αποφάσεως, ο/η καθ' ου επέδωσε {{execution_act}}.

{{description}}

Β. ΑΣΚΗΘΕΙΣΑ ΕΦΕΣΗ

Κατά της ανωτέρω αποφάσεως άσκησα εμπρόθεσμα Έφεση, στοιχεία: {{appeal_details}}.
Η έφεση έχει ασκηθεί εμπρόθεσμα και νόμιμα και έχει πιθανότητες ευδοκίμησης για τους λόγους που αναπτύσσονται κατωτέρω (άρθρο 912 §1 ΚΠολΔ).

Γ. ΕΠΙΚΕΙΜΕΝΗ ΑΝΕΠΑΝΟΡΘΩΤΗ ΒΛΑΒΗ

{{irreparable_harm}}

Δ. ΝΟΜΙΚΗ ΘΕΜΕΛΙΩΣΗ

Άρθρα 912, 913, 914 §1 Κ.Πολ.Δ.: Η αίτηση αναστολής είναι παραδεκτή εφόσον:
α) Η εκτέλεση θα επιφέρει στον αιτούντα ανεπανόρθωτη ή δυσχερώς επανορθώσιμη βλάβη, ΚΑΙ
β) Πιθανολογείται η ευδοκίμηση του ασκηθέντος ένδικου μέσου.

ΕΑΝ ΔΕΝ ΕΧΕ ΟΡΙΣΘΕΙ ΔΙΚΑΣΙΜΟΣ: Ζητείται ο ορισμός της ταχύτερης δυνατής δικασίμου και κοινοποίηση στον καθ' ου.

Ε. ΑΙΤΗΜΑ

Για τους λόγους αυτούς,

ΖΗΤΩ

1. Να γίνει δεκτή η παρούσα αίτηση.
2. Να ανασταλεί η εκτέλεση της υπ' αριθ. {{decision_number}} αποφάσεως του {{issuing_court}} μέχρι εκδόσεως αποφάσεως επί της ασκηθείσης Εφέσεως / μέχρι τελεσιδικίας.
3. (Επικουρικώς) Να ανασταλεί η εκτέλεση υπό τον όρο καταθέσεως εγγύησης, εκτιμωμένης κατά διακριτική ευχέρεια του Δικαστηρίου.
4. Να καταδικασθεί ο/η καθ' ου στη δικαστική δαπάνη.

Ορισθείσα δικάσιμος: {{hearing_date}}

{{city}}, {{date}}

Ο Πληρεξούσιος Δικηγόρος


_________________________________
{{lawyer_name}}"""
    },

    # ── 5. ΥΠΕΥΘΥΝΗ ΔΗΛΩΣΗ (Ν. 1599/1986) ────────────────────────────────────
    "ypeuthinidilosi": {
        "id": "ypeuthinidilosi",
        "name": "Υπεύθυνη Δήλωση",
        "description": "Υπεύθυνη δήλωση κατ' άρθρο 8 Ν. 1599/1986",
        "category": "Γενικό",
        "fields": [
            {"name":"_case_id","label":"Υπόθεση (συμπληρώνει αυτόματα)","type":"linked","linked_to":"cases","required":False},
            {"name":"_client_id","label":"Εντολέας (αν δεν υπάρχει υπόθεση)","type":"linked","linked_to":"clients","required":False},
            {"name":"client_name","label":"Ονοματεπώνυμο δηλούντος","type":"text","required":True},
            {"name":"client_father","label":"Πατρώνυμο","type":"text","required":True},
            {"name":"client_mother","label":"Μητρώνυμο","type":"text","required":True},
            {"name":"client_birth_date","label":"Ημερομηνία γέννησης","type":"date","required":True},
            {"name":"client_birth_place","label":"Τόπος γέννησης","type":"text","required":True},
            {"name":"client_id_number","label":"Αρ. Δ.Α.Τ. / Διαβατηρίου","type":"text","required":True},
            {"name":"client_id_issuer","label":"Εκδούσα αρχή Δ.Α.Τ.","type":"text","required":False},
            {"name":"client_tax_id","label":"ΑΦΜ","type":"text","required":True},
            {"name":"client_address","label":"Διεύθυνση κατοικίας","type":"text","required":True},
            {"name":"client_phone","label":"Τηλέφωνο","type":"text","required":True},
            {"name":"content","label":"Περιεχόμενο δήλωσης","type":"textarea","required":True},
            {"name":"recipient","label":"Προς (αρχή/υπηρεσία)","type":"text","required":True},
            {"name":"city","label":"Πόλη","type":"text","required":True},
            {"name":"date","label":"Ημερομηνία","type":"date","required":True},
        ],
        "template": """ΥΠΕΥΘΥΝΗ ΔΗΛΩΣΗ
(άρθρο 8 Ν. 1599/1986)

Η ακρίβεια των στοιχείων που υποβάλλονται με αυτή τη δήλωση μπορεί να ελεγχθεί με βάση το αρχείο άλλων υπηρεσιών (άρθρο 8 παρ. 4 Ν. 1599/1986).

ΠΡΟΣ: {{recipient}}

Ο/Η κάτωθι υπογράφων/ουσα:

Επώνυμο: ........................   Όνομα: ........................
Ονοματεπώνυμο: {{client_name}}
Πατρώνυμο: {{client_father}}           Μητρώνυμο: {{client_mother}}
Ημερ. γέννησης: {{client_birth_date}}   Τόπος γέννησης: {{client_birth_place}}
Αρ. Δ.Α.Τ.: {{client_id_number}}       Εκδούσα αρχή: {{client_id_issuer}}
Α.Φ.Μ.: {{client_tax_id}}
Διεύθυνση: {{client_address}}
Τηλέφωνο: {{client_phone}}

ΔΗΛΩΝΩ ΥΠΕΥΘΥΝΩΣ

ότι:

{{content}}

Με πλήρη γνώση των συνεπειών που προβλέπει η κείμενη νομοθεσία για ψευδή δήλωση, ήτοι ποινική δίωξη κατ' άρθρο 22 Ν. 1599/1986 (φυλάκιση τουλάχιστον τριών μηνών).

{{city}}, {{date}}

Ο/Η Δηλών/ούσα


_________________________________
{{client_name}}"""
    },

    # ── 6. ΣΥΜΒΑΣΗ ΠΑΡΟΧΗΣ ΝΟΜΙΚΩΝ ΥΠΗΡΕΣΙΩΝ ─────────────────────────────────
    "symbasi_entolis": {
        "id": "symbasi_entolis",
        "name": "Σύμβαση Παροχής Νομικών Υπηρεσιών",
        "description": "Δικηγορική σύμβαση εντολής παροχής νομικών υπηρεσιών (Ν. 4194/2013)",
        "category": "Γενικό",
        "fields": [
            {"name":"_case_id","label":"Υπόθεση (συμπληρώνει αυτόματα)","type":"linked","linked_to":"cases","required":False},
            {"name":"_client_id","label":"Εντολέας (αν δεν υπάρχει υπόθεση)","type":"linked","linked_to":"clients","required":False},
            {"name":"lawyer_name","label":"Ονοματεπώνυμο δικηγόρου","type":"text","required":True},
            {"name":"lawyer_am","label":"ΑΜ δικηγόρου","type":"text","required":True},
            {"name":"lawyer_dsb","label":"Δικηγορικός Σύλλογος","type":"text","required":True},
            {"name":"lawyer_address","label":"Διεύθυνση γραφείου","type":"text","required":True},
            {"name":"lawyer_tax_id","label":"ΑΦΜ δικηγόρου","type":"text","required":True},
            {"name":"lawyer_doy","label":"ΔΟΥ δικηγόρου","type":"text","required":True},
            {"name":"client_name","label":"Ονοματεπώνυμο εντολέα","type":"text","required":True},
            {"name":"client_father","label":"Πατρώνυμο εντολέα","type":"text","required":True},
            {"name":"client_id_number","label":"Αρ. Δ.Α.Τ. εντολέα","type":"text","required":True},
            {"name":"client_tax_id","label":"ΑΦΜ εντολέα","type":"text","required":True},
            {"name":"client_address","label":"Διεύθυνση εντολέα","type":"text","required":True},
            {"name":"client_phone","label":"Τηλέφωνο εντολέα","type":"text","required":True},
            {"name":"client_email","label":"Email εντολέα","type":"text","required":False},
            {"name":"case_description","label":"Περιγραφή νομικής υπόθεσης / αντικείμενο εντολής","type":"textarea","required":True},
            {"name":"fee_amount","label":"Αμοιβή (€)","type":"text","required":True},
            {"name":"fee_schedule","label":"Τρόπος / χρόνος καταβολής αμοιβής","type":"text","required":True},
            {"name":"expenses_note","label":"Δαπάνες (παράβολα, γραμμάτια κ.λπ.)","type":"text","required":False},
            {"name":"duration","label":"Διάρκεια σύμβασης","type":"text","required":False},
            {"name":"city","label":"Πόλη σύναψης","type":"text","required":True},
            {"name":"date","label":"Ημερομηνία","type":"date","required":True},
        ],
        "template": """ΣΥΜΒΑΣΗ ΠΑΡΟΧΗΣ ΝΟΜΙΚΩΝ ΥΠΗΡΕΣΙΩΝ
(κατ' άρθρα 92 επ. Ν. 4194/2013 – Κώδικας Δικηγόρων)

Στην {{city}}, σήμερα {{date}}, μεταξύ των:

Α) ΔΙΚΗΓΟΡΟΥ: {{lawyer_name}}, Δικηγόρου {{lawyer_dsb}} (ΑΜ {{lawyer_am}}), κατοίκου/εδρεύοντος στη διεύθυνση {{lawyer_address}}, με Α.Φ.Μ. {{lawyer_tax_id}}, ΔΟΥ {{lawyer_doy}} (εφεξής «Δικηγόρος»),

ΚΑΙ

Β) ΕΝΤΟΛΕΑ: {{client_name}}, του/της {{client_father}}, κατοίκου {{client_address}}, κατόχου Δ.Α.Τ. αριθ. {{client_id_number}}, Α.Φ.Μ. {{client_tax_id}}, τηλ. {{client_phone}}, email: {{client_email}} (εφεξής «Εντολέας»),

συμφωνούνται και γίνονται αμοιβαία αποδεκτά τα εξής:

ΑΡΘΡΟ 1 – ΑΝΤΙΚΕΙΜΕΝΟ ΕΝΤΟΛΗΣ

Ο Εντολέας αναθέτει στον Δικηγόρο και ο Δικηγόρος αναλαμβάνει να παράσχει τις παρακάτω νομικές υπηρεσίες:

{{case_description}}

ΑΡΘΡΟ 2 – ΑΜΟΙΒΗ

2.1. Η αμοιβή του Δικηγόρου συμφωνείται σε {{fee_amount}} € (πλέον ΦΠΑ 24% βάσει του άρθρου 21 Ν. 4194/2013).
2.2. Καταβολή: {{fee_schedule}}.
2.3. Δαπάνες (δικαστικά έξοδα, παράβολα, γραμμάτιο ΕΦΚΑ, τέλη κ.λπ.): {{expenses_note}} — βαρύνουν τον Εντολέα επιπλέον της αμοιβής.

ΑΡΘΡΟ 3 – ΥΠΟΧΡΕΩΣΕΙΣ ΔΙΚΗΓΟΡΟΥ

3.1. Ο Δικηγόρος υποχρεούται να παρέχει τις ανωτέρω υπηρεσίες με επιμέλεια, ειλικρίνεια και εχεμύθεια, σύμφωνα με τον Κώδικα Δικηγόρων (Ν. 4194/2013) και τον Κώδικα Δεοντολογίας.
3.2. Ο Δικηγόρος θα ενημερώνει τακτικά τον Εντολέα για την πορεία της υπόθεσης.
3.3. Ο Δικηγόρος τηρεί απόρρητο για κάθε πληροφορία που λαμβάνει εντός της σχέσης εντολής.

ΑΡΘΡΟ 4 – ΥΠΟΧΡΕΩΣΕΙΣ ΕΝΤΟΛΕΑ

4.1. Ο Εντολέας υποχρεούται να καταβάλλει εμπρόθεσμα την αμοιβή και τις δαπάνες.
4.2. Να παρέχει στον Δικηγόρο όλα τα αναγκαία στοιχεία, έγγραφα και πληροφορίες.
4.3. Να ειδοποιεί εγκαίρως τον Δικηγόρο για κάθε μεταβολή στοιχείων επικοινωνίας.

ΑΡΘΡΟ 5 – ΔΙΑΡΚΕΙΑ – ΛΥΣΗ

5.1. Η παρούσα ισχύει: {{duration}}.
5.2. Οποιοσδήποτε συμβαλλόμενος μπορεί να καταγγείλει τη σύμβαση με έγγραφη ειδοποίηση 15 ημερών.
5.3. Σε περίπτωση καταγγελίας, ο Εντολέας υποχρεούται να καταβάλει αμοιβή αναλογική του πεπραγμένου έργου.

ΑΡΘΡΟ 6 – ΕΦΑΡΜΟΣΤΕΟ ΔΙΚΑΙΟ – ΔΙΚΑΙΟΔΟΣΙΑ

Εφαρμοστέο δίκαιο είναι το Ελληνικό. Για κάθε διαφορά από την παρούσα αρμόδια είναι τα Δικαστήρια της {{city}}.

ΑΡΘΡΟ 7 – ΠΡΟΣΤΑΣΙΑ ΔΕΔΟΜΕΝΩΝ (GDPR)

Τα προσωπικά δεδομένα του Εντολέα χρησιμοποιούνται αποκλειστικά για τους σκοπούς εκτέλεσης της παρούσας σύμβασης, σύμφωνα με τον Κανονισμό (ΕΕ) 2016/679 (GDPR) και τον Ν. 4624/2019.

Η παρούσα σύμβαση έχει συνταχθεί σε δύο (2) αντίτυπα, έλαβε δε ο κάθε συμβαλλόμενος από ένα.

ΟΙ ΣΥΜΒΑΛΛΟΜΕΝΟΙ


Ο ΔΙΚΗΓΟΡΟΣ                                  Ο/Η ΕΝΤΟΛΕΑΣ


_____________________                        _____________________
{{lawyer_name}}                               {{client_name}}
(ΑΜ {{lawyer_am}})"""
    },

    # ── 7. ΑΓΩΓΗ (ΑΣΤΙΚΗ) ────────────────────────────────────────────────────
    "agogi": {
        "id": "agogi",
        "name": "Αγωγή (Αστική)",
        "description": "Αστική αγωγή ενώπιον Πρωτοδικείου / Ειρηνοδικείου",
        "category": "Αστικό",
        "fields": [
            {"name":"_case_id","label":"Υπόθεση (συμπληρώνει αυτόματα)","type":"linked","linked_to":"cases","required":False},
            {"name":"court","label":"Δικαστήριο (π.χ. Πρωτοδικείο Αθηνών)","type":"text","required":True},
            {"name":"procedure","label":"Διαδικασία (τακτική / ειδικές / ασφαλιστικά)","type":"text","required":True},
            {"name":"client_name","label":"Ονοματεπώνυμο ενάγοντος","type":"text","required":True},
            {"name":"client_father","label":"Πατρώνυμο ενάγοντος","type":"text","required":True},
            {"name":"client_id_number","label":"Αρ. Δ.Α.Τ. ενάγοντος","type":"text","required":True},
            {"name":"client_tax_id","label":"ΑΦΜ ενάγοντος","type":"text","required":True},
            {"name":"client_address","label":"Διεύθυνση ενάγοντος","type":"text","required":True},
            {"name":"opponent_name","label":"Ονοματεπώνυμο εναγομένου","type":"text","required":True},
            {"name":"opponent_father","label":"Πατρώνυμο εναγομένου","type":"text","required":False},
            {"name":"opponent_address","label":"Διεύθυνση εναγομένου","type":"text","required":True},
            {"name":"lawyer_name","label":"Ονοματεπώνυμο δικηγόρου","type":"text","required":True},
            {"name":"lawyer_am","label":"ΑΜ δικηγόρου","type":"text","required":True},
            {"name":"cause_of_action","label":"Νομική βάση (αίτιο αγωγής)","type":"text","required":True},
            {"name":"law_articles","label":"Άρθρα νόμου (ΑΚ, ΚΠολΔ κ.λπ.)","type":"text","required":True},
            {"name":"facts","label":"Ιστορική βάση – πραγματικά περιστατικά","type":"textarea","required":True},
            {"name":"claim_amount","label":"Ποσό αγωγής (€)","type":"text","required":False},
            {"name":"claim_description","label":"Αίτημα αγωγής (αναλυτικά)","type":"textarea","required":True},
            {"name":"evidence","label":"Αποδεικτικά μέσα","type":"textarea","required":False},
            {"name":"city","label":"Πόλη κατάθεσης","type":"text","required":True},
            {"name":"date","label":"Ημερομηνία","type":"date","required":True},
        ],
        "template": """ΕΝΩΠΙΟΝ ΤΟΥ {{court}}
(Διαδικασία: {{procedure}})

Α Γ Ω Γ Η

ΤΟΥ/ΤΗΣ: {{client_name}}, του/της {{client_father}}, κατοίκου {{client_address}}, Δ.Α.Τ. {{client_id_number}}, Α.Φ.Μ. {{client_tax_id}}, ο/η οποίος/α εκπροσωπείται από τον/την Δικηγόρο {{lawyer_name}} (ΑΜ {{lawyer_am}})
(εφεξής «Ενάγων/ουσα»)

Κ Α Τ Α

ΤΟΥ/ΤΗΣ: {{opponent_name}}, του/της {{opponent_father}}, κατοίκου {{opponent_address}}
(εφεξής «Εναγόμενος/η»)

Α. ΙΣΤΟΡΙΚΗ ΒΑΣΗ

{{facts}}

Β. ΝΟΜΙΚΗ ΒΑΣΗ

{{cause_of_action}}

Άρθρα: {{law_articles}}

Γ. ΑΠΟΔΕΙΚΤΙΚΑ ΜΕΣΑ

{{evidence}}

Δ. ΑΠΟ ΤΑ ΑΝΩΤΕΡΩ ΠΡΟΚΥΠΤΕΙ ότι:
– Η αγωγή είναι νόμιμη, ορισμένη και βάσιμη.
– Αρμόδιο καθ' ύλην και κατά τόπον δικαστήριο είναι το ανωτέρω βάσει των άρθρων 7 επ. και 22 ΚΠολΔ.

ΔΙΑ ΤΑΥΤΑ

ΖΗΤΩ

{{claim_description}}

(Αντικείμενο αγωγής: {{claim_amount}} €)

Να καταδικαστεί ο/η εναγόμενος/η στη δικαστική δαπάνη.
Να κηρυχθεί η απόφαση προσωρινά εκτελεστή.

{{city}}, {{date}}

Ο Πληρεξούσιος Δικηγόρος


_________________________________
{{lawyer_name}} (ΑΜ {{lawyer_am}})"""
    },

    # ── 8. ΕΦΕΣΗ ──────────────────────────────────────────────────────────────
    "efesi": {
        "id": "efesi",
        "name": "Έφεση",
        "description": "Έφεση κατά πρωτόδικης απόφασης (άρθρα 495-524 ΚΠολΔ)",
        "category": "Δικαστικό",
        "fields": [
            {"name":"_case_id","label":"Υπόθεση (συμπληρώνει αυτόματα)","type":"linked","linked_to":"cases","required":False},
            {"name":"court","label":"Εφετείο (π.χ. Εφετείο Αθηνών)","type":"text","required":True},
            {"name":"client_name","label":"Ονοματεπώνυμο εκκαλούντος","type":"text","required":True},
            {"name":"client_father","label":"Πατρώνυμο εκκαλούντος","type":"text","required":True},
            {"name":"client_id_number","label":"Αρ. Δ.Α.Τ. εκκαλούντος","type":"text","required":True},
            {"name":"client_tax_id","label":"ΑΦΜ εκκαλούντος","type":"text","required":True},
            {"name":"client_address","label":"Διεύθυνση εκκαλούντος","type":"text","required":True},
            {"name":"opponent_name","label":"Ονοματεπώνυμο εφεσίβλητου","type":"text","required":True},
            {"name":"opponent_address","label":"Διεύθυνση εφεσίβλητου","type":"text","required":True},
            {"name":"lawyer_name","label":"Ονοματεπώνυμο δικηγόρου","type":"text","required":True},
            {"name":"lawyer_am","label":"ΑΜ δικηγόρου","type":"text","required":True},
            {"name":"first_court","label":"Πρωτόδικο Δικαστήριο","type":"text","required":True},
            {"name":"decision_number","label":"Αριθμός πρωτόδικης απόφασης","type":"text","required":True},
            {"name":"decision_date","label":"Ημερομηνία πρωτόδικης απόφασης","type":"date","required":True},
            {"name":"service_date","label":"Ημερομηνία επίδοσης απόφασης","type":"date","required":False},
            {"name":"grounds","label":"Λόγοι έφεσης (ουσιαστικοί / δικονομικοί)","type":"textarea","required":True},
            {"name":"prejudiced_party","label":"Ζημία / νομικό σφάλμα πρωτόδικης απόφασης","type":"textarea","required":True},
            {"name":"city","label":"Πόλη κατάθεσης","type":"text","required":True},
            {"name":"date","label":"Ημερομηνία","type":"date","required":True},
        ],
        "template": """ΕΝΩΠΙΟΝ ΤΟΥ {{court}}

Ε Φ Ε Σ Η
(κατ' άρθρα 495-524 Κ.Πολ.Δ.)

ΕΚΚΑΛΩΝ/ΟΥΣΑ: {{client_name}}, του/της {{client_father}}, κάτοικος {{client_address}}, Δ.Α.Τ. {{client_id_number}}, Α.Φ.Μ. {{client_tax_id}}, δια του πληρεξουσίου Δικηγόρου {{lawyer_name}} (ΑΜ {{lawyer_am}}).

ΕΦΕΣΙΒΛΗΤΟΣ/Η: {{opponent_name}}, κάτοικος {{opponent_address}}.

ΠΡΟΣΒΑΛΛΟΜΕΝΗ ΑΠΟΦΑΣΗ: Η υπ' αριθ. {{decision_number}} απόφαση του {{first_court}}, ημερομηνίας {{decision_date}}, που επιδόθηκε στις {{service_date}}.

Α. ΠΑΡΑΔΕΚΤΟ ΕΦΕΣΗΣ

Η παρούσα έφεση ασκείται εμπρόθεσμα, εντός της προθεσμίας του άρθρου 518 §1 ΚΠολΔ (30 ημέρες από επίδοση), νόμιμα και παραδεκτά, έχοντας έννομο συμφέρον ο/η εκκαλών/ούσα.

Β. ΛΟΓΟΙ ΕΦΕΣΗΣ

{{grounds}}

Γ. ΝΟΜΙΚΑ ΣΦΑΛΜΑΤΑ ΠΡΩΤΟΔΙΚΗΣ ΑΠΟΦΑΣΗΣ

{{prejudiced_party}}

Δ. ΑΙΤΗΜΑ

Για τους ανωτέρω λόγους,

ΖΗΤΩ

1. Να γίνει τυπικά και ουσιαστικά δεκτή η παρούσα Έφεση.
2. Να εξαφανιστεί η προσβαλλόμενη υπ' αριθ. {{decision_number}} απόφαση.
3. Να γίνει δεκτή η αγωγή / να απορριφθεί η αγωγή (κατά περίπτωση).
4. Να καταδικαστεί ο/η εφεσίβλητος/η στη δικαστική δαπάνη αμφοτέρων των βαθμών.

Επισυνάπτεται:
– Αντίγραφο της προσβαλλόμενης απόφασης
– Αποδεικτικό επίδοσης
– Παράβολο άρθρου 495 §3 ΚΠολΔ

{{city}}, {{date}}

Ο Πληρεξούσιος Δικηγόρος


_________________________________
{{lawyer_name}} (ΑΜ {{lawyer_am}})"""
    },

    # ── 9. ΔΗΛΩΣΗ ΠΟΙΝΙΚΗΣ ΔΙΑΔΙΚΑΣΙΑΣ (ΚΑΤΗΓΟΡΟΥΜΕΝΟΣ) ───────────────────────
    "dilosi_katigoro": {
        "id": "dilosi_katigoro",
        "name": "Δήλωση Παράστασης Πολιτικής Αγωγής",
        "description": "Δήλωση παράστασης πολιτικής αγωγής ενώπιον ποινικού δικαστηρίου (άρθρο 82 ΚΠΔ)",
        "category": "Ποινικό",
        "fields": [
            {"name":"_case_id","label":"Υπόθεση (συμπληρώνει αυτόματα)","type":"linked","linked_to":"cases","required":False},
            {"name":"court","label":"Ποινικό Δικαστήριο","type":"text","required":True},
            {"name":"hearing_date","label":"Δικάσιμος","type":"date","required":True},
            {"name":"case_number","label":"Αριθμός δικογραφίας / εγκλήματος","type":"text","required":False},
            {"name":"client_name","label":"Ονοματεπώνυμο παθόντος (πολιτικώς ενάγοντος)","type":"text","required":True},
            {"name":"client_father","label":"Πατρώνυμο","type":"text","required":True},
            {"name":"client_id_number","label":"Αρ. Δ.Α.Τ.","type":"text","required":True},
            {"name":"client_tax_id","label":"ΑΦΜ","type":"text","required":True},
            {"name":"client_address","label":"Διεύθυνση","type":"text","required":True},
            {"name":"lawyer_name","label":"Ονοματεπώνυμο δικηγόρου","type":"text","required":True},
            {"name":"lawyer_am","label":"ΑΜ δικηγόρου","type":"text","required":True},
            {"name":"accused_name","label":"Ονοματεπώνυμο κατηγορουμένου","type":"text","required":True},
            {"name":"offense","label":"Αδίκημα για το οποίο δικάζεται","type":"text","required":True},
            {"name":"claim_amount","label":"Χρηματική αξίωση (€) — συμβολικό ή ποσό","type":"text","required":True},
            {"name":"description","label":"Συνοπτική περιγραφή ζημίας","type":"textarea","required":True},
            {"name":"city","label":"Πόλη","type":"text","required":True},
            {"name":"date","label":"Ημερομηνία","type":"date","required":True},
        ],
        "template": """ΕΝΩΠΙΟΝ ΤΟΥ {{court}}
(Δικάσιμος: {{hearing_date}} — Αρ. δικογραφίας: {{case_number}})

ΔΗΛΩΣΗ ΠΑΡΑΣΤΑΣΗΣ ΠΟΛΙΤΙΚΗΣ ΑΓΩΓΗΣ
(άρθρα 82-89 Κ.Π.Δ.)

Του/Της: {{client_name}}, του/της {{client_father}}, κατοίκου {{client_address}}, Δ.Α.Τ. {{client_id_number}}, Α.Φ.Μ. {{client_tax_id}}, ο/η οποίος/α εκπροσωπείται από τον/την Δικηγόρο {{lawyer_name}} (ΑΜ {{lawyer_am}}).
(εφεξής «Πολιτικώς Ενάγων/ούσα»)

ΚΑΤΑ

Του/Της κατηγορουμένου/ης: {{accused_name}}, κατηγορουμένου/ης για το αδίκημα: {{offense}}.

ΔΗΛΩΝΩ

ότι παρίσταμαι ως πολιτικώς ενάγων/ούσα στην παρούσα ποινική δίκη, για να αξιώσω εύλογη χρηματική ικανοποίηση λόγω ηθικής βλάβης / περιουσιακής ζημίας.

Α. ΠΕΡΙΓΡΑΦΗ ΖΗΜΙΑΣ

{{description}}

Β. ΑΞΙΩΣΗ

Ζητώ χρηματική ικανοποίηση ποσού {{claim_amount}} €, καθώς και κάθε άλλη νόμιμη αποζημίωση.

Γ. ΝΟΜΙΚΗ ΒΑΣΗ

Άρθρα 82, 83 Κ.Π.Δ., 914, 932 ΑΚ.

Δ. ΑΙΤΗΜΑ

Να γίνει δεκτή η παρούσα δήλωση παράστασης πολιτικής αγωγής.
Να υποχρεωθεί ο/η κατηγορούμενος/η να μου καταβάλει ποσό {{claim_amount}} €.

{{city}}, {{date}}

Ο/Η Πολιτικώς Ενάγων/ούσα           Ο Πληρεξούσιος Δικηγόρος


_____________________                _____________________
{{client_name}}                       {{lawyer_name}}"""
    },
}

@app.get("/api/templates")
async def list_templates(user=Depends(get_current_user)):
    """List all available document templates."""
    return [{"id": t["id"], "name": t["name"], "description": t["description"],
             "category": t["category"], "fields": t["fields"]} for t in TEMPLATES.values()]

@app.get("/api/templates/{template_id}")
async def get_template(template_id: str, user=Depends(get_current_user)):
    if template_id not in TEMPLATES: raise HTTPException(404, "Πρότυπο δεν βρέθηκε")
    return TEMPLATES[template_id]

@app.post("/api/templates/{template_id}/fill")
async def fill_template(
    template_id: str,
    case_id: Optional[str] = None,
    client_id: Optional[str] = None,
    user=Depends(get_current_user)
):
    """Auto-fill template fields from case/client data."""
    if template_id not in TEMPLATES: raise HTTPException(404, "Πρότυπο δεν βρέθηκε")
    tmpl = TEMPLATES[template_id]
    data: dict = {"date": datetime.utcnow().strftime("%d/%m/%Y")}

    # Pull office/lawyer settings for lawyer fields
    settings = await db.settings.find_one({"_id": "office"}) or {}
    if settings:
        data["lawyer_name"]    = settings.get("lawyer_name") or user.get("name", "")
        data["lawyer_am"]      = settings.get("lawyer_am", "")
        data["lawyer_dsb"]     = settings.get("lawyer_dsb", "")
        data["lawyer_address"] = settings.get("address", "")
        data["lawyer_tax_id"]  = settings.get("afm", "")
        data["lawyer_doy"]     = settings.get("doy", "")
        data["city"]           = settings.get("city", "")
    else:
        data["lawyer_name"] = user.get("name", "")

    def _fill_client(cl: dict):
        data["client_name"]    = cl.get("full_name") or cl.get("name", "")
        data["client_tax_id"]  = cl.get("afm") or cl.get("tax_id", "")
        data["client_address"] = cl.get("address", "")
        data["client_phone"]   = cl.get("phone", "")
        data["client_email"]   = cl.get("email", "")

    # Case-based auto-fill
    if case_id:
        case = await db.cases.find_one({"_id": make_id(case_id)})
        if case:
            data["case_title"]       = case.get("title", "")
            data["case_description"] = case.get("description") or case.get("title", "")
            data["case_number"]      = case.get("case_number", "")
            data["court"]            = case.get("court", "")
            data["offense"]          = case.get("offense", "")
            data["law_articles"]     = case.get("law_articles", "")
            # Client(s) linked to case
            cid = case.get("client_id") or (case.get("client_ids") or [None])[0]
            if cid:
                cl = await db.clients.find_one({"_id": make_id(cid)})
                if cl:
                    _fill_client(cl)
            # Assigned lawyer (override generic setting)
            if case.get("assigned_lawyer_id"):
                lawyer = await db.users.find_one({"_id": make_id(case["assigned_lawyer_id"])})
                if lawyer:
                    data["lawyer_name"] = lawyer.get("full_name") or lawyer.get("name", "")
            # Opponent from parties collection
            try:
                parties = await db.case_parties.find(
                    {"case_id": case_id, "party_role": {"$in": ["opponent", "καθ' ου", "εναγόμενος", "κατηγορούμενος"]}}
                ).to_list(1)
            except Exception:
                parties = []
            if parties:
                data["opponent_name"]    = parties[0].get("name", "")
                data["opponent_address"] = parties[0].get("address", "")
                data["accused_name"]     = parties[0].get("name", "")

    # Client-only auto-fill (when no case selected)
    elif client_id:
        cl = await db.clients.find_one({"_id": make_id(client_id)})
        if cl:
            _fill_client(cl)

    return {"template": tmpl, "auto_filled": data}

@app.post("/api/templates/{template_id}/generate")
async def generate_document(template_id: str, body: dict, user=Depends(get_current_user)):
    """Generate a proper .docx file from template with provided field values."""
    if template_id not in TEMPLATES:
        raise HTTPException(404, "Πρότυπο δεν βρέθηκε")
    tmpl = TEMPLATES[template_id]

    # Frontend sends { fields: {...} } — unwrap if needed
    fields: dict = body.get("fields", body)

    # Fill placeholders (skip _-prefixed meta fields like _case_id, _client_id)
    text: str = tmpl["template"]
    for key, value in fields.items():
        if key.startswith("_"):
            continue
        text = text.replace("{{" + key + "}}", str(value) if value else "________")
    text = re.sub(r"\{\{[^}]+\}\}", "________", text)

    # Build .docx in memory using python-docx
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import io

        doc = DocxDocument()

        # Page margins (Greek legal standard: ~2.5 cm)
        for section in doc.sections:
            section.top_margin    = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin   = Cm(3.0)
            section.right_margin  = Cm(2.5)

        # Default font
        style = doc.styles["Normal"]
        font  = style.font
        font.name = "Times New Roman"
        font.size = Pt(12)

        lines = text.split("\n")
        for line in lines:
            stripped = line.strip()
            p = doc.add_paragraph()

            # Detect title / heading lines (all-caps, short, centred)
            is_title = (
                stripped.isupper()
                and len(stripped) > 3
                and len(stripped) < 80
                and not stripped.startswith("–")
                and not stripped.startswith("-")
            )
            if is_title:
                run = p.add_run(stripped)
                run.bold = True
                run.font.size = Pt(13)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif stripped.startswith("_____"):
                run = p.add_run(line)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif stripped == "":
                p.paragraph_format.space_after = Pt(4)
            else:
                run = p.add_run(line)
                run.font.size = Pt(12)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        # RFC 5987 encoding for non-ASCII filenames
        from urllib.parse import quote
        name_ascii  = tmpl["id"]  # always safe ASCII
        name_utf8   = quote(tmpl["name"].replace(" ", "_") + ".docx", safe="")
        disposition = f"attachment; filename=\"{name_ascii}.docx\"; filename*=UTF-8''{name_utf8}"
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": disposition},
        )

    except ImportError:
        import io
        from urllib.parse import quote
        buf = io.BytesIO(text.encode("utf-8"))
        buf.seek(0)
        name_utf8   = quote(tmpl["name"].replace(" ", "_") + ".txt", safe="")
        disposition = f"attachment; filename=\"{tmpl['id']}.txt\"; filename*=UTF-8''{name_utf8}"
        return StreamingResponse(
            buf,
            media_type="text/plain; charset=utf-8",
            headers={"Content-Disposition": disposition},
        )

# ══════════════════════════════════════════════════════════════════════════════
# SETTINGS (Office info, team, notifications)
# ══════════════════════════════════════════════════════════════════════════════
@app.get("/api/settings")
async def get_settings(user=Depends(get_current_user)):
    s = await db.settings.find_one({"_id": "office"})
    return serialize(s) if s else {"name":"","address":"","phone":"","email":"","afm":"","website":""}

@app.put("/api/settings")
async def update_settings(data: dict, user=Depends(require_role(UserRole.ADMIN))):
    await db.settings.update_one({"_id": "office"}, {"$set": data}, upsert=True)
    await audit("UPDATE_SETTINGS", user["id"], "settings")
    return {"ok": True}

# ══════════════════════════════════════════════════════════════════════════════
# SEPARATE EXPENSES (with categories)
# ══════════════════════════════════════════════════════════════════════════════
class ExpenseRequest(BaseModel):
    case_id: str
    description: str
    amount: float
    category: str = "other"
    date: Optional[datetime] = None
    notes: Optional[str] = None
    receipt_ref: Optional[str] = None   # αρ. απόδειξης / παραστατικού

@app.get("/api/expenses")
async def list_expenses(
    user=Depends(get_current_user),
    case_id: Optional[str] = None,
    client_id: Optional[str] = None,
    lawyer_id: Optional[str] = None,
):
    q: dict = {}
    if case_id:   q["case_id"]   = case_id
    if client_id: q["client_id"] = client_id
    if lawyer_id: q["created_by"] = lawyer_id
    # Lawyers see only their own cases
    if user["role"] == UserRole.LAWYER.value and not case_id:
        cids = [str(c) for c in await db.cases.distinct("_id", {"assigned_lawyer_id": user["id"]})]
        q["case_id"] = {"$in": cids}

    entries = await db.expenses_log.find(q).sort("date", -1).to_list(None)
    serialized = [serialize(e) for e in entries]

    # ── Aggregated summaries ──────────────────────────────────────────────────
    total = sum(e.get("amount", 0) for e in entries)

    # Per-case totals
    by_case: dict = {}
    for e in entries:
        k = e.get("case_id", "")
        if k not in by_case:
            by_case[k] = {"case_id": k, "case_title": e.get("case_title", ""), "case_number": e.get("case_number", ""), "client_name": e.get("client_name", ""), "total": 0, "count": 0}
        by_case[k]["total"] += e.get("amount", 0)
        by_case[k]["count"] += 1

    # Per-lawyer totals
    by_lawyer: dict = {}
    for e in entries:
        k = e.get("created_by", "")
        name = e.get("created_by_name", "—")
        if k not in by_lawyer:
            by_lawyer[k] = {"lawyer_id": k, "lawyer_name": name, "total": 0, "count": 0}
        by_lawyer[k]["total"] += e.get("amount", 0)
        by_lawyer[k]["count"] += 1

    # Per-category totals
    by_category: dict = {}
    for e in entries:
        k = e.get("category", "other")
        if k not in by_category:
            by_category[k] = {"category": k, "total": 0, "count": 0}
        by_category[k]["total"] += e.get("amount", 0)
        by_category[k]["count"] += 1

    # This-month total
    now = datetime.utcnow()
    month_total = sum(
        e.get("amount", 0) for e in entries
        if isinstance(e.get("date"), datetime)
        and e["date"].year == now.year and e["date"].month == now.month
    )

    return {
        "entries": serialized,
        "total": round(total, 2),
        "month_total": round(month_total, 2),
        "count": len(entries),
        "by_case":     sorted(by_case.values(),     key=lambda x: x["total"], reverse=True),
        "by_lawyer":   sorted(by_lawyer.values(),   key=lambda x: x["total"], reverse=True),
        "by_category": sorted(by_category.values(), key=lambda x: x["total"], reverse=True),
    }

@app.post("/api/expenses", status_code=201)
async def create_expense(req: ExpenseRequest, user=Depends(get_current_user)):
    await _check_case_access(req.case_id, user)
    if req.amount <= 0:
        raise HTTPException(400, "Το ποσό πρέπει να είναι θετικό")
    case = await db.cases.find_one({"_id": make_id(req.case_id)})
    client_id = case.get("client_id", "") if case else ""
    client_name = await get_client_name(client_id) if client_id else ""
    doc = {
        "case_id":          req.case_id,
        "case_title":       case.get("title", "") if case else "",
        "case_number":      case.get("case_number", "") if case else "",
        "client_id":        client_id,
        "client_name":      client_name,
        "description":      sanitize_string(req.description),
        "amount":           req.amount,
        "category":         req.category,
        "date":             req.date or datetime.utcnow(),
        "notes":            sanitize_string(req.notes) if req.notes else None,
        "receipt_ref":      req.receipt_ref,
        "created_by":       user["id"],
        "created_by_name":  user.get("name", ""),
        "created_at":       datetime.utcnow(),
    }
    r = await db.expenses_log.insert_one(doc)
    await audit("CREATE_EXPENSE", user["id"], "expense", str(r.inserted_id))
    doc["_id"] = r.inserted_id
    return serialize(doc)

@app.delete("/api/expenses/{expense_id}")
async def delete_expense(expense_id: str, user=Depends(require_role(UserRole.ADMIN))):
    await db.expenses_log.delete_one({"_id": make_id(expense_id)})
    await audit("DELETE_EXPENSE", user["id"], "expense", expense_id)
    return {"ok": True}

@app.get("/api/health")
async def health():
    try: await client.admin.command("ping"); dbs = "connected"
    except: dbs = "disconnected"
    return {"status": "ok" if dbs == "connected" else "degraded", "database": dbs, "timestamp": datetime.utcnow().isoformat()}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("server:app", host="0.0.0.0", port=8000, reload=False)


# ═══════════════════════════════════════════════════════════════════════════════
#  LEGAL OS — OPERATING MODEL EXTENSIONS
#  Leads/Pipeline · Workflow · Billing Engine · KPIs · Checklists
# ═══════════════════════════════════════════════════════════════════════════════

# ── Enums ─────────────────────────────────────────────────────────────────────

class LeadStage(str, Enum):
    NEW          = "new_lead"
    CONTACTED    = "contacted"
    BOOKED       = "consultation_booked"
    DONE         = "consultation_done"
    PROPOSAL     = "proposal_sent"
    RETAINER     = "retainer_paid"
    OPENED       = "matter_opened"
    LOST         = "lost"
    FOLLOW_UP    = "follow_up_later"

class LeadSource(str, Enum):
    REFERRAL = "referral"
    LOGISTIS = "logistis"
    GOOGLE   = "google"
    SOCIAL   = "social"
    WALK_IN  = "walk_in"
    OTHER    = "other"

class ProfitClass(str, Enum):
    A = "A"
    B = "B"
    C = "C"

class WorkflowTemplate(str, Enum):
    EXODIKO  = "exodiko"
    AGOGI    = "agogi"
    RETAINER = "retainer"
    ERGATIKO = "ergatiko"
    DIATAGI  = "diatagi"
    SYNTHESI = "synthesi"
    GENERAL  = "general"

# ── Workflow stage lists per template ─────────────────────────────────────────

WORKFLOW_STAGES: dict = {
    "exodiko":  ["Intake","Legal review","Drafting","Client approval","Dispatch","Follow-up","Closing"],
    "agogi":    ["Intake & Viability","Pre-filing","Filing","Pending hearing","Hearing","Post-hearing","Closing"],
    "retainer": ["Onboarding","Monthly open","Ad-hoc intake","Active handling","Monthly review","Monthly close","Annual review"],
    "ergatiko": ["Intake","Assessment","Pre-action","Filing","Hearing","Resolution","Closing"],
    "diatagi":  ["Intake","Documentation","Filing","Service","Monitoring","Closing"],
    "synthesi": ["Intake","Drafting","Review","Client approval","Execution","Closing"],
    "general":  ["Intake","In review","Drafting","Filed / sent","Pending response","Active handling","Closing"],
}

WORKFLOW_CHECKLISTS: dict = {
    "exodiko": [
        "Άνοιγμα client record","Conflict check","Λήψη εγγράφων","Ανάθεση owner","Προκαταβολή",
        "Νομική αξιολόγηση","Drafting εξωδίκου","Εσωτερικός έλεγχος","Έγκριση πελάτη",
        "Αποστολή / Επίδοση","Ανέβασμα αποδεικτικού","Follow-up D+3","Follow-up D+7","Follow-up D+15","Κλείσιμο matter",
    ],
    "agogi": [
        "Συνέντευξη πελάτη","Conflict check","Συλλογή εγγράφων","Νομική αξιολόγηση βιωσιμότητας",
        "Κοστολόγηση","Σύμβαση εντολής","Προκαταβολή","Συλλογή αποδεικτικών",
        "Νομική έρευνα","Drafting αγωγής","Εσωτερικός έλεγχος","Τελική έγκριση",
        "Κατάθεση αγωγής","Κοινοποίηση","Καταγραφή δικασίμου","Παράσταση","Ενημέρωση πελάτη","Κλείσιμο / Εκτέλεση",
    ],
    "retainer": [
        "Σύναψη σύμβασης retainer","Ορισμός included / extra scope","KYC πελάτη","Billing setup",
        "Άνοιγμα monthly cycle","Καταγραφή αιτημάτων","Classify: included / extra",
        "Monthly summary","Monthly invoice","Upsell check",
    ],
    "ergatiko": [
        "Συνέντευξη πελάτη","Conflict check","Σύμβαση εργασίας","Αποδεικτικά απολύσεως",
        "Νομική αξιολόγηση","Κοστολόγηση","Drafting","Filing","Παράσταση","Κλείσιμο",
    ],
    "diatagi": [
        "Λήψη εγγράφων απαίτησης","Conflict check","Προκαταβολή",
        "Drafting διαταγής","Κατάθεση","Επίδοση","Παρακολούθηση","Κλείσιμο",
    ],
    "synthesi": [
        "Intake","Conflict check","Λήψη σχεδίου / briefing","Νομική ανάλυση",
        "Drafting","Εσωτερικός έλεγχος","Έγκριση πελάτη","Υπογραφή","Κλείσιμο",
    ],
    "general": [
        "Intake","Conflict check","Ανάθεση owner","Προκαταβολή","Νομική εργασία","Follow-up","Κλείσιμο",
    ],
}

BILLING_TRIGGERS: dict = {
    "exodiko":  {"Intake": "Consultation fee", "Drafting": "Drafting fee", "Dispatch": "Dispatch fee", "Closing": "Final balance"},
    "agogi":    {"Intake & Viability": "Assessment fee + Advance", "Pre-filing": "Drafting fee", "Filing": "Filing fee", "Pending hearing": "Pre-hearing milestone", "Hearing": "Hearing fee", "Closing": "Final balance"},
    "retainer": {"Onboarding": "Πρώτο μηνιαίο invoice", "Monthly close": "Monthly invoice", "Annual review": "Renewal"},
    "ergatiko": {"Intake": "Consultation fee", "Filing": "Filing fee", "Hearing": "Hearing fee", "Closing": "Final balance"},
    "diatagi":  {"Intake": "Fixed drafting + Advance", "Filing": "Filing fee", "Service": "Service fee"},
    "synthesi": {"Intake": "Consultation fee", "Execution": "Fixed fee"},
    "general":  {"Intake": "Advance", "Closing": "Final balance"},
}

# ── Lead Models ───────────────────────────────────────────────────────────────

class LeadRequest(BaseModel):
    name: str
    phone: Optional[str] = None
    email: Optional[str] = None
    source: str = "other"
    referral_partner: Optional[str] = None
    case_type: Optional[str] = None
    stage: str = "new_lead"
    urgency: str = "normal"
    assigned_to: Optional[str] = None
    notes: Optional[str] = None
    next_action: Optional[str] = None
    next_action_date: Optional[datetime] = None
    consultation_fee: Optional[float] = None

# ── Lead Endpoints ────────────────────────────────────────────────────────────

@app.get("/api/leads")
async def list_leads(stage: Optional[str] = None, user=Depends(get_current_user)):
    q = {} if not stage else {"stage": stage}
    leads = await db.leads.find(q).sort("created_at", -1).to_list(None)
    return [serialize(l) for l in leads]

@app.get("/api/leads/pipeline")
async def leads_pipeline(user=Depends(get_current_user)):
    leads = await db.leads.find({}).to_list(None)
    stages = [s.value for s in LeadStage]
    pipeline = {s: [] for s in stages}
    for lead in leads:
        s = lead.get("stage", "new_lead")
        if s in pipeline:
            pipeline[s].append(serialize(lead))
    return [{"stage": s, "count": len(pipeline[s]), "leads": pipeline[s]} for s in stages]

@app.get("/api/leads/{lead_id}")
async def get_lead(lead_id: str, user=Depends(get_current_user)):
    lead = await db.leads.find_one({"_id": make_id(lead_id)})
    if not lead: raise HTTPException(404)
    return serialize(lead)

@app.post("/api/leads", status_code=201)
async def create_lead(req: LeadRequest, user=Depends(get_current_user)):
    doc = req.dict()
    doc["created_at"] = datetime.utcnow()
    doc["created_by"] = user["id"]
    doc["updated_at"] = datetime.utcnow()
    result = await db.leads.insert_one(doc)
    await log_action(user["id"], "lead_created", str(result.inserted_id), {"name": req.name})
    return serialize(await db.leads.find_one({"_id": result.inserted_id}))

@app.put("/api/leads/{lead_id}")
async def update_lead(lead_id: str, req: LeadRequest, user=Depends(get_current_user)):
    doc = req.dict(); doc["updated_at"] = datetime.utcnow()
    result = await db.leads.update_one({"_id": make_id(lead_id)}, {"$set": doc})
    if result.matched_count == 0: raise HTTPException(404)
    return serialize(await db.leads.find_one({"_id": make_id(lead_id)}))

@app.patch("/api/leads/{lead_id}/stage")
async def move_lead_stage(lead_id: str, stage: str, user=Depends(get_current_user)):
    await db.leads.update_one({"_id": make_id(lead_id)}, {"$set": {"stage": stage, "updated_at": datetime.utcnow()}})
    await log_action(user["id"], "lead_stage_changed", lead_id, {"stage": stage})
    return serialize(await db.leads.find_one({"_id": make_id(lead_id)}))

@app.delete("/api/leads/{lead_id}")
async def delete_lead(lead_id: str, user=Depends(get_current_user)):
    if user["role"] not in [UserRole.ADMIN.value, UserRole.LAWYER.value]: raise HTTPException(403)
    await db.leads.delete_one({"_id": make_id(lead_id)})
    return {"ok": True}

# ── Workflow Extensions for Cases ─────────────────────────────────────────────

class WorkflowUpdateRequest(BaseModel):
    profit_class: Optional[str] = None
    workflow_template: Optional[str] = None
    matter_stage: Optional[str] = None
    last_client_update: Optional[datetime] = None
    next_billing_trigger: Optional[str] = None
    estimated_value: Optional[float] = None
    complexity_score: Optional[int] = None
    success_likelihood: Optional[str] = None
    opposing_party: Optional[str] = None
    counterparty_lawyer: Optional[str] = None
    claim_value: Optional[float] = None
    missing_documents: Optional[str] = None
    client_source: Optional[str] = None
    risk_level: Optional[str] = None

@app.get("/api/workflow/templates")
async def get_workflow_templates(user=Depends(get_current_user)):
    return {"stages": WORKFLOW_STAGES, "billing_triggers": BILLING_TRIGGERS}

@app.patch("/api/cases/{case_id}/workflow")
async def update_case_workflow(case_id: str, req: WorkflowUpdateRequest, user=Depends(get_current_user)):
    case = await db.cases.find_one({"_id": make_id(case_id)})
    if not case: raise HTTPException(404)
    update = {k: v for k, v in req.dict().items() if v is not None}
    update["updated_at"] = datetime.utcnow()
    await db.cases.update_one({"_id": make_id(case_id)}, {"$set": update})
    await log_action(user["id"], "workflow_updated", case_id, update)
    return serialize(await db.cases.find_one({"_id": make_id(case_id)}))

@app.get("/api/cases/workflow/stuck")
async def get_stuck_workflow_cases(days: int = 14, user=Depends(get_current_user)):
    cutoff = datetime.utcnow() - timedelta(days=days)
    q = {"status": {"$in": ["active","pending"]},
         "$or": [{"updated_at": {"$lt": cutoff}}, {"updated_at": {"$exists": False}, "created_at": {"$lt": cutoff}}]}
    cases = await db.cases.find(q).to_list(None)
    return [serialize(c) for c in cases]

@app.get("/api/cases/workflow/no-next-action")
async def cases_without_next_action(user=Depends(get_current_user)):
    q = {"status": {"$in": ["active","pending"]},
         "$or": [{"next_action": None},{"next_action": ""},{"next_action": {"$exists": False}}]}
    cases = await db.cases.find(q).to_list(None)
    return [serialize(c) for c in cases]

# ── Checklists ────────────────────────────────────────────────────────────────

@app.get("/api/checklists/templates")
async def get_checklist_templates(user=Depends(get_current_user)):
    return WORKFLOW_CHECKLISTS

@app.get("/api/cases/{case_id}/checklist")
async def get_case_checklist(case_id: str, user=Depends(get_current_user)):
    checklist = await db.checklists.find_one({"case_id": case_id})
    if not checklist:
        case = await db.cases.find_one({"_id": make_id(case_id)})
        if not case: raise HTTPException(404)
        template = case.get("workflow_template", "general")
        items = WORKFLOW_CHECKLISTS.get(template, WORKFLOW_CHECKLISTS["general"])
        doc = {
            "case_id": case_id,
            "template": template,
            "items": [{"text": i, "done": False, "done_at": None, "done_by": None} for i in items],
            "created_at": datetime.utcnow(),
        }
        res = await db.checklists.insert_one(doc)
        checklist = await db.checklists.find_one({"_id": res.inserted_id})
    return serialize(checklist)

@app.patch("/api/cases/{case_id}/checklist/{item_index}")
async def toggle_checklist_item(case_id: str, item_index: int, done: bool, user=Depends(get_current_user)):
    checklist = await db.checklists.find_one({"case_id": case_id})
    if not checklist: raise HTTPException(404)
    items = checklist.get("items", [])
    if item_index >= len(items): raise HTTPException(400, "Invalid index")
    items[item_index]["done"] = done
    items[item_index]["done_at"] = datetime.utcnow().isoformat() if done else None
    items[item_index]["done_by"] = user.get("name", "") if done else None
    await db.checklists.update_one({"case_id": case_id}, {"$set": {"items": items}})
    return {"ok": True, "items": items}

@app.post("/api/cases/{case_id}/checklist/reset")
async def reset_checklist(case_id: str, template: str = "general", user=Depends(get_current_user)):
    items = WORKFLOW_CHECKLISTS.get(template, WORKFLOW_CHECKLISTS["general"])
    doc = {"case_id": case_id, "template": template,
           "items": [{"text": i, "done": False, "done_at": None, "done_by": None} for i in items],
           "updated_at": datetime.utcnow()}
    await db.checklists.update_one({"case_id": case_id}, {"$set": doc}, upsert=True)
    return serialize(await db.checklists.find_one({"case_id": case_id}))

# ── Billing Reminders ─────────────────────────────────────────────────────────

class BillingReminderRequest(BaseModel):
    invoice_id: str
    case_id: Optional[str] = None
    client_name: str
    amount: float
    due_date: datetime
    status: str = "sent"
    notes: Optional[str] = None

@app.get("/api/billing/reminders")
async def list_reminders(status: Optional[str] = None, user=Depends(get_current_user)):
    q = {} if not status else {"status": status}
    reminders = await db.billing_reminders.find(q).sort("due_date", 1).to_list(None)
    now = datetime.utcnow()
    result = []
    for r in reminders:
        s = serialize(r)
        due = r.get("due_date")
        if due:
            s["days_overdue"] = max(0, (now - due).days) if due < now else 0
        result.append(s)
    return result

@app.post("/api/billing/reminders", status_code=201)
async def create_reminder(req: BillingReminderRequest, user=Depends(get_current_user)):
    doc = req.dict(); doc["created_at"] = datetime.utcnow(); doc["last_action_at"] = datetime.utcnow()
    res = await db.billing_reminders.insert_one(doc)
    return serialize(await db.billing_reminders.find_one({"_id": res.inserted_id}))

@app.patch("/api/billing/reminders/{reminder_id}/advance")
async def advance_reminder(reminder_id: str, notes: Optional[str] = None, user=Depends(get_current_user)):
    CYCLE = ["sent","reminder_1","reminder_2","escalated","collection","paid"]
    reminder = await db.billing_reminders.find_one({"_id": make_id(reminder_id)})
    if not reminder: raise HTTPException(404)
    current = reminder.get("status","sent")
    idx = CYCLE.index(current) if current in CYCLE else 0
    next_status = CYCLE[min(idx + 1, len(CYCLE) - 1)]
    update = {"status": next_status, "last_action_at": datetime.utcnow()}
    if notes: update["notes"] = notes
    await db.billing_reminders.update_one({"_id": make_id(reminder_id)}, {"$set": update})
    return serialize(await db.billing_reminders.find_one({"_id": make_id(reminder_id)}))

@app.patch("/api/billing/reminders/{reminder_id}/status")
async def set_reminder_status(reminder_id: str, status: str, user=Depends(get_current_user)):
    await db.billing_reminders.update_one({"_id": make_id(reminder_id)}, {"$set": {"status": status, "last_action_at": datetime.utcnow()}})
    return serialize(await db.billing_reminders.find_one({"_id": make_id(reminder_id)}))

@app.get("/api/billing/collection-rate")
async def get_collection_rate(user=Depends(get_current_user)):
    invoices = await db.invoices.find({}).to_list(None)
    total = sum(i.get("amount", 0) for i in invoices)
    paid = sum(i.get("amount", 0) for i in invoices if i.get("payment_status") == "paid")
    partial = sum(i.get("amount_paid", 0) for i in invoices if i.get("payment_status") == "partial")
    collected = paid + partial
    return {
        "total_billed": total, "total_collected": collected,
        "collection_rate": round(collected / total * 100, 1) if total else 0,
        "invoice_count": len(invoices),
        "paid_count": sum(1 for i in invoices if i.get("payment_status") == "paid"),
    }

# ── KPI Summary ───────────────────────────────────────────────────────────────

@app.get("/api/kpi/summary")
async def kpi_summary(user=Depends(get_current_user)):
    now = datetime.utcnow()
    month_start = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
    week_start  = now - timedelta(days=now.weekday())
    cutoff_14   = now - timedelta(days=14)
    cutoff_72h  = now - timedelta(hours=72)

    # Leads
    total_leads    = await db.leads.count_documents({})
    new_leads      = await db.leads.count_documents({"stage": "new_lead"})
    week_leads     = await db.leads.count_documents({"created_at": {"$gte": week_start}})
    converted      = await db.leads.count_documents({"stage": "matter_opened"})
    stale_leads    = await db.leads.count_documents({"stage": "new_lead", "created_at": {"$lt": cutoff_72h}})
    conv_rate      = round(converted / total_leads * 100, 1) if total_leads else 0

    # Cases
    active_cases   = await db.cases.count_documents({"status": {"$in": ["active","pending"]}})
    stuck_cases    = await db.cases.count_documents({
        "status": {"$in": ["active","pending"]},
        "$or": [{"updated_at": {"$lt": cutoff_14}}, {"updated_at": {"$exists": False}, "created_at": {"$lt": cutoff_14}}]
    })
    no_next        = await db.cases.count_documents({
        "status": {"$in": ["active","pending"]},
        "$or": [{"next_action": None},{"next_action": ""},{"next_action": {"$exists": False}}]
    })

    # Financials
    invoices       = await db.invoices.find({}).to_list(None)
    total_billed   = sum(i.get("amount",0) for i in invoices)
    total_paid_all = sum(i.get("amount",0) for i in invoices if i.get("payment_status")=="paid")
    month_billed   = sum(i.get("amount",0) for i in invoices if i.get("created_at", datetime.min) >= month_start)
    overdue_amount = sum(i.get("amount",0) for i in invoices if i.get("payment_status") in ["pending","partial"] and i.get("due_date") and i["due_date"] < now)
    overdue_count  = sum(1 for i in invoices if i.get("payment_status") in ["pending","partial"] and i.get("due_date") and i["due_date"] < now)
    coll_rate      = round(total_paid_all / total_billed * 100, 1) if total_billed else 0

    # Revenue by case type
    cases_list = await db.cases.find({}).to_list(None)
    case_map   = {str(c["_id"]): c for c in cases_list}
    rev_by_type: dict = {}
    for inv in invoices:
        if inv.get("payment_status") == "paid":
            c = case_map.get(str(inv.get("case_id","")), {})
            cat = c.get("legal_category","Άλλο") or "Άλλο"
            rev_by_type[cat] = rev_by_type.get(cat,0) + inv.get("amount",0)

    # Revenue by source (leads → matters)
    rev_by_source: dict = {}
    leads_converted = await db.leads.find({"stage": "matter_opened"}).to_list(None)
    for l in leads_converted:
        src = l.get("source","other")
        rev_by_source[src] = rev_by_source.get(src,0) + 1

    return {
        "leads": {"total": total_leads, "new": new_leads, "this_week": week_leads,
                  "converted": converted, "conversion_rate": conv_rate, "stale": stale_leads},
        "cases": {"active": active_cases, "stuck_14d": stuck_cases, "no_next_action": no_next},
        "financials": {"month_billed": month_billed, "total_outstanding": total_billed - total_paid_all,
                       "overdue_amount": overdue_amount, "overdue_count": overdue_count, "collection_rate": coll_rate},
        "revenue_by_type": rev_by_type,
        "revenue_by_source": rev_by_source,
    }

@app.get("/api/billing/overdue")
async def get_overdue_invoices(user=Depends(get_current_user)):
    now = datetime.utcnow()
    invoices = await db.invoices.find({
        "payment_status": {"$in": ["pending", "partial"]},
        "due_date": {"$lt": now}
    }).sort("due_date", 1).to_list(None)
    result = []
    for inv in invoices:
        s = serialize(inv)
        s["days_overdue"] = max(0, (now - inv["due_date"]).days) if inv.get("due_date") else 0
        # Try to enrich with client name
        if inv.get("client_id"):
            try:
                client = await db.clients.find_one({"_id": make_id(str(inv["client_id"]))})
            except:
                client = None
            if client:
                s["client_name"] = client.get("name", "")
        result.append(s)
    return result


# ═══════════════════════════════════════════════════════════════════════════════
#  NOMOS ONE — EXTENSIONS v2: Payments · Hearings · Email · Lindy · PDF · Excel
# ═══════════════════════════════════════════════════════════════════════════════

import smtplib
import io
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from fpdf import FPDF
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment
from openpyxl.utils import get_column_letter
import httpx as _httpx

FIRM_DISPLAY_NAME = os.getenv("ADMIN_NAME", "Σκοτάνης & Συνεργάτες")
FIRM_EMAIL_DISPLAY = os.getenv("ADMIN_EMAIL", "christos@skotanislaw.com")
LINDY_WEBHOOK_URL = os.getenv("LINDY_WEBHOOK_URL", "https://chat.lindy.ai/christos-skotaniss-workspace/lindy/legal-document-extractor-69db65ca7cf5099909310fa8/tasks")

# ── Payment Models ─────────────────────────────────────────────────────────────

class PaymentRequest(BaseModel):
    invoice_id: Optional[str] = None
    case_id: Optional[str] = None
    client_id: Optional[str] = None
    client_name: Optional[str] = None
    amount: float
    payment_method: str = "bank_transfer"
    payment_date: datetime
    notes: Optional[str] = None
    reference: Optional[str] = None

# ── Payment Endpoints ──────────────────────────────────────────────────────────

@app.get("/api/payments")
async def list_payments(case_id: Optional[str] = None, client_id: Optional[str] = None, user=Depends(get_current_user)):
    q: dict = {}
    if case_id: q["case_id"] = case_id
    if client_id: q["client_id"] = client_id
    payments = await db.payments.find(q).sort("payment_date", -1).to_list(None)
    return [serialize(p) for p in payments]

@app.post("/api/payments", status_code=201)
async def create_payment(req: PaymentRequest, user=Depends(get_current_user)):
    doc = req.dict()
    doc["created_at"] = datetime.utcnow()
    doc["created_by"] = user["id"]
    if req.invoice_id:
        try:
            invoice = await db.invoices.find_one({"_id": make_id(req.invoice_id)})
            if invoice:
                total = float(invoice.get("total", invoice.get("amount", 0)))
                existing = await db.payments.find({"invoice_id": req.invoice_id}).to_list(None)
                paid_so_far = sum(float(p.get("amount", 0)) for p in existing)
                new_paid = paid_so_far + req.amount
                new_status = "paid" if new_paid >= total else "partial"
                await db.invoices.update_one(
                    {"_id": make_id(req.invoice_id)},
                    {"$set": {"payment_status": new_status, "amount_paid": new_paid}}
                )
        except Exception as e:
            logger.warning(f"Could not update invoice status: {e}")
    r = await db.payments.insert_one(doc)
    await audit("CREATE_PAYMENT", user["id"], "payment", str(r.inserted_id))
    return serialize(await db.payments.find_one({"_id": r.inserted_id}))

@app.get("/api/payments/{payment_id}")
async def get_payment(payment_id: str, user=Depends(get_current_user)):
    p = await db.payments.find_one({"_id": make_id(payment_id)})
    if not p: raise HTTPException(404)
    return serialize(p)

@app.delete("/api/payments/{payment_id}")
async def delete_payment(payment_id: str, user=Depends(require_role(UserRole.ADMIN))):
    await db.payments.delete_one({"_id": make_id(payment_id)})
    await audit("DELETE_PAYMENT", user["id"], "payment", payment_id)
    return {"ok": True}

@app.get("/api/cases/{case_id}/payments")
async def get_case_payments(case_id: str, user=Depends(get_current_user)):
    payments = await db.payments.find({"case_id": case_id}).sort("payment_date", -1).to_list(None)
    return [serialize(p) for p in payments]

# ── Hearing Models ─────────────────────────────────────────────────────────────

class HearingRequest(BaseModel):
    case_id: str
    court: str
    hearing_date: datetime
    judge: Optional[str] = None
    notes: Optional[str] = None
    outcome: Optional[str] = None
    next_hearing: Optional[datetime] = None
    status: str = "scheduled"

# ── Hearing Endpoints ──────────────────────────────────────────────────────────

@app.get("/api/hearings")
async def list_hearings(user=Depends(get_current_user)):
    hearings = await db.hearings.find({}).sort("hearing_date", 1).to_list(None)
    return [serialize(h) for h in hearings]

@app.get("/api/cases/{case_id}/hearings")
async def get_case_hearings(case_id: str, user=Depends(get_current_user)):
    hearings = await db.hearings.find({"case_id": case_id}).sort("hearing_date", -1).to_list(None)
    return [serialize(h) for h in hearings]

@app.post("/api/hearings", status_code=201)
async def create_hearing(req: HearingRequest, user=Depends(get_current_user)):
    await _check_payment_gate(req.case_id, f"Δικάσιμος: {req.court}")
    doc = req.dict()
    doc["created_at"] = datetime.utcnow()
    doc["created_by"] = user["id"]
    r = await db.hearings.insert_one(doc)
    await audit("CREATE_HEARING", user["id"], "hearing", str(r.inserted_id))
    return serialize(await db.hearings.find_one({"_id": r.inserted_id}))

@app.put("/api/hearings/{hearing_id}")
async def update_hearing(hearing_id: str, req: HearingRequest, user=Depends(get_current_user)):
    doc = req.dict()
    doc["updated_at"] = datetime.utcnow()
    result = await db.hearings.update_one({"_id": make_id(hearing_id)}, {"$set": doc})
    if result.matched_count == 0: raise HTTPException(404)
    return serialize(await db.hearings.find_one({"_id": make_id(hearing_id)}))

@app.delete("/api/hearings/{hearing_id}")
async def delete_hearing(hearing_id: str, user=Depends(get_current_user)):
    await db.hearings.delete_one({"_id": make_id(hearing_id)})
    return {"ok": True}

# ── Email ──────────────────────────────────────────────────────────────────────

class EmailSendRequest(BaseModel):
    to_email: str
    to_name: Optional[str] = None
    subject: str
    body_html: str
    body_text: Optional[str] = None
    invoice_id: Optional[str] = None

@app.post("/api/email/send")
async def send_email(req: EmailSendRequest, user=Depends(get_current_user)):
    settings_doc = await db.settings.find_one({"_id": "global"}) or {}
    smtp_host = settings_doc.get("smtp_host") or os.getenv("SMTP_HOST", "")
    smtp_port = int(settings_doc.get("smtp_port") or os.getenv("SMTP_PORT", "587"))
    smtp_user = settings_doc.get("smtp_user") or os.getenv("SMTP_USER", "")
    smtp_pass = settings_doc.get("smtp_pass") or os.getenv("SMTP_PASS", "")
    from_email = settings_doc.get("notification_email") or smtp_user or FIRM_EMAIL_DISPLAY
    log_doc = {
        "to": req.to_email, "subject": req.subject,
        "sent_by": user["id"], "sent_at": datetime.utcnow(),
        "invoice_id": req.invoice_id, "status": "pending"
    }
    # Sender identity: noreply address + sending user's name + reply-to their personal email
    noreply_addr  = os.getenv("SMTP_FROM", from_email)
    sender_name   = user.get("full_name") or user.get("name") or FIRM_DISPLAY_NAME
    sender_email  = user.get("email") or ""
    display_from  = f"{sender_name} — {FIRM_DISPLAY_NAME} <{noreply_addr}>"

    if not smtp_host or not smtp_user:
        log_doc["status"] = "placeholder_logged"
        await db.email_logs.insert_one(log_doc)
        logger.info(f"EMAIL PLACEHOLDER → {req.to_email}: {req.subject}")
        return {"ok": True, "status": "placeholder", "message": f"SMTP not configured. Logged email to {req.to_email}"}
    try:
        msg = MIMEMultipart("alternative")
        msg["Subject"] = req.subject
        msg["From"]    = display_from
        msg["To"]      = f"{req.to_name or ''} <{req.to_email}>"
        if sender_email:
            msg["Reply-To"] = sender_email
        if req.body_text:
            msg.attach(MIMEText(req.body_text, "plain", "utf-8"))
        msg.attach(MIMEText(req.body_html, "html", "utf-8"))
        with smtplib.SMTP(smtp_host, smtp_port, timeout=10) as srv:
            srv.ehlo()
            srv.starttls()
            srv.login(smtp_user, smtp_pass)
            srv.sendmail(noreply_addr, req.to_email, msg.as_string())
        log_doc["status"] = "sent"
        await db.email_logs.insert_one(log_doc)
        await audit("EMAIL_SENT", user["id"], "email", req.to_email)
        return {"ok": True, "status": "sent"}
    except Exception as e:
        log_doc["status"] = "failed"
        log_doc["error"] = str(e)
        await db.email_logs.insert_one(log_doc)
        logger.error(f"Email send failed: {e}")
        raise HTTPException(500, f"Αποτυχία αποστολής: {str(e)}")

@app.get("/api/email/logs")
async def get_email_logs(user=Depends(require_role(UserRole.ADMIN))):
    logs = await db.email_logs.find({}).sort("sent_at", -1).limit(100).to_list(None)
    return [serialize(l) for l in logs]

# ── Lindy AI Proxy ─────────────────────────────────────────────────────────────

class LindyForwardRequest(BaseModel):
    message: Optional[str] = None
    source: str = "manual"
    metadata: Optional[dict] = None

@app.post("/api/lindy/forward")
async def forward_to_lindy(req: LindyForwardRequest, user=Depends(get_current_user)):
    webhook_url = os.getenv("LINDY_WEBHOOK_URL", LINDY_WEBHOOK_URL)
    payload = {
        "message": req.message,
        "source": req.source,
        "metadata": req.metadata or {},
        "sent_by": user.get("email"),
        "timestamp": datetime.utcnow().isoformat(),
    }
    try:
        async with _httpx.AsyncClient(timeout=15) as http:
            resp = await http.post(webhook_url, json=payload)
            return {"ok": True, "status": resp.status_code, "response": resp.text[:500]}
    except Exception as e:
        logger.error(f"Lindy proxy failed: {e}")
        return {"ok": False, "error": str(e)}

# ── Invoice PDF ────────────────────────────────────────────────────────────────

@app.get("/api/invoices/{invoice_id}/pdf")
async def generate_invoice_pdf(invoice_id: str, user=Depends(get_current_user)):
    inv = await db.invoices.find_one({"_id": make_id(invoice_id)})
    if not inv: raise HTTPException(404)
    settings_doc = await db.settings.find_one({"_id": "global"}) or {}
    firm = settings_doc.get("firm_name", FIRM_DISPLAY_NAME)
    firm_afm = settings_doc.get("firm_afm", "—")
    firm_address = settings_doc.get("firm_address", "Αθήνα, Ελλάδα")
    firm_email = settings_doc.get("notification_email", FIRM_EMAIL_DISPLAY)

    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_fill_color(7, 18, 32)
    pdf.rect(0, 0, 210, 48, 'F')
    pdf.set_font("Helvetica", "B", 18)
    pdf.set_text_color(198, 167, 94)
    pdf.set_xy(15, 10)
    pdf.cell(0, 10, firm, ln=True)
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(138, 160, 184)
    pdf.set_xy(15, 23)
    pdf.cell(0, 5, f"{firm_address}  |  {firm_email}  |  AFM: {firm_afm}", ln=True)
    inv_number = inv.get("invoice_number", str(inv["_id"])[-6:])
    pdf.set_font("Helvetica", "B", 13)
    pdf.set_text_color(198, 167, 94)
    pdf.set_xy(130, 13)
    pdf.cell(0, 8, f"TIMOΛΟΓΙΟ #{inv_number}")
    inv_date = inv.get("created_at", datetime.utcnow())
    if isinstance(inv_date, str):
        try: inv_date = datetime.fromisoformat(inv_date.replace("Z",""))
        except: inv_date = datetime.utcnow()
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(138, 160, 184)
    pdf.set_xy(130, 24)
    pdf.cell(0, 5, f"Ημ/νια: {inv_date.strftime('%d/%m/%Y')}")
    pdf.set_text_color(30, 30, 30)
    pdf.set_y(58)
    pdf.set_font("Helvetica", "B", 10)
    pdf.set_fill_color(237, 241, 245)
    pdf.set_x(15)
    pdf.cell(85, 7, "ΠΡΟΣ:", fill=True)
    pdf.set_x(110)
    pdf.cell(85, 7, "ΥΠΟΘΕΣΗ:", fill=True)
    pdf.ln(7)
    pdf.set_font("Helvetica", "", 10)
    pdf.set_x(15)
    pdf.cell(85, 6, (inv.get("client_name") or "—")[:40])
    pdf.set_x(110)
    pdf.cell(85, 6, (inv.get("case_title") or "—")[:40])
    pdf.ln(6)
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(100, 100, 100)
    if inv.get("client_afm"):
        pdf.set_x(15)
        pdf.cell(85, 5, f"AFM: {inv['client_afm']}")
    if inv.get("is_professional"):
        pdf.set_x(110)
        pdf.cell(85, 5, "[Επιτηδευματιας - Παρακρατηση 20%]")
    pdf.ln(12)
    pdf.set_text_color(30, 30, 30)
    pdf.set_fill_color(7, 18, 32)
    pdf.set_text_color(198, 167, 94)
    pdf.set_font("Helvetica", "B", 9)
    pdf.set_x(15)
    pdf.cell(120, 8, "ΠΕΡΙΓΡΑΦΗ ΥΠΗΡΕΣΙΩΝ", fill=True)
    pdf.cell(55, 8, "ΠΟΣΟ", fill=True, align="R")
    pdf.ln(8)
    pdf.set_text_color(30, 30, 30)
    pdf.set_font("Helvetica", "", 10)
    pdf.set_fill_color(250, 251, 252)
    pdf.set_x(15)
    desc = (inv.get("description") or "Νομικές υπηρεσίες")[:80]
    amount = float(inv.get("amount", 0))
    pdf.cell(120, 8, desc, fill=True)
    pdf.cell(55, 8, f"EUR {amount:,.2f}", fill=True, align="R")
    pdf.ln(16)
    vat_rate = float(inv.get("vat_rate", 24))
    vat = float(inv.get("vat_amount", amount * vat_rate / 100))
    withholding = float(inv.get("withholding_tax", 0))
    total = float(inv.get("total", amount + vat))
    net_payable = float(inv.get("net_payable", total - withholding))
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(80, 80, 80)
    for lbl, val, sign in [(f"ΦΠΑ {int(vat_rate)}%", vat, "+")]:
        pdf.set_x(110)
        pdf.cell(55, 6, lbl)
        pdf.cell(20, 6, f"{sign} EUR {val:,.2f}", align="R")
        pdf.ln(6)
    if withholding > 0:
        pdf.set_x(110)
        pdf.cell(55, 6, "Παρακρατηση 20%")
        pdf.cell(20, 6, f"- EUR {withholding:,.2f}", align="R")
        pdf.ln(6)
    pdf.set_fill_color(7, 18, 32)
    pdf.set_text_color(198, 167, 94)
    pdf.set_font("Helvetica", "B", 11)
    pdf.set_x(110)
    pdf.cell(75, 9, f"ΣΥΝΟΛΟ: EUR {total:,.2f}", fill=True, align="R")
    pdf.ln(9)
    if withholding > 0:
        pdf.set_fill_color(22, 163, 74)
        pdf.set_text_color(255, 255, 255)
        pdf.set_font("Helvetica", "B", 10)
        pdf.set_x(110)
        pdf.cell(75, 9, f"ΕΙΣΠΡΑΚΤΕΟ: EUR {net_payable:,.2f}", fill=True, align="R")
        pdf.ln(14)
    if withholding > 0:
        pdf.set_text_color(100, 100, 100)
        pdf.set_font("Helvetica", "I", 7)
        pdf.set_x(15)
        pdf.multi_cell(180, 4, "* Η παρακρατηση φορου 20% αποδιδεται απο τον πελατη (επιτηδευματια) στην ΑΑΔΕ.")
    pdf.set_y(-45)
    pdf.set_text_color(30, 30, 30)
    pdf.set_font("Helvetica", "", 9)
    pdf.set_x(15)
    pdf.cell(80, 5, "_" * 28)
    pdf.set_x(115)
    pdf.cell(80, 5, "_" * 28)
    pdf.ln(6)
    pdf.set_x(15)
    pdf.cell(80, 5, "Υπογραφη Δικηγορου")
    pdf.set_x(115)
    pdf.cell(80, 5, "Υπογραφη Πελατη")
    pdf.set_y(-18)
    pdf.set_fill_color(237, 241, 245)
    pdf.rect(0, pdf.get_y(), 210, 18, 'F')
    pdf.set_text_color(120, 120, 120)
    pdf.set_font("Helvetica", "", 8)
    pdf.set_x(15)
    pdf.cell(0, 9, f"{firm}  |  {firm_address}  |  AFM: {firm_afm}")
    pdf_bytes = pdf.output()
    return StreamingResponse(
        io.BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename=invoice_{inv_number}.pdf"}
    )

# ── Excel Exports ──────────────────────────────────────────────────────────────

def _make_wb_header(ws, headers):
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.font = Font(bold=True, color="FFFFFF", size=10)
        cell.fill = PatternFill("solid", fgColor="071220")
        cell.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 18

@app.get("/api/invoices/export/excel")
async def export_invoices_excel(user=Depends(get_current_user)):
    invoices = await db.invoices.find({}).sort("created_at", -1).to_list(None)
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Τιμολόγια"
    headers = ["Αρ.Τιμολ.", "Υπόθεση", "Πελάτης", "Ημερομηνία", "Ακαθάριστα", "ΦΠΑ", "Παρακράτηση", "Σύνολο", "Καθαρό", "Κατάσταση"]
    _make_wb_header(ws, headers)
    alt = PatternFill("solid", fgColor="F5F7FA")
    for i, inv in enumerate(invoices, 2):
        d = inv.get("created_at")
        dstr = d.strftime("%d/%m/%Y") if isinstance(d, datetime) else (str(d)[:10] if d else "")
        row = [
            inv.get("invoice_number", str(inv["_id"])[-6:]),
            inv.get("case_title", ""),
            inv.get("client_name", ""),
            dstr,
            float(inv.get("amount", 0)),
            float(inv.get("vat_amount", 0)),
            float(inv.get("withholding_tax", 0)),
            float(inv.get("total", inv.get("amount", 0))),
            float(inv.get("net_payable", inv.get("total", 0))),
            inv.get("payment_status", "pending"),
        ]
        for col, val in enumerate(row, 1):
            cell = ws.cell(row=i, column=col, value=val)
            if i % 2 == 0: cell.fill = alt
            if col in (5, 6, 7, 8, 9): cell.number_format = '#,##0.00 €'
    for col, w in enumerate([14,26,22,13,14,12,16,14,14,12], 1):
        ws.column_dimensions[get_column_letter(col)].width = w
    ws2 = wb.create_sheet("Μηνιαία")
    _make_wb_header(ws2, ["Μήνας","Πλήθος","Ακαθάριστα","ΦΠΑ 24%","Παρακράτηση","Καθαρό"])
    monthly: dict = {}
    for inv in invoices:
        d2 = inv.get("created_at", datetime.utcnow())
        if isinstance(d2, str):
            try: d2 = datetime.fromisoformat(d2.replace("Z",""))
            except: d2 = datetime.utcnow()
        key = f"{d2.year}-{d2.month:02d}"
        label = d2.strftime("%B %Y")
        if key not in monthly:
            monthly[key] = {"label": label, "n": 0, "g": 0, "v": 0, "w": 0, "net": 0}
        monthly[key]["n"] += 1
        monthly[key]["g"] += float(inv.get("amount", 0))
        monthly[key]["v"] += float(inv.get("vat_amount", 0))
        monthly[key]["w"] += float(inv.get("withholding_tax", 0))
        monthly[key]["net"] += float(inv.get("net_payable", inv.get("total", 0)))
    for i, (_, m) in enumerate(sorted(monthly.items(), reverse=True), 2):
        row2 = [m["label"], m["n"], m["g"], m["v"], m["w"], m["net"]]
        for col, val in enumerate(row2, 1):
            cell = ws2.cell(row=i, column=col, value=val)
            if col >= 3: cell.number_format = '#,##0.00 €'
    out = io.BytesIO()
    wb.save(out); out.seek(0)
    return StreamingResponse(out, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=invoices.xlsx"})

@app.get("/api/clients/export/excel")
async def export_clients_excel(user=Depends(get_current_user)):
    cl = await db.clients.find({}).sort("full_name", 1).to_list(None)
    wb = openpyxl.Workbook()
    ws = wb.active; ws.title = "Πελάτες"
    _make_wb_header(ws, ["Ονοματεπώνυμο","ΑΦΜ","Τηλέφωνο","Email","Διεύθυνση","Τύπος","Κατάσταση","Εγγραφή"])
    alt = PatternFill("solid", fgColor="F5F7FA")
    for i, c in enumerate(cl, 2):
        d = c.get("created_at")
        row = [c.get("full_name",""), c.get("afm",""), c.get("phone",""), c.get("email",""),
               c.get("address",""), c.get("client_type","individual"),
               "Ενεργός" if c.get("is_active") != False else "Ανενεργός",
               d.strftime("%d/%m/%Y") if isinstance(d, datetime) else ""]
        for col, val in enumerate(row, 1):
            cell = ws.cell(row=i, column=col, value=val)
            if i % 2 == 0: cell.fill = alt
    for col, w in enumerate([26,14,14,26,26,16,12,14],1):
        ws.column_dimensions[get_column_letter(col)].width = w
    out = io.BytesIO(); wb.save(out); out.seek(0)
    return StreamingResponse(out, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=clients.xlsx"})

@app.get("/api/cases/export/excel")
async def export_cases_excel(user=Depends(get_current_user)):
    cases = await db.cases.find({}).sort("created_at", -1).to_list(None)
    wb = openpyxl.Workbook()
    ws = wb.active; ws.title = "Υποθέσεις"
    _make_wb_header(ws, ["Κωδικός","Τίτλος","Πελάτης","Κατηγορία","Κατάσταση","Δικηγόρος","Ημ.Δημ."])
    alt = PatternFill("solid", fgColor="F5F7FA")
    for i, c in enumerate(cases, 2):
        d = c.get("created_at")
        row = [c.get("case_number",""), c.get("title",""), c.get("client_name",""),
               c.get("category", c.get("legal_category","")), c.get("status",""),
               c.get("assigned_lawyer",""),
               d.strftime("%d/%m/%Y") if isinstance(d, datetime) else ""]
        for col, val in enumerate(row, 1):
            cell = ws.cell(row=i, column=col, value=val)
            if i % 2 == 0: cell.fill = alt
    for col, w in enumerate([14,30,24,14,14,20,12],1):
        ws.column_dimensions[get_column_letter(col)].width = w
    out = io.BytesIO(); wb.save(out); out.seek(0)
    return StreamingResponse(out, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=cases.xlsx"})

# ══════════════════════════════════════════════════════════════════════════════
# API v1 ROUTES - MOBILE & PWA SUPPORT
# ══════════════════════════════════════════════════════════════════════════════

@app.post("/api/v1/auth/register-device", status_code=201)
async def register_device_v1(
    req: RegisterDeviceRequest,
    user=Depends(get_current_user),
    request: Request = None
):
    """Register a new device for push notifications"""
    if req.device_type not in ["ios", "android", "web", "desktop"]:
        raise HTTPException(status_code=400, detail="Invalid device_type")

    try:
        device_service = get_device_service(db)
        result = await device_service.register_device(
            user_id=user["id"],
            device_name=req.device_name,
            device_type=req.device_type,
            push_token=req.push_token,
            app_version=req.app_version,
            user_agent=request.headers.get("user-agent", "") if request else ""
        )

        if "error" in result:
            raise HTTPException(status_code=400, detail=result["error"])

        await audit("device_registered", user["id"], "device", details={
            "device_type": req.device_type,
            "is_new": result.get("is_new", False)
        })

        return result
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Device registration failed: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to register device")


@app.get("/api/v1/auth/register-device")
async def list_user_devices_v1(user=Depends(get_current_user)):
    """Get all devices registered for current user"""
    try:
        device_service = get_device_service(db)
        devices = await device_service.get_user_devices(user["id"])
        return devices
    except Exception as e:
        logger.error(f"Failed to list devices: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to retrieve devices")


@app.post("/api/v1/auth/register-device/{device_id}/trust")
async def trust_device_v1(
    device_id: str,
    req: TrustDeviceRequest,
    user=Depends(get_current_user)
):
    """Mark device as trusted (skip 2FA for 30 days)"""
    try:
        device_service = get_device_service(db)
        success = await device_service.trust_device(
            device_id=device_id,
            user_id=user["id"],
            device_name=req.device_name
        )

        if not success:
            raise HTTPException(status_code=404, detail="Device not found")

        await audit("device_trusted", user["id"], "device", resource_id=device_id)
        return {"status": "trusted"}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to trust device: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to trust device")


@app.delete("/api/v1/auth/register-device/{device_id}")
async def unregister_device_v1(
    device_id: str,
    user=Depends(get_current_user)
):
    """Unregister/delete a device"""
    try:
        device_service = get_device_service(db)
        success = await device_service.unregister_device(
            device_id=device_id,
            user_id=user["id"]
        )

        if not success:
            raise HTTPException(status_code=404, detail="Device not found")

        await audit("device_unregistered", user["id"], "device", resource_id=device_id)
        return {"status": "unregistered"}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to unregister device: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to unregister device")


@app.get("/api/v1/cases/sync")
async def delta_sync_v1(
    last_sync: Optional[datetime] = None,
    device_id: Optional[str] = None,
    user=Depends(get_current_user)
):
    """Delta sync: Get only case data modified since last sync (optimized for mobile)"""
    try:
        sync_time = last_sync or (datetime.utcnow() - timedelta(days=30))

        # Get cases assigned to user that have been modified
        cases_cursor = db.cases.find({
            "$or": [
                {"assigned_lawyer_id": ObjectId(user["_id"])},
                {"assigned_secretary_id": ObjectId(user["_id"])}
            ],
            "updated_at": {"$gte": sync_time}
        }).sort("updated_at", -1).limit(100)

        cases = []
        async for case in cases_cursor:
            cases.append(serialize(case))

        # Get documents modified since last sync
        documents_cursor = db.documents.find({
            "updated_at": {"$gte": sync_time}
        }).sort("updated_at", -1).limit(100)

        documents = []
        async for doc in documents_cursor:
            documents.append(serialize(doc))

        return {
            "cases": cases,
            "documents": documents,
            "messages": [],
            "updates": [],
            "last_sync_at": datetime.utcnow().isoformat(),
            "has_more": False
        }
    except Exception as e:
        logger.error(f"Delta sync failed: {str(e)}")
        raise HTTPException(status_code=500, detail="Sync failed")


@app.post("/api/v1/auth/logout")
async def logout_device_v1(
    device_id: Optional[str] = None,
    user=Depends(get_current_user)
):
    """Logout from a specific device or all devices"""
    try:
        if device_id:
            device_service = get_device_service(db)
            await device_service.revoke_device_trust(device_id, user["id"])
            await audit("device_logout", user["id"], "device", resource_id=device_id)
        else:
            # Logout from all devices
            await db.users.update_one(
                {"_id": user["_id"]},
                {"$set": {"trusted_devices": []}}
            )
            await audit("logout_all_devices", user["id"], "user")

        return {"status": "logged_out"}
    except Exception as e:
        logger.error(f"Logout failed: {str(e)}")
        raise HTTPException(status_code=500, detail="Logout failed")


@app.get("/api/v1/auth/me")
async def get_current_user_info_v1(user=Depends(get_current_user)):
    """Get current authenticated user info (v1 enhanced with devices)"""
    try:
        device_service = get_device_service(db)
        devices = await device_service.get_user_devices(user["id"])

        user_copy = serialize(user)
        user_copy["devices"] = devices
        user_copy["app_preferences"] = user.get("app_preferences", {
            "notification_enabled": True,
            "offline_mode_enabled": True
        })
        user_copy.pop("password", None)

        return user_copy
    except Exception as e:
        logger.error(f"Failed to get user info: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to retrieve user info")


@app.get("/api/v1/health")
async def health_check_v1():
    """API health check and version information"""
    return {
        "status": "ok",
        "version": "1.0.0",
        "min_app_version": "1.0.0",
        "required_app_version": None,
        "timestamp": datetime.utcnow().isoformat()
    }


@app.get("/api/v1/config/app")
async def get_app_config_v1(user: Optional[Dict] = None):
    """Get app configuration for mobile clients"""
    return {
        "api_version": "1.0.0",
        "features": {
            "offline_mode": True,
            "push_notifications": True,
            "websocket_messaging": True,
            "two_factor_auth": False
        },
        "websocket_url": "ws://localhost:8000/ws"
    }


# ═══════════════════════════════════════════════════════════════════════════
#  GREEK CALENDAR & DEADLINE ENGINE
# ═══════════════════════════════════════════════════════════════════════════

from datetime import date as _date_t

def _orthodox_easter(year: int) -> _date_t:
    """Compute Greek Orthodox Easter (Julian → Gregorian +13 days)."""
    c = year % 19
    d = (19 * c + 15) % 30
    e = (2 * (year % 4) + 4 * (year % 7) - d + 34) % 7
    month = (d + e + 114) // 31
    day = (d + e + 114) % 31 + 1
    return _date_t(year, month, day) + timedelta(days=13)

def _greek_holidays_set(year: int) -> set:
    e = _orthodox_easter(year)
    return {
        _date_t(year, 1, 1),   # Πρωτοχρονιά
        _date_t(year, 1, 6),   # Θεοφάνεια
        e - timedelta(48),      # Καθαρά Δευτέρα
        _date_t(year, 3, 25),  # 25η Μαρτίου
        e - timedelta(2),       # Μεγάλη Παρασκευή
        e,                      # Κυριακή Πάσχα
        e + timedelta(1),       # Δευτέρα Πάσχα
        _date_t(year, 5, 1),   # Εργατική Πρωτομαγιά
        e + timedelta(50),      # Αγίου Πνεύματος
        _date_t(year, 8, 15),  # Κοίμηση Θεοτόκου
        _date_t(year, 10, 28), # 28η Οκτωβρίου
        _date_t(year, 12, 25), # Χριστούγεννα
        _date_t(year, 12, 26), # Σύναξη Θεοτόκου
    }

def _greek_holidays_list(year: int) -> list:
    e = _orthodox_easter(year)
    return [
        {"date": _date_t(year, 1, 1).isoformat(),  "name": "Πρωτοχρονιά"},
        {"date": _date_t(year, 1, 6).isoformat(),  "name": "Θεοφάνεια"},
        {"date": (e - timedelta(48)).isoformat(),   "name": "Καθαρά Δευτέρα"},
        {"date": _date_t(year, 3, 25).isoformat(), "name": "25η Μαρτίου"},
        {"date": (e - timedelta(2)).isoformat(),    "name": "Μεγάλη Παρασκευή"},
        {"date": e.isoformat(),                     "name": "Κυριακή του Πάσχα"},
        {"date": (e + timedelta(1)).isoformat(),    "name": "Δευτέρα του Πάσχα"},
        {"date": _date_t(year, 5, 1).isoformat(),  "name": "Εργατική Πρωτομαγιά"},
        {"date": (e + timedelta(50)).isoformat(),   "name": "Αγίου Πνεύματος"},
        {"date": _date_t(year, 8, 15).isoformat(), "name": "Κοίμηση Θεοτόκου"},
        {"date": _date_t(year, 10, 28).isoformat(),"name": "28η Οκτωβρίου"},
        {"date": _date_t(year, 12, 25).isoformat(),"name": "Χριστούγεννα"},
        {"date": _date_t(year, 12, 26).isoformat(),"name": "Σύναξη Θεοτόκου"},
    ]

def _court_vacations_list(year: int) -> list:
    e = _orthodox_easter(year)
    next_y = year + 1 if year < 9999 else year
    return [
        {
            "start": (e - timedelta(4)).isoformat(),
            "end":   (e + timedelta(8)).isoformat(),
            "name":  "Πασχαλινές Δικαστικές Διακοπές",
            "type":  "easter",
        },
        {
            "start": _date_t(year, 7, 1).isoformat(),
            "end":   _date_t(year, 9, 15).isoformat(),
            "name":  "Θερινές Δικαστικές Διακοπές",
            "type":  "summer",
        },
        {
            "start": _date_t(year, 12, 24).isoformat(),
            "end":   _date_t(next_y, 1, 6).isoformat(),
            "name":  "Χριστουγεννιάτικες Δικαστικές Διακοπές",
            "type":  "christmas",
        },
    ]

def _in_vacation(d: _date_t, vacations: list) -> bool:
    for v in vacations:
        if _date_t.fromisoformat(v["start"]) <= d <= _date_t.fromisoformat(v["end"]):
            return True
    return False

def _is_working_day(d: _date_t, holidays: set) -> bool:
    return d.weekday() < 5 and d not in holidays

def _next_working_day(d: _date_t, holidays: set) -> _date_t:
    while not _is_working_day(d, holidays):
        d += timedelta(days=1)
    return d

def _compute_deadline(start_iso: str, days: int, suspended: bool) -> dict:
    start = _date_t.fromisoformat(start_iso)
    # Collect holidays/vacations for up to 3 years ahead
    all_holidays: set = set()
    all_vacations: list = []
    for yr in range(start.year, start.year + 4):
        all_holidays |= _greek_holidays_set(yr)
        all_vacations += _court_vacations_list(yr)

    # ΚΠολΔ 144: counting starts from day AFTER the trigger event
    current = start + timedelta(days=1)
    remaining = days
    suspended_days = 0

    while remaining > 0:
        if suspended and _in_vacation(current, all_vacations):
            suspended_days += 1  # clock pauses during recess
        else:
            remaining -= 1
        if remaining > 0:
            current += timedelta(days=1)

    original = current
    # ΚΠολΔ 145: if last day is weekend/holiday, extend to next working day
    final = _next_working_day(current, all_holidays)

    # Identify which vacations overlap
    overlapping = [v["name"] for v in all_vacations
                   if _date_t.fromisoformat(v["start"]) <= final
                   and _date_t.fromisoformat(v["end"]) >= start]

    return {
        "start_date": start_iso,
        "deadline_date": final.isoformat(),
        "original_date": original.isoformat(),
        "extended_due_to_holiday": final != original,
        "suspended_days": suspended_days,
        "total_calendar_days": (final - start).days,
        "overlapping_vacations": overlapping,
    }

# ── Deadline catalog ────────────────────────────────────────────────────────
DEADLINE_CATALOG = {
    "ΚΠολΔ": [
        {"id": "efesi",        "label": "Έφεση (ΚΠολΔ 518)",                   "days": 30,   "suspended": True,  "note": "30 ημέρες από επίδοση απόφασης. Αναστέλλεται κατά δικαστικές διακοπές (ΚΠολΔ 147)."},
        {"id": "anairesi",     "label": "Αναίρεση (ΚΠολΔ 564)",                "days": 60,   "suspended": True,  "note": "60 ημέρες. Αναστέλλεται κατά διακοπές."},
        {"id": "tritanakopi",  "label": "Τριτανακοπή (ΚΠολΔ 586)",             "days": 60,   "suspended": True,  "note": "60 ημέρες από δημοσίευση απόφασης."},
        {"id": "anakopi_erim", "label": "Ανακοπή Ερημοδικίας (ΚΠολΔ 501)",    "days": 15,   "suspended": True,  "note": "15 ημέρες από επίδοση ερήμην απόφασης. Αναστέλλεται."},
        {"id": "anakopi_ektl", "label": "Ανακοπή κατ' Εκτέλεσης (ΚΠολΔ 934)", "days": 45,   "suspended": False, "note": "45 ημέρες από επίδοση. ΔΕΝ αναστέλλεται κατά διακοπές."},
        {"id": "antifesi",     "label": "Αντέφεση (ΚΠολΔ 523)",                "days": 30,   "suspended": True,  "note": "Ίδια προθεσμία με έφεση."},
        {"id": "prosthiki",    "label": "Προσθήκη-Αντίκρουση (ΚΠολΔ 237§2)",  "days": 100,  "suspended": True,  "note": "100 ημέρες από κατάθεση αγωγής."},
        {"id": "efesi_erim",   "label": "Έφεση κατά Ερήμην (ΚΠολΔ 501§2)",    "days": 15,   "suspended": True,  "note": "15 ημέρες από επίδοση."},
        {"id": "anakopi_dpl",  "label": "Ανακοπή Διαταγής Πληρωμής (ΚΠολΔ 632)", "days": 15, "suspended": True, "note": "15 εργάσιμες ημέρες από επίδοση."},
        {"id": "klitefsi",     "label": "Κλήτευση πριν δικάσιμο (ΚΠολΔ 228)", "days": 60,   "suspended": False, "note": "Τουλάχιστον 60 ημέρες πριν δικάσιμο."},
    ],
    "ΚΠΔ": [
        {"id": "efesi_kpd",    "label": "Έφεση (ΚΠΔ 474)",                     "days": 10,   "suspended": False, "note": "10 ημέρες. ΔΕΝ αναστέλλεται. Υπέρ κατηγορουμένου από δημοσίευση."},
        {"id": "anairesi_kpd", "label": "Αναίρεση (ΚΠΔ 509)",                  "days": 30,   "suspended": False, "note": "30 ημέρες. ΔΕΝ αναστέλλεται."},
        {"id": "anakopi_kpd",  "label": "Ανακοπή Ερημοδικίας (ΚΠΔ 341)",      "days": 5,    "suspended": False, "note": "5 ημέρες από επίδοση ερήμην."},
        {"id": "prosfygi_kpd", "label": "Προσφυγή (ΚΠΔ 464)",                  "days": 10,   "suspended": False, "note": "10 ημέρες."},
    ],
    "ΑΚ": [
        {"id": "p20",  "label": "Γενική Παραγραφή 20ετής (ΑΚ 249)",  "days": 7300, "suspended": False, "note": "20 χρόνια γενική παραγραφή."},
        {"id": "p5",   "label": "5ετής Παραγραφή (ΑΚ 250)",          "days": 1825, "suspended": False, "note": "5 χρόνια: ενοίκια, τόκοι, μισθοί."},
        {"id": "p3",   "label": "3ετής Παραγραφή (ΑΚ 867)",          "days": 1095, "suspended": False, "note": "3 χρόνια: μισθώματα, αποζημίωση εργολάβου."},
        {"id": "p2",   "label": "2ετής Παραγραφή (ΑΚ 937)",          "days": 730,  "suspended": False, "note": "2 χρόνια από γνώση ζημίας (αδικοπραξία)."},
        {"id": "p1",   "label": "1ετής Παραγραφή (ΑΚ 682)",          "days": 365,  "suspended": False, "note": "1 χρόνο."},
    ],
    "ΔΔ": [
        {"id": "prosfygi_dd",   "label": "Προσφυγή (ΚΔΔ 66)",            "days": 60, "suspended": False, "note": "60 ημέρες από κοινοποίηση πράξης."},
        {"id": "efesi_dd",      "label": "Έφεση Διοικητικού (ΚΔΔ 92)",   "days": 60, "suspended": False, "note": "60 ημέρες από επίδοση απόφασης."},
        {"id": "aithisi_akyr",  "label": "Αίτηση Ακύρωσης (ΠΔ 18/89)",   "days": 60, "suspended": False, "note": "60 ημέρες από δημοσίευση/κοινοποίηση."},
        {"id": "anakopi_fte",   "label": "Ανακοπή κατά ΦΤΕ",             "days": 30, "suspended": False, "note": "30 ημέρες."},
    ],
}


@app.get("/api/calendar/holidays")
async def get_holidays_api(year: int = Query(default=None), user: Dict = Depends(get_current_user)):
    if year is None:
        year = datetime.utcnow().year
    return {
        "year": year,
        "holidays": _greek_holidays_list(year),
        "easter": _orthodox_easter(year).isoformat(),
    }


@app.get("/api/calendar/vacations")
async def get_court_vacations_api(year: int = Query(default=None), user: Dict = Depends(get_current_user)):
    if year is None:
        year = datetime.utcnow().year
    return {
        "year": year,
        "vacations": _court_vacations_list(year),
    }


@app.get("/api/calendar/deadline-types")
async def get_deadline_types(user: Dict = Depends(get_current_user)):
    return DEADLINE_CATALOG


class DeadlineCalcRequest(BaseModel):
    start_date: str          # ISO date string YYYY-MM-DD
    law_code: str            # e.g. "ΚΠολΔ"
    deadline_type_id: str    # e.g. "efesi"
    notes: Optional[str] = None


@app.post("/api/calendar/calculate")
async def calculate_deadline_api(req: DeadlineCalcRequest, user: Dict = Depends(get_current_user)):
    catalog = DEADLINE_CATALOG.get(req.law_code)
    if not catalog:
        raise HTTPException(400, f"Άγνωστος κωδικός νόμου: {req.law_code}")
    dtype = next((d for d in catalog if d["id"] == req.deadline_type_id), None)
    if not dtype:
        raise HTTPException(400, f"Άγνωστος τύπος προθεσμίας: {req.deadline_type_id}")

    result = _compute_deadline(req.start_date, dtype["days"], dtype["suspended"])
    return {
        **result,
        "law_code": req.law_code,
        "deadline_type": dtype["label"],
        "nominal_days": dtype["days"],
        "suspended_during_recesses": dtype["suspended"],
        "legal_note": dtype["note"],
        "user_notes": req.notes or "",
    }


@app.get("/api/calendar/week")
async def get_week_data(date: str = Query(default=None), user: Dict = Depends(get_current_user)):
    """Return all hearings + deadlines for the week containing the given date."""
    if date is None:
        date = datetime.utcnow().date().isoformat()
    try:
        pivot = _date_t.fromisoformat(date)
    except ValueError:
        raise HTTPException(400, "Μη έγκυρη ημερομηνία")

    # Monday of that week
    monday = pivot - timedelta(days=pivot.weekday())
    sunday = monday + timedelta(days=6)

    # Collect hearings
    hearings_cursor = db.hearings.find({
        "hearing_date": {
            "$gte": datetime.combine(monday, datetime.min.time()),
            "$lte": datetime.combine(sunday, datetime.max.time()),
        }
    })
    hearings = []
    async for h in hearings_cursor:
        h["_id"] = str(h["_id"])
        hearings.append(h)

    # Collect deadlines
    deadlines_cursor = db.deadlines.find({
        "$or": [
            {"due_date": {"$gte": monday.isoformat(), "$lte": sunday.isoformat()}},
            {"date": {
                "$gte": datetime.combine(monday, datetime.min.time()),
                "$lte": datetime.combine(sunday, datetime.max.time()),
            }},
        ]
    })
    deadlines = []
    async for d in deadlines_cursor:
        d["_id"] = str(d["_id"])
        deadlines.append(d)

    # Enrich with case titles
    case_ids = list({h.get("case_id") for h in hearings} | {d.get("case_id") for d in deadlines} - {None})
    case_map = {}
    for cid in case_ids:
        try:
            c = await db.cases.find_one({"_id": ObjectId(cid)})
            if c:
                case_map[cid] = c.get("title") or c.get("offense") or "—"
        except Exception:
            pass

    for h in hearings:
        h["case_title"] = case_map.get(h.get("case_id"), "—")
        h["day"] = datetime.fromisoformat(str(h["hearing_date"]).replace("Z","")).date().isoformat() if h.get("hearing_date") else None
    for d in deadlines:
        d["case_title"] = case_map.get(d.get("case_id"), "—")
        d["day"] = d.get("due_date") or (str(d["date"])[:10] if d.get("date") else None)

    return {
        "week_start": monday.isoformat(),
        "week_end": sunday.isoformat(),
        "hearings": hearings,
        "deadlines": deadlines,
        "holidays": _greek_holidays_list(monday.year),
        "vacations": _court_vacations_list(monday.year),
    }


# ═══════════════════════════════════════════════════════════════════════════════
# ΠΙΝΑΚΕΙΑ — Court Schedule Management
# ═══════════════════════════════════════════════════════════════════════════════

@app.post("/api/pinakia/upload")
async def upload_pinakio(
    file: UploadFile = File(...),
    current_user: dict = Depends(get_current_user),
):
    """Upload and process a court schedule document (PDF/DOCX/XLSX/image)."""
    file_bytes = await file.read()
    media_type = file.content_type or "application/pdf"
    filename = file.filename or "pinakio"

    # Extract hearings via Claude
    try:
        extracted = extract_pinakio(ANTHROPIC_API_KEY, file_bytes, media_type, MODEL_CHAT)
    except Exception as e:
        logger.error(f"Pinakio extraction failed: {e}")
        raise HTTPException(status_code=500, detail=f"Αποτυχία εξαγωγής: {str(e)}")

    hearings_raw = extracted.get("hearings") or []
    hearing_date = extracted.get("hearing_date") or datetime.now(timezone.utc).date().isoformat()
    court_name = extracted.get("court_name") or "Άγνωστο Δικαστήριο"

    # Match against open cases
    hearings = await match_pinakio_hearings(db, hearings_raw)
    match_count = sum(1 for h in hearings if h.get("matched_case_id"))

    doc = {
        "court_name": court_name,
        "hearing_date": hearing_date,
        "file_name": filename,
        "media_type": media_type,
        "uploaded_at": datetime.now(timezone.utc),
        "uploaded_by": current_user.get("email", "unknown"),
        "source": "web",
        "hearings": hearings,
        "hearing_count": len(hearings),
        "match_count": match_count,
    }
    result = await db.pinakia.insert_one(doc)
    doc["_id"] = str(result.inserted_id)
    doc["uploaded_at"] = doc["uploaded_at"].isoformat()

    await log_action(current_user["id"], "pinakio_upload", {
        "court": court_name, "date": hearing_date, "hearings": len(hearings), "matches": match_count
    })
    return doc


@app.post("/api/pinakia/intake")
async def intake_pinakio(
    file: UploadFile = File(...),
    source_label: str = "gas",
    x_api_key: str = Header(default="", alias="X-API-Key"),
):
    """Server-to-server pinakio intake (Google Apps Script, etc.). Auth via X-API-Key header."""
    if not INTAKE_API_KEY or x_api_key != INTAKE_API_KEY:
        raise HTTPException(status_code=401, detail="Invalid or missing API key")

    file_bytes = await file.read()
    media_type = file.content_type or "application/pdf"
    filename = file.filename or "pinakio"

    try:
        extracted = extract_pinakio(ANTHROPIC_API_KEY, file_bytes, media_type, MODEL_CHAT)
    except Exception as e:
        logger.error(f"Pinakio intake extraction failed: {e}")
        raise HTTPException(status_code=500, detail=f"Αποτυχία εξαγωγής: {str(e)}")

    hearings_raw = extracted.get("hearings") or []
    hearing_date = extracted.get("hearing_date") or datetime.now(timezone.utc).date().isoformat()
    court_name = extracted.get("court_name") or "Άγνωστο Δικαστήριο"

    hearings = await match_pinakio_hearings(db, hearings_raw)
    match_count = sum(1 for h in hearings if h.get("matched_case_id"))

    doc = {
        "court_name": court_name,
        "hearing_date": hearing_date,
        "file_name": filename,
        "media_type": media_type,
        "uploaded_at": datetime.now(timezone.utc),
        "uploaded_by": source_label,
        "source": source_label,
        "hearings": hearings,
        "hearing_count": len(hearings),
        "match_count": match_count,
    }
    result = await db.pinakia.insert_one(doc)
    doc["_id"] = str(result.inserted_id)
    doc["uploaded_at"] = doc["uploaded_at"].isoformat()

    logger.info(f"Pinakio intake ({source_label}): {court_name} {hearing_date} — {len(hearings)} hearings, {match_count} matches")

    # Notify Telegram
    from email_intake_service import _notify_telegram
    await _notify_telegram(db, doc)

    return {"ok": True, "id": doc["_id"], "court_name": court_name,
            "hearing_date": hearing_date, "hearing_count": len(hearings), "match_count": match_count}


@app.get("/api/pinakia")
async def list_pinakia(
    limit: int = 50,
    current_user: dict = Depends(get_current_user),
):
    """List all court schedules, newest first."""
    docs = await db.pinakia.find(
        {}, {"hearings": 0}  # exclude full hearings list for list view
    ).sort("hearing_date", -1).limit(limit).to_list(limit)
    return [serialize(d) for d in docs]


@app.get("/api/pinakia/hearings")
async def get_hearings_by_date(
    date: str = Query(..., description="YYYY-MM-DD"),
    current_user: dict = Depends(get_current_user),
):
    """Get all hearings for a specific date across all pinakia."""
    docs = await db.pinakia.find({"hearing_date": date}).to_list(20)
    all_hearings = []
    for doc in docs:
        court = doc.get("court_name", "—")
        for h in doc.get("hearings", []):
            all_hearings.append({**h, "court_name": court, "pinakio_id": str(doc["_id"])})
    return {"date": date, "hearings": all_hearings, "count": len(all_hearings)}


@app.get("/api/pinakia/search")
async def search_pinakia(
    q: str = Query(..., min_length=2),
    current_user: dict = Depends(get_current_user),
):
    """Search hearings by party name or case number."""
    docs = await db.pinakia.find({
        "$or": [
            {"hearings.parties": {"$regex": q, "$options": "i"}},
            {"hearings.case_number": {"$regex": q, "$options": "i"}},
            {"court_name": {"$regex": q, "$options": "i"}},
        ]
    }).sort("hearing_date", -1).limit(20).to_list(20)

    results = []
    for doc in docs:
        court = doc.get("court_name", "—")
        date = doc.get("hearing_date", "—")
        for h in doc.get("hearings", []):
            parties_str = " ".join(h.get("parties", []))
            case_num = h.get("case_number", "")
            if (re.search(q, parties_str, re.IGNORECASE) or
                    re.search(q, case_num, re.IGNORECASE)):
                results.append({
                    **h,
                    "court_name": court,
                    "hearing_date": date,
                    "pinakio_id": str(doc["_id"]),
                })
    return {"query": q, "results": results, "count": len(results)}


@app.get("/api/pinakia/{pinakio_id}")
async def get_pinakio(
    pinakio_id: str,
    current_user: dict = Depends(get_current_user),
):
    """Get a single pinakio with all hearings."""
    try:
        oid = ObjectId(pinakio_id)
    except Exception:
        raise HTTPException(status_code=400, detail="Μη έγκυρο ID")
    doc = await db.pinakia.find_one({"_id": oid})
    if not doc:
        raise HTTPException(status_code=404, detail="Πινάκειο δεν βρέθηκε")
    return serialize(doc)


@app.delete("/api/pinakia/{pinakio_id}")
async def delete_pinakio(
    pinakio_id: str,
    current_user: dict = Depends(require_role("admin", "lawyer")),
):
    try:
        oid = ObjectId(pinakio_id)
    except Exception:
        raise HTTPException(status_code=400, detail="Μη έγκυρο ID")
    result = await db.pinakia.delete_one({"_id": oid})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Πινάκειο δεν βρέθηκε")
    return {"ok": True}


# ── Criminal Cases Module ──────────────────────────────────────────────────────

from criminal_case_service import (
    build_case_context, generate_output as cc_generate_output,
    summarize_document as cc_summarize_document,
    compute_health as cc_compute_health, days_until as cc_days_until,
    render_pdf as cc_render_pdf, render_docx as cc_render_docx,
    PROMPTS as CC_PROMPTS,
)

CC_OUTPUT_TITLES = {
    "case_summary": "Περίληψη Υπόθεσης",
    "chronology": "Χρονολόγιο Γεγονότων",
    "missing_documents": "Checklist Ελλειπόντων Εγγράφων",
    "client_questions": "Ερωτήσεις προς Πελάτη",
    "witness_questions": "Ερωτήσεις προς Μάρτυρες",
    "risk_analysis": "Ανάλυση Κινδύνου",
    "court_brief": "Court Preparation Brief",
    "client_email": "Draft Email προς Πελάτη",
    "internal_memo": "Εσωτερικό Memo",
    "defence_strategy": "Στρατηγική Υπεράσπισης (Draft)",
    "prosecution_support": "Υποστήριξη Κατηγορίας (Draft)",
    "legal_issues": "Νομικά Ζητήματα (Draft)",
}

CC_UPLOAD_DIR = Path(os.getenv("DOCUMENT_STORAGE_PATH", "/data/documents")) / "criminal"
CC_UPLOAD_DIR.mkdir(parents=True, exist_ok=True)


async def _cc_get_case_or_404(case_id: str) -> dict:
    case = await db.cc_cases.find_one({"id": case_id}, {"_id": 0})
    if not case:
        raise HTTPException(status_code=404, detail="Ποινική υπόθεση δεν βρέθηκε")
    return case


def _cc_extract_text(file_path: Path) -> str:
    try:
        ext = file_path.suffix.lower()
        if ext == ".pdf":
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(str(file_path))
                return "\n".join((p.extract_text() or "") for p in reader.pages[:50])
            except Exception:
                return ""
        if ext == ".docx":
            try:
                from docx import Document as _DocxDoc
                d = _DocxDoc(str(file_path))
                return "\n".join(p.text for p in d.paragraphs)
            except Exception:
                return ""
        if ext in (".txt", ".md"):
            try:
                return file_path.read_text(encoding="utf-8", errors="ignore")[:50000]
            except Exception:
                return ""
        return ""
    except Exception as e:
        logger.warning(f"cc_extract_text failed: {e}")
        return ""


async def _cc_overdue_task_count(case_id: str) -> int:
    today = datetime.utcnow().date().isoformat()
    return await db.cc_tasks.count_documents({
        "case_id": case_id,
        "status": {"$in": ["open", "in_progress"]},
        "due_date": {"$lt": today, "$ne": None},
    })


async def _cc_unapproved_critical_outputs(case_id: str) -> int:
    return await db.cc_outputs.count_documents({
        "case_id": case_id,
        "output_type": {"$in": ["court_brief", "client_email"]},
        "status": {"$ne": "approved"},
    })


# ── Cases ────────────────────────────────────────────────────────────────────

class CCCaseCreate(BaseModel):
    case_title: str
    client_name: str
    client_email: Optional[str] = None
    client_phone: Optional[str] = None
    client_role: str = "accused"
    opposing_party: Optional[str] = None
    matter_type: str
    court: Optional[str] = None
    hearing_date: Optional[str] = None
    urgency_level: str = "medium"
    status: str = "intake"
    short_description: str


class CCCaseUpdate(BaseModel):
    case_title: Optional[str] = None
    client_name: Optional[str] = None
    client_email: Optional[str] = None
    client_phone: Optional[str] = None
    client_role: Optional[str] = None
    opposing_party: Optional[str] = None
    matter_type: Optional[str] = None
    court: Optional[str] = None
    hearing_date: Optional[str] = None
    urgency_level: Optional[str] = None
    status: Optional[str] = None
    short_description: Optional[str] = None


@app.get("/api/criminal/cases")
async def cc_list_cases(user=Depends(get_current_user)):
    docs = await db.cc_cases.find({}, {"_id": 0}).sort("created_at", -1).to_list(500)
    return docs


@app.post("/api/criminal/cases", status_code=201)
async def cc_create_case(payload: CCCaseCreate, user=Depends(require_role(UserRole.ADMIN, UserRole.LAWYER, UserRole.SECRETARY))):
    if not payload.case_title.strip() or not payload.client_name.strip() \
            or not payload.matter_type.strip() or not payload.short_description.strip():
        raise HTTPException(400, "Λείπουν υποχρεωτικά πεδία")
    now = datetime.utcnow().isoformat()
    case_id = str(uuid.uuid4())
    doc = {
        "id": case_id,
        **payload.model_dump(),
        "case_title": sanitize_string(payload.case_title),
        "client_name": sanitize_string(payload.client_name),
        "matter_type": sanitize_string(payload.matter_type),
        "short_description": sanitize_string(payload.short_description),
        "created_by": user["id"],
        "created_at": now,
        "updated_at": now,
    }
    await db.cc_cases.insert_one(doc)
    doc.pop("_id", None)
    await audit("CC_CASE_CREATED", user["id"], "cc_case", case_id, {"title": doc["case_title"]})
    return doc


@app.get("/api/criminal/cases/{case_id}")
async def cc_get_case(case_id: str, user=Depends(get_current_user)):
    return await _cc_get_case_or_404(case_id)


@app.patch("/api/criminal/cases/{case_id}")
async def cc_update_case(case_id: str, payload: CCCaseUpdate, user=Depends(require_role(UserRole.ADMIN, UserRole.LAWYER, UserRole.SECRETARY))):
    await _cc_get_case_or_404(case_id)
    updates = {k: v for k, v in payload.model_dump().items() if v is not None}
    if "status" in updates and updates["status"] in ("court_preparation", "closed"):
        if user["role"] not in (UserRole.ADMIN.value, UserRole.LAWYER.value):
            raise HTTPException(403, "Μόνο δικηγόρος/admin μπορεί να αλλάξει αυτή την κατάσταση")
    updates["updated_at"] = datetime.utcnow().isoformat()
    await db.cc_cases.update_one({"id": case_id}, {"$set": updates})
    await audit("CC_CASE_UPDATED", user["id"], "cc_case", case_id, list(updates.keys()))
    return await db.cc_cases.find_one({"id": case_id}, {"_id": 0})


@app.delete("/api/criminal/cases/{case_id}")
async def cc_delete_case(case_id: str, user=Depends(require_role(UserRole.ADMIN))):
    await _cc_get_case_or_404(case_id)
    await db.cc_cases.delete_one({"id": case_id})
    for col in ("cc_parties", "cc_events", "cc_documents", "cc_evidence", "cc_issues", "cc_tasks", "cc_outputs"):
        await db[col].delete_many({"case_id": case_id})
    await audit("CC_CASE_DELETED", user["id"], "cc_case", case_id)
    return {"ok": True}


# ── Parties ──────────────────────────────────────────────────────────────────

class CCPartyCreate(BaseModel):
    name: str
    role: str = "other"
    contact_details: Optional[str] = None
    notes: Optional[str] = None


@app.get("/api/criminal/cases/{case_id}/parties")
async def cc_list_parties(case_id: str, user=Depends(get_current_user)):
    await _cc_get_case_or_404(case_id)
    return await db.cc_parties.find({"case_id": case_id}, {"_id": 0}).to_list(500)


@app.post("/api/criminal/cases/{case_id}/parties", status_code=201)
async def cc_create_party(case_id: str, payload: CCPartyCreate, user=Depends(require_role(UserRole.ADMIN, UserRole.LAWYER, UserRole.SECRETARY))):
    await _cc_get_case_or_404(case_id)
    now = datetime.utcnow().isoformat()
    doc = {"id": str(uuid.uuid4()), "case_id": case_id, **payload.model_dump(), "created_at": now}
    await db.cc_parties.insert_one(doc)
    doc.pop("_id", None)
    await audit("CC_PARTY_ADDED", user["id"], "cc_party", doc["id"], {"name": payload.name})
    return doc


@app.delete("/api/criminal/cases/{case_id}/parties/{party_id}")
async def cc_delete_party(case_id: str, party_id: str, user=Depends(require_role(UserRole.ADMIN, UserRole.LAWYER, UserRole.SECRETARY))):
    await db.cc_parties.delete_one({"id": party_id, "case_id": case_id})
    await audit("CC_PARTY_DELETED", user["id"], "cc_party", party_id)
    return {"ok": True}


# ── Events (Timeline) ────────────────────────────────────────────────────────

class CCEventCreate(BaseModel):
    event_date: str
    event_time: Optional[str] = None
    event_description: str
    source: Optional[str] = None
    confidence_level: str = "alleged"
    notes: Optional[str] = None


@app.get("/api/criminal/cases/{case_id}/events")
async def cc_list_events(case_id: str, user=Depends(get_current_user)):
    await _cc_get_case_or_404(case_id)
    return await db.cc_events.find({"case_id": case_id}, {"_id": 0}).sort("event_date", 1).to_list(500)


@app.post("/api/criminal/cases/{case_id}/events", status_code=201)
async def cc_create_event(case_id: str, payload: CCEventCreate, user=Depends(require_role(UserRole.ADMIN, UserRole.LAWYER, UserRole.SECRETARY))):
    await _cc_get_case_or_404(case_id)
    now = datetime.utcnow().isoformat()
    doc = {"id": str(uuid.uuid4()), "case_id": case_id, **payload.model_dump(), "created_at": now}
    await db.cc_events.insert_one(doc)
    doc.pop("_id", None)
    await audit("CC_EVENT_ADDED", user["id"], "cc_event", doc["id"])
    return doc


@app.delete("/api/criminal/cases/{case_id}/events/{event_id}")
async def cc_delete_event(case_id: str, event_id: str, user=Depends(require_role(UserRole.ADMIN, UserRole.LAWYER, UserRole.SECRETARY))):
    await db.cc_events.delete_one({"id": event_id, "case_id": case_id})
    await audit("CC_EVENT_DELETED", user["id"], "cc_event", event_id)
    return {"ok": True}


# ── Documents ────────────────────────────────────────────────────────────────

@app.get("/api/criminal/cases/{case_id}/documents")
async def cc_list_documents(case_id: str, user=Depends(get_current_user)):
    await _cc_get_case_or_404(case_id)
    return await db.cc_documents.find({"case_id": case_id}, {"_id": 0}).sort("upload_date", -1).to_list(500)


@app.post("/api/criminal/cases/{case_id}/documents", status_code=201)
async def cc_upload_document(
    case_id: str,
    file: UploadFile = File(...),
    category: str = Form("other"),
    importance_level: str = Form("medium"),
    user=Depends(require_role(UserRole.ADMIN, UserRole.LAWYER, UserRole.SECRETARY)),
):
    await _cc_get_case_or_404(case_id)
    if not file.filename:
        raise HTTPException(400, "Κενό αρχείο")
    case_dir = CC_UPLOAD_DIR / case_id
    case_dir.mkdir(parents=True, exist_ok=True)

    valid_categories = {"police_report", "witness_statement", "medical_report", "photo", "video", "court_document", "email", "other"}
    valid_importance = {"low", "medium", "high", "critical"}
    doc_id = str(uuid.uuid4())
    now = datetime.utcnow().isoformat()

    target = case_dir / f"{doc_id}__{file.filename}"
    raw = await file.read()
    target.write_bytes(raw)

    text = _cc_extract_text(target)
    summary = None
    if text:
        try:
            summary = await cc_summarize_document(file.filename, text[:50000], "el")
        except Exception as e:
            logger.warning(f"CC AI summary failed: {e}")

    doc = {
        "id": doc_id,
        "case_id": case_id,
        "file_name": file.filename,
        "file_type": file.content_type or "application/octet-stream",
        "category": category if category in valid_categories else "other",
        "importance_level": importance_level if importance_level in valid_importance else "medium",
        "upload_date": now,
        "summary": summary,
        "extracted_text": (text or "")[:50000],
        "stored_path": str(target),
        "size_bytes": len(raw),
    }
    await db.cc_documents.insert_one(doc)
    doc.pop("_id", None)
    await audit("CC_DOCUMENT_UPLOADED", user["id"], "cc_document", doc_id, {"file_name": file.filename})
    return doc


@app.delete("/api/criminal/cases/{case_id}/documents/{doc_id}")
async def cc_delete_document(case_id: str, doc_id: str, user=Depends(require_role(UserRole.ADMIN, UserRole.LAWYER))):
    doc = await db.cc_documents.find_one({"id": doc_id, "case_id": case_id}, {"_id": 0})
    if doc and doc.get("stored_path"):
        try:
            Path(doc["stored_path"]).unlink(missing_ok=True)
        except Exception:
            pass
    await db.cc_documents.delete_one({"id": doc_id, "case_id": case_id})
    await audit("CC_DOCUMENT_DELETED", user["id"], "cc_document", doc_id)
    return {"ok": True}


# ── Evidence ─────────────────────────────────────────────────────────────────

class CCEvidenceCreate(BaseModel):
    title: str
    description: Optional[str] = None
    source: Optional[str] = None
    supports: str = "neutral"
    weakens: Optional[str] = None
    reliability: str = "unverified"


@app.get("/api/criminal/cases/{case_id}/evidence")
async def cc_list_evidence(case_id: str, user=Depends(get_current_user)):
    await _cc_get_case_or_404(case_id)
    return await db.cc_evidence.find({"case_id": case_id}, {"_id": 0}).to_list(500)


@app.post("/api/criminal/cases/{case_id}/evidence", status_code=201)
async def cc_create_evidence(case_id: str, payload: CCEvidenceCreate, user=Depends(require_role(UserRole.ADMIN, UserRole.LAWYER, UserRole.SECRETARY))):
    await _cc_get_case_or_404(case_id)
    now = datetime.utcnow().isoformat()
    doc = {"id": str(uuid.uuid4()), "case_id": case_id, **payload.model_dump(), "created_at": now}
    await db.cc_evidence.insert_one(doc)
    doc.pop("_id", None)
    await audit("CC_EVIDENCE_ADDED", user["id"], "cc_evidence", doc["id"], {"title": payload.title})
    return doc


@app.delete("/api/criminal/cases/{case_id}/evidence/{evidence_id}")
async def cc_delete_evidence(case_id: str, evidence_id: str, user=Depends(require_role(UserRole.ADMIN, UserRole.LAWYER, UserRole.SECRETARY))):
    await db.cc_evidence.delete_one({"id": evidence_id, "case_id": case_id})
    await audit("CC_EVIDENCE_DELETED", user["id"], "cc_evidence", evidence_id)
    return {"ok": True}


# ── Legal Issues ──────────────────────────────────────────────────────────────

class CCIssueCreate(BaseModel):
    issue_title: str
    facts_supporting: Optional[str] = None
    missing_facts: Optional[str] = None
    risk_level: str = "medium"
    lawyer_notes: Optional[str] = None


@app.get("/api/criminal/cases/{case_id}/issues")
async def cc_list_issues(case_id: str, user=Depends(get_current_user)):
    await _cc_get_case_or_404(case_id)
    return await db.cc_issues.find({"case_id": case_id}, {"_id": 0}).to_list(500)


@app.post("/api/criminal/cases/{case_id}/issues", status_code=201)
async def cc_create_issue(case_id: str, payload: CCIssueCreate, user=Depends(require_role(UserRole.ADMIN, UserRole.LAWYER, UserRole.SECRETARY))):
    await _cc_get_case_or_404(case_id)
    now = datetime.utcnow().isoformat()
    doc = {"id": str(uuid.uuid4()), "case_id": case_id, **payload.model_dump(), "created_at": now}
    await db.cc_issues.insert_one(doc)
    doc.pop("_id", None)
    await audit("CC_ISSUE_ADDED", user["id"], "cc_issue", doc["id"], {"title": payload.issue_title})
    return doc


@app.delete("/api/criminal/cases/{case_id}/issues/{issue_id}")
async def cc_delete_issue(case_id: str, issue_id: str, user=Depends(require_role(UserRole.ADMIN, UserRole.LAWYER, UserRole.SECRETARY))):
    await db.cc_issues.delete_one({"id": issue_id, "case_id": case_id})
    await audit("CC_ISSUE_DELETED", user["id"], "cc_issue", issue_id)
    return {"ok": True}


# ── Tasks ────────────────────────────────────────────────────────────────────

class CCTaskCreate(BaseModel):
    title: str
    description: Optional[str] = None
    assigned_to: Optional[str] = None
    due_date: Optional[str] = None
    priority: str = "medium"
    status: str = "open"


class CCTaskUpdate(BaseModel):
    title: Optional[str] = None
    description: Optional[str] = None
    assigned_to: Optional[str] = None
    due_date: Optional[str] = None
    priority: Optional[str] = None
    status: Optional[str] = None


@app.get("/api/criminal/cases/{case_id}/tasks")
async def cc_list_tasks(case_id: str, user=Depends(get_current_user)):
    await _cc_get_case_or_404(case_id)
    return await db.cc_tasks.find({"case_id": case_id}, {"_id": 0}).to_list(500)


@app.post("/api/criminal/cases/{case_id}/tasks", status_code=201)
async def cc_create_task(case_id: str, payload: CCTaskCreate, user=Depends(require_role(UserRole.ADMIN, UserRole.LAWYER, UserRole.SECRETARY))):
    await _cc_get_case_or_404(case_id)
    now = datetime.utcnow().isoformat()
    doc = {"id": str(uuid.uuid4()), "case_id": case_id, **payload.model_dump(), "created_at": now}
    await db.cc_tasks.insert_one(doc)
    doc.pop("_id", None)
    await audit("CC_TASK_ADDED", user["id"], "cc_task", doc["id"], {"title": payload.title})
    return doc


@app.patch("/api/criminal/cases/{case_id}/tasks/{task_id}")
async def cc_update_task(case_id: str, task_id: str, payload: CCTaskUpdate, user=Depends(require_role(UserRole.ADMIN, UserRole.LAWYER, UserRole.SECRETARY))):
    doc = await db.cc_tasks.find_one({"id": task_id, "case_id": case_id}, {"_id": 0})
    if not doc:
        raise HTTPException(404, "Task δεν βρέθηκε")
    updates = {k: v for k, v in payload.model_dump().items() if v is not None}
    if updates:
        await db.cc_tasks.update_one({"id": task_id}, {"$set": updates})
        doc.update(updates)
    await audit("CC_TASK_UPDATED", user["id"], "cc_task", task_id)
    return doc


@app.delete("/api/criminal/cases/{case_id}/tasks/{task_id}")
async def cc_delete_task(case_id: str, task_id: str, user=Depends(require_role(UserRole.ADMIN, UserRole.LAWYER, UserRole.SECRETARY))):
    await db.cc_tasks.delete_one({"id": task_id, "case_id": case_id})
    await audit("CC_TASK_DELETED", user["id"], "cc_task", task_id)
    return {"ok": True}


# ── AI Outputs ────────────────────────────────────────────────────────────────

class CCGenerateRequest(BaseModel):
    output_type: str
    language: str = "el"
    extra_context: Optional[str] = None


class CCOutputUpdate(BaseModel):
    content: Optional[str] = None
    status: Optional[str] = None


@app.get("/api/criminal/cases/{case_id}/outputs")
async def cc_list_outputs(case_id: str, user=Depends(get_current_user)):
    await _cc_get_case_or_404(case_id)
    return await db.cc_outputs.find({"case_id": case_id}, {"_id": 0}).sort("created_at", -1).to_list(500)


@app.post("/api/criminal/cases/{case_id}/outputs/generate", status_code=201)
async def cc_generate_ai_output(case_id: str, payload: CCGenerateRequest, user=Depends(require_role(UserRole.ADMIN, UserRole.LAWYER, UserRole.SECRETARY))):
    case = await _cc_get_case_or_404(case_id)
    parties = await db.cc_parties.find({"case_id": case_id}, {"_id": 0}).to_list(500)
    events = await db.cc_events.find({"case_id": case_id}, {"_id": 0}).sort("event_date", 1).to_list(500)
    documents = await db.cc_documents.find({"case_id": case_id}, {"_id": 0}).to_list(500)
    evidence = await db.cc_evidence.find({"case_id": case_id}, {"_id": 0}).to_list(500)
    issues = await db.cc_issues.find({"case_id": case_id}, {"_id": 0}).to_list(500)

    context = build_case_context(case, parties, events, documents, evidence, issues)
    try:
        content = await cc_generate_output(payload.output_type, payload.language, context, payload.extra_context)
    except ValueError as e:
        raise HTTPException(400, str(e))
    except Exception as e:
        logger.exception("CC AI generation failed")
        raise HTTPException(502, f"AI generation failed: {e}")

    now = datetime.utcnow().isoformat()
    out = {
        "id": str(uuid.uuid4()),
        "case_id": case_id,
        "output_type": payload.output_type,
        "title": CC_OUTPUT_TITLES.get(payload.output_type, payload.output_type),
        "content": content,
        "language": payload.language,
        "status": "draft",
        "created_by": user["id"],
        "approved_by": None,
        "created_at": now,
        "updated_at": now,
    }
    await db.cc_outputs.insert_one(out)
    out.pop("_id", None)
    await audit("CC_AI_GENERATED", user["id"], "cc_output", out["id"], {"output_type": payload.output_type})
    return out


@app.patch("/api/criminal/cases/{case_id}/outputs/{output_id}")
async def cc_update_output(case_id: str, output_id: str, payload: CCOutputUpdate, user=Depends(get_current_user)):
    doc = await db.cc_outputs.find_one({"id": output_id, "case_id": case_id}, {"_id": 0})
    if not doc:
        raise HTTPException(404, "Output δεν βρέθηκε")
    updates: dict = {}
    if payload.content is not None:
        if user["role"] not in (UserRole.ADMIN.value, UserRole.LAWYER.value, UserRole.SECRETARY.value):
            raise HTTPException(403, "Δεν επιτρέπεται επεξεργασία")
        updates["content"] = payload.content
        updates["status"] = "revised"
    if payload.status is not None:
        if payload.status in ("approved", "rejected"):
            if user["role"] not in (UserRole.ADMIN.value, UserRole.LAWYER.value):
                raise HTTPException(403, "Μόνο δικηγόρος/admin μπορεί να εγκρίνει/απορρίψει")
            updates["status"] = payload.status
            updates["approved_by"] = user["id"] if payload.status == "approved" else None
        else:
            updates["status"] = payload.status
    if updates:
        updates["updated_at"] = datetime.utcnow().isoformat()
        await db.cc_outputs.update_one({"id": output_id}, {"$set": updates})
        doc.update(updates)
        await audit(f"CC_OUTPUT_{updates.get('status', 'updated').upper()}", user["id"], "cc_output", output_id)
    return doc


@app.delete("/api/criminal/cases/{case_id}/outputs/{output_id}")
async def cc_delete_output(case_id: str, output_id: str, user=Depends(require_role(UserRole.ADMIN, UserRole.LAWYER))):
    await db.cc_outputs.delete_one({"id": output_id, "case_id": case_id})
    await audit("CC_OUTPUT_DELETED", user["id"], "cc_output", output_id)
    return {"ok": True}


# ── Export ────────────────────────────────────────────────────────────────────

@app.get("/api/criminal/cases/{case_id}/outputs/{output_id}/export")
async def cc_export_output(case_id: str, output_id: str, format: str = "pdf", user=Depends(get_current_user)):
    case = await _cc_get_case_or_404(case_id)
    doc = await db.cc_outputs.find_one({"id": output_id, "case_id": case_id}, {"_id": 0})
    if not doc:
        raise HTTPException(404, "Output δεν βρέθηκε")

    if doc["output_type"] in ("client_email", "court_brief") and doc["status"] != "approved":
        if user["role"] not in (UserRole.ADMIN.value, UserRole.LAWYER.value):
            raise HTTPException(403, "Output δεν έχει εγκριθεί")

    fmt = (format or "pdf").lower()
    if fmt == "pdf":
        data = cc_render_pdf(case["case_title"], doc["title"], doc["content"], doc["status"])
        media = "application/pdf"
        ext = "pdf"
    elif fmt == "docx":
        data = cc_render_docx(case["case_title"], doc["title"], doc["content"], doc["status"])
        media = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ext = "docx"
    else:
        raise HTTPException(400, "Μη υποστηριζόμενη μορφή")

    await audit("CC_OUTPUT_EXPORTED", user["id"], "cc_output", output_id, {"format": ext})
    from urllib.parse import quote
    raw_name = "".join(c for c in doc["title"] if c.isalnum() or c in " -_").strip().replace(" ", "_") or "output"
    ascii_fallback = raw_name.encode("ascii", "ignore").decode("ascii") or doc["output_type"]
    utf8_quoted = quote(raw_name, safe="")
    headers = {
        "Content-Disposition": (
            f'attachment; filename="{ascii_fallback}.{ext}"; '
            f"filename*=UTF-8''{utf8_quoted}.{ext}"
        )
    }
    return StreamingResponse(iter([data]), media_type=media, headers=headers)


# ── Health ────────────────────────────────────────────────────────────────────

@app.get("/api/criminal/cases/{case_id}/health")
async def cc_case_health(case_id: str, user=Depends(get_current_user)):
    case = await _cc_get_case_or_404(case_id)
    doc_count = await db.cc_documents.count_documents({"case_id": case_id})
    overdue_tasks = await _cc_overdue_task_count(case_id)
    unapproved_critical = await _cc_unapproved_critical_outputs(case_id)
    missing_docs = doc_count == 0
    health = cc_compute_health(case, doc_count, overdue_tasks, unapproved_critical, missing_docs)
    health["hearing_days_left"] = cc_days_until(case.get("hearing_date"))
    return health


# ── Audit log per criminal case ───────────────────────────────────────────────

@app.get("/api/criminal/cases/{case_id}/audit")
async def cc_case_audit(case_id: str, user=Depends(get_current_user)):
    await _cc_get_case_or_404(case_id)
    docs = await db.audit_logs.find(
        {"resource": "cc_case", "resource_id": case_id},
        {"_id": 0},
    ).sort("timestamp", -1).to_list(200)
    return [serialize(d) for d in docs]


# ── VAPID Web Push Endpoints ──────────────────────────────────────────────────

class PushSubscribeRequest(BaseModel):
    endpoint: str
    auth: str
    p256dh: str
    user_agent: Optional[str] = None

@app.get("/api/v1/push/vapid-public-key")
async def get_vapid_public_key():
    """Return VAPID public key for browser push subscription."""
    key = os.getenv("VAPID_PUBLIC_KEY", "")
    if not key:
        raise HTTPException(503, "Push notifications not configured")
    return {"public_key": key}

@app.post("/api/v1/push/subscribe")
async def subscribe_push(req: PushSubscribeRequest, request: Request, user=Depends(get_current_user)):
    """Store a Web Push subscription for the current user."""
    if not req.endpoint or not req.auth or not req.p256dh:
        raise HTTPException(400, "Missing subscription fields")

    await db.push_subscriptions.update_one(
        {"user_id": user["id"], "endpoint": req.endpoint},
        {"$set": {
            "user_id": user["id"],
            "endpoint": req.endpoint,
            "auth": req.auth,
            "p256dh": req.p256dh,
            "user_agent": req.user_agent or request.headers.get("user-agent", ""),
            "updated_at": datetime.utcnow(),
        }, "$setOnInsert": {"created_at": datetime.utcnow()}},
        upsert=True,
    )
    return {"status": "subscribed"}

@app.delete("/api/v1/push/unsubscribe")
async def unsubscribe_push(endpoint: str, user=Depends(get_current_user)):
    """Remove a push subscription."""
    await db.push_subscriptions.delete_one({"user_id": user["id"], "endpoint": endpoint})
    return {"status": "unsubscribed"}

@app.post("/api/v1/push/test")
async def test_push(user=Depends(get_current_user)):
    """Send a test push notification to the current user."""
    svc = get_push_service(db)
    result = await svc.send_to_user(
        user["id"],
        title="Nomos One — Δοκιμή",
        body="Οι push notifications λειτουργούν κανονικά!",
        path="/",
    )
    return result
